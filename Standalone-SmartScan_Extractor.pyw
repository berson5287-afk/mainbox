#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SmartScan Extractor v2.0
Standalone RFQ / material scanner designed to later plug into Mainbox.

Goals:
- Fast native extraction first.
- Optional OCR engines when installed.
- Slow/Normal/Fast scan modes.
- Highly customizable settings.
- Background scanning so UI stays responsive.
- SQLite cache and correction learning.
- Export to CSV / JSON.

Optional dependencies supported automatically:
    pip install pymupdf pillow pytesseract opencv-python
    pip install paddleocr
    pip install surya-ocr   # if available in your environment

External programs optionally supported:
    tesseract.exe on PATH, or configured in Settings.

This file intentionally avoids requiring optional OCR packages just to start.

v2.0 additions:
- Anthropic Claude per-item real-time review (streams green/yellow/red as scan runs)
- Persistent source-text colour highlights for every identified line
- original_* fields on MaterialLine capture scanner's first guess for DB learning
- Corrections DB stores both the original scan guess AND the user's corrected values
- SettingsWindow has masked API-key entry with show/hide toggle
"""

from __future__ import annotations

import csv
import dataclasses
import hashlib
import json
import os
import queue
import re
import sqlite3
import subprocess
import shutil
import sys
import tempfile
import threading
import time
import traceback
import urllib.request
from difflib import SequenceMatcher
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

import tkinter as tk
from tkinter import filedialog, messagebox, ttk

# Optional imports -----------------------------------------------------------
try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover
    fitz = None

try:
    from PIL import Image, ImageEnhance, ImageFilter, ImageOps, ImageTk, ImageDraw, ImageFont
except Exception:  # pragma: no cover
    Image = None
    ImageEnhance = None
    ImageFilter = None
    ImageOps = None
    ImageTk = None
    ImageDraw = None
    ImageFont = None

try:
    import pytesseract
except Exception:  # pragma: no cover
    pytesseract = None

try:
    import cv2
    import numpy as np
except Exception:  # pragma: no cover
    cv2 = None
    np = None

try:
    from paddleocr import PaddleOCR  # type: ignore
except Exception:  # pragma: no cover
    PaddleOCR = None

# Office document readers (optional). Spreadsheets and Word docs are common RFQ
# formats; if these libraries are missing we degrade gracefully instead of
# rejecting the file. Install with: pip install openpyxl python-docx xlrd
try:
    import openpyxl  # .xlsx / .xlsm
except Exception:  # pragma: no cover
    openpyxl = None

try:
    import docx  # python-docx for .docx
except Exception:  # pragma: no cover
    docx = None

try:
    import xlrd  # legacy .xls
except Exception:  # pragma: no cover
    xlrd = None


def _smartscan_resolve_tesseract_cmd(settings=None) -> str:
    """Find tesseract.exe without relying only on Windows PATH.

    MaINbox/SmartScan is often launched from an older Windows process/shortcut
    that does not inherit a freshly edited PATH.  This checks Settings, env vars,
    PATH, normal Windows install folders, and portable copies next to this file.
    """
    candidates: List[str] = []

    try:
        configured = (getattr(settings, "tesseract_cmd", "") or "").strip().strip('"') if settings is not None else ""
        if configured:
            candidates.append(configured)
    except Exception:
        pass

    env_cmd = (os.environ.get("TESSERACT_CMD") or "").strip().strip('"')
    if env_cmd:
        candidates.append(env_cmd)

    try:
        which_cmd = shutil.which("tesseract") or shutil.which("tesseract.exe")
        if which_cmd:
            candidates.append(which_cmd)
    except Exception:
        pass

    program_files = [
        os.environ.get("ProgramFiles", r"C:\Program Files"),
        os.environ.get("ProgramFiles(x86)", r"C:\Program Files (x86)"),
        os.environ.get("LocalAppData", ""),
    ]
    for base in program_files:
        if not base:
            continue
        candidates.extend([
            str(Path(base) / "Tesseract-OCR" / "tesseract.exe"),
            str(Path(base) / "Programs" / "Tesseract-OCR" / "tesseract.exe"),
        ])

    try:
        app_dir = Path(__file__).resolve().parent
        candidates.append(str(app_dir / "Tesseract-OCR" / "tesseract.exe"))
        candidates.append(str(app_dir / "tesseract.exe"))
    except Exception:
        pass

    seen = set()
    for cand in candidates:
        if not cand or cand in seen:
            continue
        seen.add(cand)
        try:
            cpath = Path(cand)
            if cpath.exists() and cpath.is_file():
                return str(cpath)
        except Exception:
            continue
    return ""


def _smartscan_configure_pytesseract(settings=None) -> str:
    """Point pytesseract at the discovered executable and return the path used."""
    if pytesseract is None:
        return ""
    cmd = _smartscan_resolve_tesseract_cmd(settings)
    if cmd:
        try:
            pytesseract.pytesseract.tesseract_cmd = cmd
        except Exception:
            pass
        try:
            if settings is not None and not (getattr(settings, "tesseract_cmd", "") or "").strip():
                settings.tesseract_cmd = cmd
        except Exception:
            pass
    return cmd

# tkinterdnd2 is optional. The app works without it.
try:
    from tkinterdnd2 import DND_FILES, TkinterDnD  # type: ignore
except Exception:  # pragma: no cover
    DND_FILES = None
    TkinterDnD = None

APP_NAME = "SmartScan Extractor"
APP_VERSION = "2.8.6"
DEFAULT_DB_NAME = "smartscan_data.db"
SUPPORTED_TEXT_EXTS = {".txt", ".csv", ".json", ".xml", ".html", ".htm", ".eml"}
SUPPORTED_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}
SUPPORTED_SPREADSHEET_EXTS = {".xlsx", ".xlsm", ".xls", ".csv", ".tsv"}
SUPPORTED_WORD_EXTS = {".docx"}
# Safety cap so a giant spreadsheet cannot lock up extraction.
SPREADSHEET_MAX_ROWS_PER_SHEET = 5000
SUPPORTED_DOC_EXTS = (
    SUPPORTED_TEXT_EXTS | SUPPORTED_IMAGE_EXTS
    | SUPPORTED_SPREADSHEET_EXTS | SUPPORTED_WORD_EXTS | {".pdf"}
)

ELECTRICAL_UNITS = {
    "ea", "each", "pc", "pcs", "piece", "pieces", "ft", "feet", "foot", "roll", "rolls",
    "box", "boxes", "bag", "bags", "case", "cases", "length", "lengths", "stick", "sticks",
    "set", "sets", "lot", "lots", "pair", "pairs", "reel", "reels", "spool", "spools",
    "coil", "coils", "bucket", "buckets", "dozen", "dozens",
    "m", "meter", "meters", "yd", "yard", "yards", "lb", "lbs", "pack", "packs", "pkg", "pkgs",
}

ELECTRICAL_KEYWORDS = [
    "emt", "imc", "rigid", "pvc", "conduit", "mc cable", "bx", "romex", "thhn", "xhhw",
    "wire", "cable", "breaker", "panel", "disconnect", "switch", "receptacle", "gfi", "gfci",
    "cover", "plate", "box", "handy box", "4s", "1900", "connector", "coupling", "strap",
    "bushing", "locknut", "elbow", "lb", "pull box", "junction", "fixture", "led", "driver",
    # v3.8.09: lighting fixtures/units are electrical material (e.g. emergency light unit,
    # luminaire, ballast) - keeps them from a spurious "not electrical" review warning.
    "light", "lighting", "lamp", "luminaire", "ballast",
    "transformer", "contactor", "relay", "fuse", "fused", "unistrut", "strut", "ground", "bond",
    "lug", "meter", "ct", "photocell", "occupancy", "dimmer", "toggle", "decora", "gang",
    # SmartScan v3.8.1: include common PO hardware/support terms seen in clean text PDFs.
    "rod", "threaded rod", "kindorf", "washer", "lock washer", "nut", "anchor", "drill bit",
    "hammer drill", "tool", "bolt", "screw", "sds",
]

MANUFACTURER_HINTS = [
    "leviton", "lithonia", "cooper", "eaton", "siemens", "square d", "schneider", "ge",
    "hubbell", "wiremold", "legrand", "appleton", "crouse", "halco", "rab", "nora",
    "lutron", "ideal", "greenlee", "arlington", "bridgeport", "southwire", "cerrowire",
]


def best_manufacturer_hint(text: str) -> str:
    """Return a conservative manufacturer hint from a description.

    v2.8.1 added OCR quantity-list recovery and called this helper, but the
    helper was not present in the file, causing a runtime NameError on some
    OCR scans. Keep this global helper intentionally small and side-effect free
    so it matches the existing parser behavior without touching other logic.
    """
    low = (text or "").lower()
    for mf in MANUFACTURER_HINTS:
        if re.search(r"\b" + re.escape(mf) + r"\b", low):
            return mf.title()
    return ""

PART_NO_RE = re.compile(r"\b[A-Z0-9][A-Z0-9._\-/]{3,}\b", re.I)
SIZE_RE = re.compile(
    r"(?P<size>(?:\d+\s*/\s*\d+|\d+(?:\.\d+)?)(?:\s*[\"'])?|#\s?\d+|\d+\s?(?:awg|mcm|kcmil)|\d+\s?[xX]\s?\d+)"
)
QTY_LEADING_RE = re.compile(
    r"^\s*(?:qty\s*[:\-]?\s*)?(?P<qty>\(?\d+(?:\.\d+)?\)?)(?:\s*(?P<unit>ea|each|pcs?|ft|feet|foot|rolls?|boxes|box|bags?|cases?|lengths?|sticks?|sets?|lots?|pairs?|reels?|spools?))?\s*(?:[-–—xX:]\s*)?(?P<desc>.+)$",
    re.I,
)
QTY_TRAILING_RE = re.compile(
    r"^(?P<desc>.+?)\s+(?:qty\s*[:\-]?\s*)?(?P<qty>\d+(?:\.\d+)?)(?:\s*(?P<unit>ea|each|pcs?|ft|feet|foot|rolls?|boxes|box|bags?|cases?|lengths?|sticks?|sets?|lots?|pairs?|reels?|spools?))?\s*$",
    re.I,
)
BULLET_RE = re.compile(r"^\s*(?:[-*•●▪]|\d+[.)])\s+")

# Contact/address exclusion. These lines often appear near RFQs and can trick a
# quantity-first parser because phone numbers and street numbers look like QTYs.
PHONE_RE = re.compile(
    r"(?<!\w)(?:\+?1[\s\-.]*)?(?:\(?\d{3}\)?[\s\-.]*)\d{3}[\s\-.]*\d{4}(?:\s*(?:x|ext\.?|extension)\s*\d+)?(?!\w)",
    re.I,
)
EMAIL_RE = re.compile(r"\b[A-Z0-9._%+\-]+@[A-Z0-9.\-]+\.[A-Z]{2,}\b", re.I)
URL_RE = re.compile(r"\b(?:https?://|www\.)\S+", re.I)
ZIP_STATE_RE = re.compile(r"\b[A-Z]{2}\s+\d{5}(?:-\d{4})?\b", re.I)
STREET_SUFFIX_RE = re.compile(
    r"\b(?:street|st\.?|avenue|ave\.?|road|rd\.?|drive|dr\.?|lane|ln\.?|boulevard|blvd\.?|court|ct\.?|place|pl\.?|terrace|ter\.?|way|parkway|pkwy\.?|circle|cir\.?|highway|hwy\.?)\b",
    re.I,
)
ADDRESS_LABEL_RE = re.compile(
    r"^\s*(?:ship\s*to|bill\s*to|sold\s*to|deliver\s*to|delivery|job\s*site|address|location|contact|phone|fax|tel|email|e-mail|from|to)\s*[:\-]",
    re.I,
)

# Strong business-document/header/footer noise. These are common on quotes/PDFs and
# often contain numbers that look like quantities, but they are not material rows.
DOC_ADMIN_RE = re.compile(
    r"\b(?:order\s*date|order\s*number|customer\s*(?:number|no|#|po)|cust\.?\s*(?:no|#)|"
    r"purchase\s*order|po\s*(?:number|no|#)|sold\s*to|ship\s*to|bill\s*to|remit\s*to|"
    r"writer|ship\s*via|page\s*\d+|quote\s*(?:number|no|#)|invoice\s*(?:number|no|#)|"
    r"account\s*(?:number|no|#)|terms|salesperson|prepared\s*by|requested\s*by|"
    r"job\s*(?:name|number|no|#)|project\s*(?:name|number|no|#)|phone|fax|tel|email|e-mail)\b",
    re.I,
)
CUSTOMER_INFO_RE = re.compile(
    r"\b(?:customer|writer|ship via|sold to|ship to|bill to|address|phone|fax|email|order date|order number)\b",
    re.I,
)
PRICE_TAIL_RE = re.compile(
    r"\s+(?:\$?\d{1,3}(?:,\d{3})*(?:\.\d{2,4})\s*/\s*(?:ea|each|c|m|ft|lf|pc|pcs|box|roll)\s+\$?\d{1,3}(?:,\d{3})*(?:\.\d{2,4})|"
    r"\$?\d{1,3}(?:,\d{3})*(?:\.\d{2,4})\s+\$?\d{1,3}(?:,\d{3})*(?:\.\d{2,4}))\s*$",
    re.I,
)
HEADER_QTY_RE = re.compile(r"\b(?:qty|quantity|qnty|q'ty|count|amount)\b", re.I)
HEADER_DESC_RE = re.compile(r"\b(?:description|desc|material|materials|item|product|catalog|cat\.?\s*no\.?|part\s*(?:#|no|number)?)\b", re.I)
HEADER_UNIT_RE = re.compile(r"\b(?:unit|uom|um)\b", re.I)
HEADER_MFG_RE = re.compile(r"\b(?:mfg|manufacturer|brand|make)\b", re.I)
PRICE_HEADER_RE = re.compile(r"\b(?:price|unit\s*price|extension|ext\.?\s*price|total|subtotal|tax|amount)\b", re.I)
LINE_NO_HEADER_RE = re.compile(r"^\s*(?:#|line|ln|row|no\.?|item\s*(?:#|no\.?)?)\s*$", re.I)
MONEY_RE = re.compile(r"(?:\$\s*)?\b\d{1,3}(?:,\d{3})*(?:\.\d{2})\b")
ROWNUM_QTY_UNIT_RE = re.compile(
    r"^\s*(?P<rownum>\d{1,4})\s+(?P<qty>\d+(?:\.\d+)?)\s+(?P<unit>ea|each|pcs?|pc|ft|feet|foot|rolls?|boxes|box|bags?|cases?|lengths?|sticks?|sets?|lots?|pairs?|reels?|spools?|m|meters?|yd|yards?|lbs?|packs?|pkgs?|dozens?)\b\s*(?P<rest>.+)$",
    re.I,
)


def now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")


def app_data_dir() -> Path:
    base = os.environ.get("APPDATA") or os.environ.get("LOCALAPPDATA") or str(Path.home())
    p = Path(base) / "SmartScanExtractor"
    p.mkdir(parents=True, exist_ok=True)
    return p


def safe_read_text(path: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=enc, errors="replace")
        except Exception:
            continue
    return ""




def _smartscan_mainbox_theme_config() -> Dict[str, Any]:
    """Read optional MaINbox theme settings passed through the environment."""
    raw = os.environ.get("MAINBOX_SMARTSCAN_THEME_JSON", "")
    cfg: Dict[str, Any] = {}
    if raw:
        try:
            loaded = json.loads(raw)
            if isinstance(loaded, dict):
                cfg.update(loaded)
        except Exception:
            cfg = {}
    theme = str(cfg.get("theme", "dark") or "dark").lower()
    if theme == "light":
        defaults = {
            "bg": "#f4f6fa", "bg2": "#e9eef6", "bg3": "#ffffff", "accent": "#1f6fd1",
            "accent2": "#d7e7fb", "button_bg": "#dbe7f6", "button_fg": "#102033",
            "fg": "#182230", "fg_dim": "#4d5f78", "sel_bg": "#b9d7ff", "sel_fg": "#0d1b2a",
            "entry_bg": "#ffffff", "font_family": "Segoe UI", "font_size": 9, "font_weight": "normal",
        }
    else:
        defaults = {
            "bg": "#1e2229", "bg2": "#252b34", "bg3": "#2c3340", "accent": "#4a9eff",
            "accent2": "#3a7fd5", "button_bg": "#2e3a4e", "button_fg": "#d6e4f7",
            "fg": "#dde6f0", "fg_dim": "#8fa0b8", "sel_bg": "#3a5a8a", "sel_fg": "#ffffff",
            "entry_bg": "#252b34", "font_family": "Segoe UI", "font_size": 9, "font_weight": "normal",
        }
    defaults.update({k: v for k, v in cfg.items() if v not in [None, ""]})
    try:
        defaults["font_size"] = max(7, min(18, int(defaults.get("font_size", 9))))
    except Exception:
        defaults["font_size"] = 9
    defaults["font_weight"] = "bold" if str(defaults.get("font_weight", "normal")).lower() == "bold" or bool(defaults.get("font_bold", False)) else "normal"
    return defaults


def _apply_mainbox_theme_to_tk_widget(widget: Any, cfg: Dict[str, Any]) -> None:
    """Make SmartScan's native Tk widgets match the MaINbox visual settings."""
    try:
        cls = widget.winfo_class()
    except Exception:
        cls = ""
    common_bg = cfg.get("bg2", "#252b34")
    text_bg = cfg.get("entry_bg", cfg.get("bg3", "#2c3340"))
    fg = cfg.get("fg", "#dde6f0")
    sel_bg = cfg.get("sel_bg", "#3a5a8a")
    sel_fg = cfg.get("sel_fg", "#ffffff")
    font_tuple = (cfg.get("font_family", "Segoe UI"), int(cfg.get("font_size", 9)), cfg.get("font_weight", "normal"))
    try:
        if cls in ["Tk", "Toplevel", "Frame", "Labelframe"]:
            widget.configure(bg=common_bg)
        elif cls in ["Label"]:
            widget.configure(bg=common_bg, fg=fg, font=font_tuple)
        elif cls in ["Button"]:
            widget.configure(bg=cfg.get("button_bg", common_bg), fg=cfg.get("button_fg", fg), activebackground=cfg.get("accent2", common_bg), activeforeground=fg, font=font_tuple, relief="flat")
        elif cls in ["Entry", "Text", "Listbox"]:
            widget.configure(bg=text_bg, fg=fg, insertbackground=fg, selectbackground=sel_bg, selectforeground=sel_fg, font=font_tuple)
        elif cls in ["Canvas"]:
            # Keep the document preview itself light when it is displaying a scanned page, but darken empty canvas chrome.
            widget.configure(highlightbackground=cfg.get("bg3", common_bg))
    except Exception:
        pass
    try:
        for child in widget.winfo_children():
            _apply_mainbox_theme_to_tk_widget(child, cfg)
    except Exception:
        pass

def normalize_space(s: str) -> str:
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"[\t\x0b\x0c]+", " ", s)
    s = re.sub(r"[ \u00a0]{2,}", " ", s)
    s = re.sub(r"\n{4,}", "\n\n\n", s)
    return s.strip()


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def split_lines(text: str) -> List[str]:
    return [ln.strip() for ln in normalize_space(text).splitlines() if ln.strip()]


def split_lines_preserve(text: str) -> List[str]:
    """Split lines but preserve internal spacing for header/column detection."""
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    return [ln.strip() for ln in text.splitlines() if ln.strip()]


@dataclass
class ScanSettings:
    mode: str = "Normal"  # Fast, Normal, Slow Careful
    enable_native_pdf: bool = True
    enable_text_files: bool = True
    enable_tesseract: bool = True
    enable_paddleocr: bool = False
    enable_surya_cli: bool = False
    enable_preprocess: bool = True
    enable_table_guess: bool = True
    enable_header_aware_extraction: bool = True
    exclude_contact_info: bool = True
    strict_address_filter: bool = True
    enable_ai_cleanup_placeholder: bool = False
    enable_cache: bool = True
    learn_from_corrections: bool = True
    auto_deep_scan_low_confidence: bool = True
    # AI review is intentionally non-destructive: it marks rows as OK/warning/error
    # and suggests fixes, but the user chooses what to accept.
    auto_ai_review_after_scan: bool = True
    # Smart review runs during extraction, before the table is shown, so the user
    # immediately sees green/yellow/red evidence on the original document.
    smart_review_during_scan: bool = True
    hide_definite_noise_rows: bool = True
    show_all_review_highlights: bool = True
    enable_rule_review: bool = True
    enable_ollama_ai_review: bool = False
    ai_review_mode: str = "Normal"  # Fast = rules only, Normal = suspicious rows, Slow Careful = every row
    ai_review_only_suspicious: bool = True
    ai_review_ollama_model: str = "qwen2.5:14b-instruct-q8_0"
    ai_review_ollama_url: str = "http://127.0.0.1:11434/api/generate"
    # Anthropic Claude API key for per-item real-time review (v2.0)
    # Leave blank to use ANTHROPIC_API_KEY env var, or disable below.
    anthropic_api_key: str = ""
    enable_anthropic_review: bool = False  # set True to activate
    confidence_threshold: int = 70
    max_pages_fast: int = 8
    max_pages_normal: int = 25
    max_pages_careful: int = 100
    render_dpi_fast: int = 160
    render_dpi_normal: int = 220
    render_dpi_careful: int = 300
    tesseract_cmd: str = ""
    tesseract_lang: str = "eng"
    surya_command: str = "surya_ocr"
    paddle_lang: str = "en"
    output_dir: str = ""
    # --- v2.8.5: catalog grounding + optional Ollama vision OCR ---
    # Grounding is safe + on by default: if the catalog DB is missing it is a
    # silent no-op. Ollama vision is opt-in (off) and falls back to Tesseract.
    enable_catalog_grounding: bool = True
    catalog_db_path: str = ""  # blank -> env MAINBOX_CATALOG_DB, else file next to the app
    enable_ollama_vision: bool = False
    ollama_vision_model: str = "llama3.2-vision"
    ollama_vision_url: str = "http://127.0.0.1:11434/api/generate"

    def as_dict(self) -> Dict[str, Any]:
        return dataclasses.asdict(self)

    @classmethod
    def from_dict(cls, d: Dict[str, Any]) -> "ScanSettings":
        base = cls()
        for k, v in d.items():
            if hasattr(base, k):
                setattr(base, k, v)
        return base

    def max_pages(self) -> int:
        if self.mode == "Fast":
            return int(self.max_pages_fast)
        if self.mode == "Slow Careful":
            return int(self.max_pages_careful)
        return int(self.max_pages_normal)

    def render_dpi(self) -> int:
        if self.mode == "Fast":
            return int(self.render_dpi_fast)
        if self.mode == "Slow Careful":
            return int(self.render_dpi_careful)
        return int(self.render_dpi_normal)


@dataclass
class MaterialLine:
    qty: str = ""
    unit: str = ""
    description: str = ""
    part_number: str = ""
    manufacturer: str = ""
    notes: str = ""
    confidence: int = 0
    source_line: str = ""
    # Evidence fields let the UI jump back to the source text/file preview.
    source_start: int = -1
    source_end: int = -1
    source_page: int = 0  # 1-based page number for PDFs/images when known
    source_bbox: str = ""  # x0,y0,x1,y1 in PDF/page coordinates or rendered image coordinates when known
    source_bbox_kind: str = "pdf"  # pdf = PDF page points; image = rendered/preview image pixels
    source_view: str = "merged"
    scan_index: int = 0  # 1-based order found in the document/text scan
    review_status: str = "pending"  # pending, confirmed, rejected, corrected
    ai_status: str = ""  # ok, warning, error, suggestion
    ai_message: str = ""
    ai_suggest_qty: str = ""
    ai_suggest_unit: str = ""
    ai_suggest_description: str = ""
    ai_suggest_part_number: str = ""
    ai_suggest_manufacturer: str = ""
    user_decision: str = ""  # accepted_ai, kept_original, manual, rejected
    # v2.0: snapshot of what the scanner first guessed (before any user correction)
    original_qty: str = ""
    original_unit: str = ""
    original_description: str = ""
    original_part_number: str = ""
    original_manufacturer: str = ""
    # v2.0: highlight colour assigned by AI review
    highlight_color: str = ""  # "green" | "yellow" | "red" | ""
    # v2.6: extraction-watchdog/source-ledger fields
    source_ledger_index: int = 0
    watchdog_status: str = ""
    watchdog_message: str = ""
    # v2.7: stable review/document order, frozen once after the scan completes.  The
    # review list must NOT reshuffle just because a row later gains a bbox from a
    # highlight click, so ordering keys off this fixed value rather than bbox presence.
    review_order: int = 0

    def snapshot_original(self) -> None:
        """Freeze scanner-guess values before the user can edit them."""
        if not self.original_description and not self.original_qty:
            self.original_qty = self.qty
            self.original_unit = self.unit
            self.original_description = self.description
            self.original_part_number = self.part_number
            self.original_manufacturer = self.manufacturer

    def as_dict(self) -> Dict[str, Any]:
        return dataclasses.asdict(self)


def material_from_dict(d: Dict[str, Any]) -> MaterialLine:
    allowed = {f.name for f in dataclasses.fields(MaterialLine)}
    return MaterialLine(**{k: v for k, v in dict(d).items() if k in allowed})


def normalize_source_signature(text: str) -> str:
    """Stable fuzzy key for learned source-line parse patterns."""
    s = (text or "").lower()
    s = re.sub(r"\d+(?:\.\d+)?", "#", s)
    s = re.sub(r"[^a-z0-9#/'\" -]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s[:220]


@dataclass
class ScanResult:
    file_path: str
    file_hash: str
    engine_summary: str = ""
    native_text: str = ""
    ocr_text: str = ""
    table_text: str = ""
    merged_text: str = ""
    materials: List[MaterialLine] = field(default_factory=list)
    confidence: int = 0
    warnings: List[str] = field(default_factory=list)
    elapsed_sec: float = 0.0
    created_at: str = field(default_factory=now_iso)

    def as_dict(self) -> Dict[str, Any]:
        d = dataclasses.asdict(self)
        d["materials"] = [m.as_dict() for m in self.materials]
        return d


class SmartScanDB:
    def __init__(self, db_path: Optional[Path] = None) -> None:
        self.db_path = db_path or (app_data_dir() / DEFAULT_DB_NAME)
        self._lock = threading.RLock()
        self._init_db()

    def _connect(self) -> sqlite3.Connection:
        con = sqlite3.connect(str(self.db_path), timeout=30)
        con.row_factory = sqlite3.Row
        return con

    def _init_db(self) -> None:
        with self._lock, self._connect() as con:
            con.execute(
                """
                CREATE TABLE IF NOT EXISTS settings (
                    key TEXT PRIMARY KEY,
                    value TEXT NOT NULL
                )
                """
            )
            con.execute(
                """
                CREATE TABLE IF NOT EXISTS scans (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_hash TEXT NOT NULL,
                    file_path TEXT NOT NULL,
                    file_name TEXT NOT NULL,
                    engine_summary TEXT,
                    native_text TEXT,
                    ocr_text TEXT,
                    table_text TEXT,
                    merged_text TEXT,
                    materials_json TEXT,
                    confidence INTEGER,
                    warnings_json TEXT,
                    elapsed_sec REAL,
                    created_at TEXT NOT NULL
                )
                """
            )
            con.execute("CREATE INDEX IF NOT EXISTS idx_scans_hash ON scans(file_hash)")
            con.execute(
                """
                CREATE TABLE IF NOT EXISTS corrections (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_hash TEXT,
                    source_line TEXT,
                    old_json TEXT,
                    original_scan_json TEXT,
                    corrected_json TEXT,
                    user_decision TEXT,
                    ai_status TEXT,
                    sender_hint TEXT,
                    created_at TEXT NOT NULL
                )
                """
            )
            # v2.2: fast lookup table for user-drawn bbox corrections.
            # On future scans of the same file the scanner reads these first.
            con.execute(
                """
                CREATE TABLE IF NOT EXISTS bbox_corrections (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_hash TEXT NOT NULL,
                    source_line TEXT,
                    description TEXT,
                    source_page INTEGER,
                    source_bbox TEXT NOT NULL,
                    source_bbox_kind TEXT DEFAULT 'image',
                    created_at TEXT NOT NULL
                )
                """
            )
            con.execute(
                """
                CREATE TABLE IF NOT EXISTS correction_patterns (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_hash TEXT NOT NULL,
                    source_signature TEXT NOT NULL,
                    source_line TEXT NOT NULL,
                    corrected_qty TEXT,
                    corrected_unit TEXT,
                    corrected_description TEXT,
                    corrected_part_number TEXT,
                    corrected_manufacturer TEXT,
                    created_at TEXT NOT NULL
                )
                """
            )
            con.execute("CREATE INDEX IF NOT EXISTS idx_bbox_hash ON bbox_corrections(file_hash)")
            con.commit()

    def load_settings(self) -> ScanSettings:
        with self._lock, self._connect() as con:
            row = con.execute("SELECT value FROM settings WHERE key='scan_settings'").fetchone()
            if not row:
                return ScanSettings()
            try:
                return ScanSettings.from_dict(json.loads(row["value"]))
            except Exception:
                return ScanSettings()

    def save_settings(self, settings: ScanSettings) -> None:
        with self._lock, self._connect() as con:
            con.execute(
                "INSERT OR REPLACE INTO settings(key, value) VALUES('scan_settings', ?)",
                (json.dumps(settings.as_dict(), indent=2),),
            )
            con.commit()

    def latest_cached_scan(self, file_hash: str) -> Optional[ScanResult]:
        with self._lock, self._connect() as con:
            row = con.execute(
                "SELECT * FROM scans WHERE file_hash=? ORDER BY id DESC LIMIT 1", (file_hash,)
            ).fetchone()
            if not row:
                return None
            try:
                materials = [material_from_dict(m) for m in json.loads(row["materials_json"] or "[]")]
                return ScanResult(
                    file_path=row["file_path"],
                    file_hash=row["file_hash"],
                    engine_summary=row["engine_summary"] or "cached",
                    native_text=row["native_text"] or "",
                    ocr_text=row["ocr_text"] or "",
                    table_text=row["table_text"] or "",
                    merged_text=row["merged_text"] or "",
                    materials=materials,
                    confidence=int(row["confidence"] or 0),
                    warnings=json.loads(row["warnings_json"] or "[]"),
                    elapsed_sec=float(row["elapsed_sec"] or 0.0),
                    created_at=row["created_at"],
                )
            except Exception:
                return None

    def save_scan(self, result: ScanResult) -> None:
        with self._lock, self._connect() as con:
            con.execute(
                """
                INSERT INTO scans(
                    file_hash, file_path, file_name, engine_summary,
                    native_text, ocr_text, table_text, merged_text,
                    materials_json, confidence, warnings_json, elapsed_sec, created_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    result.file_hash,
                    result.file_path,
                    Path(result.file_path).name,
                    result.engine_summary,
                    result.native_text,
                    result.ocr_text,
                    result.table_text,
                    result.merged_text,
                    json.dumps([m.as_dict() for m in result.materials], ensure_ascii=False),
                    int(result.confidence),
                    json.dumps(result.warnings, ensure_ascii=False),
                    float(result.elapsed_sec),
                    result.created_at,
                ),
            )
            con.commit()

    def save_corrections(self, file_hash: str, rows: List[MaterialLine]) -> None:
        with self._lock, self._connect() as con:
            for r in rows:
                original_dict = {
                    "qty": r.original_qty or r.qty,
                    "unit": r.original_unit or r.unit,
                    "description": r.original_description or r.description,
                    "part_number": r.original_part_number or r.part_number,
                    "manufacturer": r.original_manufacturer or r.manufacturer,
                    "source_line": r.source_line,
                    "confidence": r.confidence,
                    "ai_status": r.ai_status,
                }
                con.execute(
                    """
                    INSERT INTO corrections(
                        file_hash, source_line, old_json, original_scan_json,
                        corrected_json, user_decision, ai_status, sender_hint, created_at
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        file_hash,
                        r.source_line,
                        json.dumps(original_dict, ensure_ascii=False),
                        json.dumps(original_dict, ensure_ascii=False),
                        json.dumps(r.as_dict(), ensure_ascii=False),
                        r.user_decision or "",
                        r.ai_status or "",
                        "",
                        now_iso(),
                    ),
                )
            con.commit()

    def save_bbox_corrections(self, file_hash: str, rows: List["MaterialLine"]) -> None:
        """Persist user-drawn bbox for every row that has one."""
        with self._lock, self._connect() as con:
            for r in rows:
                if not r.source_bbox:
                    continue
                con.execute(
                    """
                    INSERT INTO bbox_corrections
                        (file_hash, source_line, description, source_page,
                         source_bbox, source_bbox_kind, created_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        file_hash,
                        r.source_line or "",
                        r.description or "",
                        int(r.source_page or 0),
                        r.source_bbox,
                        r.source_bbox_kind or "image",
                        now_iso(),
                    ),
                )
            con.commit()

    def load_bbox_corrections(self, file_hash: str) -> List[Dict]:
        """Return stored bbox corrections for a file so the UI can pre-fill them."""
        with self._lock, self._connect() as con:
            rows = con.execute(
                """SELECT source_line, description, source_page, source_bbox, source_bbox_kind
                   FROM bbox_corrections WHERE file_hash=?
                   ORDER BY id DESC""",
                (file_hash,),
            ).fetchall()
        return [dict(r) for r in rows]

    def recent_corrections(self, limit: int = 200) -> List["MaterialLine"]:
        with self._lock, self._connect() as con:
            rows = con.execute(
                "SELECT corrected_json FROM corrections ORDER BY id DESC LIMIT ?", (limit,)
            ).fetchall()
        out = []
        for row in rows:
            try:
                out.append(material_from_dict(json.loads(row["corrected_json"])))
            except Exception:
                pass
        return out

    def save_correction_patterns(self, file_hash: str, rows: List["MaterialLine"]) -> None:
        """Persist high-confidence before/after parsing examples.

        These patterns are intentionally conservative: they only auto-apply later
        when a future source line is extremely similar to a confirmed/corrected
        source line.  This teaches repeated vendor/customer formats without
        letting one correction rewrite unrelated material rows.
        """
        with self._lock, self._connect() as con:
            for r in rows:
                if (r.review_status or "").lower() not in {"confirmed", "corrected"}:
                    continue
                if (r.user_decision or "").lower() == "rejected":
                    continue
                src = normalize_space(r.source_line or "")
                if not src or len(src) < 5:
                    continue
                sig = normalize_source_signature(src)
                con.execute(
                    """
                    INSERT INTO correction_patterns(
                        file_hash, source_signature, source_line, corrected_qty, corrected_unit,
                        corrected_description, corrected_part_number, corrected_manufacturer, created_at
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        file_hash, sig, src, r.qty or "", normalize_unit(r.unit or ""),
                        r.description or "", r.part_number or "", r.manufacturer or "", now_iso(),
                    ),
                )
            con.commit()

    def recent_correction_patterns(self, limit: int = 400) -> List[Dict[str, str]]:
        with self._lock, self._connect() as con:
            rows = con.execute(
                """SELECT source_signature, source_line, corrected_qty, corrected_unit,
                          corrected_description, corrected_part_number, corrected_manufacturer
                   FROM correction_patterns ORDER BY id DESC LIMIT ?""",
                (limit,),
            ).fetchall()
        return [dict(r) for r in rows]


class FileExtractor:
    def __init__(self, settings: ScanSettings, log=None) -> None:
        self.settings = settings
        self.log = log or (lambda msg: None)
        self._paddle_instance = None

    def extract(self, path: Path) -> Tuple[str, str, str, List[str], str]:
        """Return native_text, ocr_text, table_text, warnings, engine_summary."""
        ext = path.suffix.lower()
        warnings: List[str] = []
        native_text = ""
        ocr_text = ""
        table_text = ""
        engines: List[str] = []
        # Side channel: when the source is a true grid (spreadsheet / Word table),
        # we keep the raw cell matrix so ScanEngine can use the dedicated structured
        # parser (column mapping) instead of re-parsing flattened text with fuzzy,
        # electrical-tuned heuristics that drop generic rows. None = not structured.
        self.last_structured = None

        # Spreadsheets (xlsx/xlsm/xls/csv/tsv) -> structured rows as aligned table text.
        # Checked before the plain-text branch so .csv/.tsv get real column parsing.
        if ext in SUPPORTED_SPREADSHEET_EXTS and self.settings.enable_text_files:
            native_text, sheet_table, warn, matrix = self.extract_spreadsheet(path)
            warnings.extend(warn)
            engines.append("spreadsheet")
            table_text = sheet_table
            if matrix:
                self.last_structured = {"kind": "spreadsheet", "matrix": matrix}
            if self.settings.enable_table_guess and not table_text:
                table_text = self.guess_tables_from_text(native_text)
            return native_text, ocr_text, table_text, warnings, ", ".join(engines)

        # Word .docx -> paragraphs + table rows.
        if ext in SUPPORTED_WORD_EXTS and self.settings.enable_text_files:
            native_text, doc_table, warn, matrices = self.extract_docx(path)
            warnings.extend(warn)
            engines.append("docx")
            table_text = doc_table
            if matrices:
                # Use the largest table as the structured grid (most likely the line items).
                best = max(matrices, key=lambda mx: (len(mx), max((len(r) for r in mx), default=0)))
                if best and len(best) >= 2:
                    self.last_structured = {"kind": "docx_table", "matrix": best}
            if self.settings.enable_table_guess and not table_text:
                table_text = self.guess_tables_from_text(native_text)
            return native_text, ocr_text, table_text, warnings, ", ".join(engines)

        if ext in SUPPORTED_TEXT_EXTS and self.settings.enable_text_files:
            native_text = safe_read_text(path)
            engines.append("text-file")
            table_text = self.guess_tables_from_text(native_text) if self.settings.enable_table_guess else ""
            return native_text, ocr_text, table_text, warnings, ", ".join(engines)

        if ext == ".pdf":
            if self.settings.enable_native_pdf:
                txt, tbl, warn = self.extract_pdf_native(path)
                native_text += txt
                table_text += tbl
                warnings.extend(warn)
                engines.append("PyMuPDF-native" if fitz else "pdf-native-unavailable")

            needs_ocr = self.needs_ocr(native_text)
            if self.settings.mode == "Slow Careful":
                needs_ocr = True
            ocr_captured = False
            if needs_ocr and getattr(self.settings, "enable_ollama_vision", False):
                ov, ov_tbl, ov_warn, ov_used = self.ollama_vision_ocr(path, "pdf")
                warnings.extend(ov_warn)
                if ov.strip():
                    ocr_text += ov
                    if ov_tbl:
                        table_text += "\n" + ov_tbl
                    engines.extend(ov_used)
                    ocr_captured = True
            if needs_ocr and not ocr_captured and (self.settings.enable_tesseract or self.settings.enable_paddleocr or self.settings.enable_surya_cli):
                ocr, tbl2, warn2, used = self.ocr_pdf(path)
                ocr_text += ocr
                if tbl2:
                    table_text += "\n" + tbl2
                warnings.extend(warn2)
                engines.extend(used)

        elif ext in SUPPORTED_IMAGE_EXTS:
            ocr_captured = False
            if getattr(self.settings, "enable_ollama_vision", False):
                ov, ov_tbl, ov_warn, ov_used = self.ollama_vision_ocr(path, "image")
                warnings.extend(ov_warn)
                if ov.strip():
                    ocr_text += ov
                    if ov_tbl:
                        table_text += ov_tbl
                    engines.extend(ov_used)
                    ocr_captured = True
            if not ocr_captured and (self.settings.enable_tesseract or self.settings.enable_paddleocr or self.settings.enable_surya_cli):
                ocr, tbl, warn, used = self.ocr_image_file(path)
                ocr_text += ocr
                table_text += tbl
                warnings.extend(warn)
                engines.extend(used)
            elif not ocr_captured:
                warnings.append("Image file found, but OCR engines are disabled.")
        else:
            warnings.append(f"Unsupported file type: {ext}")

        if self.settings.enable_table_guess:
            guessed = self.guess_tables_from_text("\n".join([native_text, ocr_text]))
            if guessed and guessed not in table_text:
                table_text += "\n" + guessed

        return native_text, ocr_text, table_text, warnings, ", ".join([e for e in engines if e]) or "none"

    def ollama_vision_ocr(self, path: Path, kind: str = "pdf") -> Tuple[str, str, List[str], List[str]]:
        """Optional (off by default): use a local Ollama vision model to transcribe
        a scanned page into plain text that then feeds the normal parser. Returns
        (text, table_text, warnings, engines_used). Any failure returns empty text
        so the caller falls back to Tesseract -- the app is never worse off."""
        import base64 as _b64
        warnings: List[str] = []
        model = (getattr(self.settings, "ollama_vision_model", "") or "llama3.2-vision").strip()
        url = (getattr(self.settings, "ollama_vision_url", "") or "http://127.0.0.1:11434/api/generate").strip()
        try:
            if kind == "image":
                img_paths: List[Path] = [path]
            else:
                imgs, warn = self.render_pdf_pages(path)  # same temp-page render as ocr_pdf
                warnings.extend(warn)
                img_paths = list(imgs)
            if not img_paths:
                return "", "", warnings, []
            prompt = (
                "You are transcribing an electrical purchase-order / RFQ page. Read EVERY "
                "material line item and output it as plain text, one item per line, exactly "
                "as written. For each line keep the quantity, unit of measure, the full "
                "description, any catalog/part number, and the manufacturer if shown. Do not "
                "invent values, do not add commentary, do not summarise. Output only the lines."
            )
            chunks: List[str] = []
            for ip in img_paths:
                try:
                    with open(ip, "rb") as fh:
                        b64 = _b64.b64encode(fh.read()).decode("ascii")
                    body = json.dumps({
                        "model": model, "prompt": prompt,
                        "images": [b64], "stream": False,
                    }).encode("utf-8")
                    req = urllib.request.Request(url, data=body,
                                                 headers={"Content-Type": "application/json"})
                    with urllib.request.urlopen(req, timeout=180) as resp:
                        data = json.loads(resp.read().decode("utf-8"))
                    txt = (data.get("response") or "").strip()
                    if txt:
                        name = getattr(ip, "name", "page")
                        chunks.append(f"\n--- VISION {name} ---\n{txt}")
                except Exception as e:
                    warnings.append(f"Ollama vision failed on a page: {e}")
            text = normalize_space("\n".join(chunks)).strip()
            used = [f"Ollama-vision:{model}"] if text else []
            return text, "", warnings, used
        except Exception as e:
            warnings.append(f"Ollama vision OCR error: {e}")
            return "", "", warnings, []

    def needs_ocr(self, native_text: str) -> bool:
        lines = split_lines(native_text)
        if len(native_text.strip()) < 120:
            return True
        if len(lines) < 4:
            return True
        material_hits = sum(1 for ln in lines if looks_materialish(ln))
        if material_hits == 0 and self.settings.mode != "Fast":
            return True
        return False

    def extract_pdf_native(self, path: Path) -> Tuple[str, str, List[str]]:
        warnings = []
        if fitz is None:
            return "", "", ["PyMuPDF is not installed. Native PDF extraction skipped."]
        parts: List[str] = []
        table_parts: List[str] = []
        try:
            doc = fitz.open(str(path))
            max_pages = min(len(doc), self.settings.max_pages())
            for i in range(max_pages):
                page = doc[i]
                text = page.get_text("text", sort=True) or ""
                blocks = page.get_text("blocks", sort=True) or []
                if text.strip():
                    parts.append(f"\n--- PAGE {i+1} NATIVE TEXT ---\n{text}")
                # Heuristic table/line reconstruction from blocks/words.
                try:
                    words = page.get_text("words", sort=True) or []
                    table_like = self.words_to_table_lines(words)
                    if table_like:
                        table_parts.append(f"\n--- PAGE {i+1} POSITIONAL LINES ---\n{table_like}")
                except Exception:
                    pass
            if len(doc) > max_pages:
                warnings.append(f"PDF has {len(doc)} pages; scanned first {max_pages} page(s) in {self.settings.mode} mode.")
            doc.close()
        except Exception as e:
            warnings.append(f"Native PDF extraction failed: {e}")
        # v2.8.6: the positional/table lines carry their column separation as runs of
        # 2+ spaces -- that separation IS the signal the header-aware table parser
        # keys on (has_table_spacing / split_table_cells).  normalize_space() collapses
        # those runs to one space, which silently disabled header-table extraction for
        # every native PDF.  Clean the table text gently instead, preserving spacing.
        table_out = "\n".join(
            ln.rstrip() for ln in "\n".join(table_parts).replace("\r\n", "\n").replace("\r", "\n").splitlines()
        )
        table_out = re.sub(r"\n{3,}", "\n\n", table_out).strip()
        return normalize_space("\n".join(parts)), table_out, warnings

    def words_to_table_lines(self, words: List[Any]) -> str:
        if not words:
            return ""
        # v2.8.6: cluster words into visual rows by y-proximity instead of fixed 4pt
        # buckets.  A fixed bucket splits one printed row into two half-rows whenever
        # its cells' baselines straddle a bucket boundary (common in filled forms
        # where qty/desc/cost-code cells are written at slightly different y) -- the
        # half-row without the quantity then fails parsing and the item is lost.
        entries: List[Tuple[float, float, float, str]] = []
        for w in words:
            try:
                x0, y0, x1, y1, text = w[:5]
                entries.append((float(y0), float(x0), float(x1), str(text)))
            except Exception:
                continue
        if not entries:
            return ""
        entries.sort(key=lambda t: (t[0], t[1]))
        clusters: List[List[Tuple[float, float, float, str]]] = []
        last_y = None
        for e in entries:
            if last_y is not None and (e[0] - last_y) <= 3.2:
                clusters[-1].append(e)
            else:
                clusters.append([e])
            last_y = e[0]
        lines = []
        for cluster in clusters:
            items = sorted(((x0, x1, text) for (_y, x0, x1, text) in cluster), key=lambda t: t[0])
            if not items:
                continue
            # v2.8.6: gap-aware column reconstruction.  The old code joined EVERY word
            # with two spaces, and extract_pdf_native() then normalize_space()d the
            # result -- collapsing all runs to ONE space.  That destroyed the column
            # separation the header-aware parser needs (has_table_spacing / split on
            # 2+ spaces), so clean "QUANTITY  DESCRIPTION  COST CODE" tables were never
            # recognized and whole tables fell through to the fuzzy text path.  Now a
            # word gap is a COLUMN break (two spaces) only when it is clearly wider
            # than this line's normal word spacing; words inside a cell stay single-
            # spaced.  The threshold adapts to the line's own word gaps so different
            # font sizes work, with sane floors/caps in PDF points.
            gaps = []
            for i in range(1, len(items)):
                g = items[i][0] - items[i - 1][1]
                if 0 <= g < 12:
                    gaps.append(g)
            gaps.sort()
            med = gaps[len(gaps) // 2] if gaps else 3.0
            col_gap = min(16.0, max(8.0, med * 3.0))
            pieces = [items[0][2]]
            for i in range(1, len(items)):
                g = items[i][0] - items[i - 1][1]
                pieces.append(("  " if g >= col_gap else " ") + items[i][2])
            line = "".join(pieces)
            if looks_materialish(line) or has_table_spacing(line):
                lines.append(line)
        return "\n".join(lines)

    def render_pdf_pages(self, path: Path) -> Tuple[List[Path], List[str]]:
        warnings = []
        out_paths: List[Path] = []
        if fitz is None:
            return out_paths, ["PyMuPDF is not installed. Cannot render PDF pages for OCR."]
        tmp = Path(tempfile.mkdtemp(prefix="smartscan_pages_"))
        try:
            doc = fitz.open(str(path))
            max_pages = min(len(doc), self.settings.max_pages())
            dpi = self.settings.render_dpi()
            for i in range(max_pages):
                page = doc[i]
                pix = page.get_pixmap(dpi=dpi, alpha=False)
                img_path = tmp / f"page_{i+1:03d}.png"
                pix.save(str(img_path))
                out_paths.append(img_path)
            if len(doc) > max_pages:
                nxt = "Slow Careful" if self.settings.mode != "Slow Careful" else self.settings.mode
                warnings.append(
                    f"PDF has {len(doc)} pages; only the first {max_pages} were scanned "
                    f"in {self.settings.mode} mode. Use {nxt} mode (Settings raises the "
                    f"page limit) to scan all pages."
                )
            doc.close()
        except Exception as e:
            warnings.append(f"PDF page rendering failed: {e}")
        return out_paths, warnings

    def ocr_pdf(self, path: Path) -> Tuple[str, str, List[str], List[str]]:
        images, warnings = self.render_pdf_pages(path)
        all_text = []
        table_text = []
        used: List[str] = []
        for img_path in images:
            txt, tbl, warn, engines = self.ocr_image_file(img_path)
            if txt.strip():
                all_text.append(f"\n--- OCR {img_path.name} ---\n{txt}")
            if tbl.strip():
                table_text.append(tbl)
            warnings.extend(warn)
            used.extend(engines)
        return normalize_space("\n".join(all_text)), normalize_space("\n".join(table_text)), warnings, sorted(set(used))

    def ocr_image_file(self, path: Path) -> Tuple[str, str, List[str], List[str]]:
        warnings = []
        text_parts: List[str] = []
        table_parts: List[str] = []
        used: List[str] = []
        img_path = path
        try:
            if self.settings.enable_preprocess:
                img_path = self.preprocess_image(path)
        except Exception as e:
            warnings.append(f"Image preprocessing failed: {e}")
            img_path = path

        if self.settings.enable_paddleocr:
            txt, warn = self.run_paddleocr(img_path)
            if txt:
                text_parts.append(txt)
                used.append("PaddleOCR")
            warnings.extend(warn)

        if self.settings.enable_tesseract:
            txt, warn = self.run_tesseract(img_path)
            if txt:
                text_parts.append(txt)
                used.append("Tesseract")
            warnings.extend(warn)

        if self.settings.enable_surya_cli:
            txt, warn = self.run_surya_cli(img_path)
            if txt:
                text_parts.append(txt)
                used.append("Surya-CLI")
            warnings.extend(warn)

        merged = normalize_space("\n".join(text_parts))
        if self.settings.enable_table_guess:
            table_parts.append(self.guess_tables_from_text(merged))
        return merged, normalize_space("\n".join(table_parts)), warnings, used

    def preprocess_image(self, path: Path) -> Path:
        if Image is None:
            return path
        img = Image.open(str(path)).convert("RGB")
        # Basic cleanup that is safe for text documents.
        img = ImageOps.exif_transpose(img)
        gray = ImageOps.grayscale(img)
        gray = ImageOps.autocontrast(gray)
        try:
            gray = ImageEnhance.Contrast(gray).enhance(1.7)
            gray = ImageEnhance.Sharpness(gray).enhance(1.4)
            gray = gray.filter(ImageFilter.MedianFilter(size=3))
        except Exception:
            pass

        # Optional OpenCV thresholding for careful mode.
        if cv2 is not None and np is not None and self.settings.mode in {"Normal", "Slow Careful"}:
            arr = np.array(gray)
            try:
                arr = cv2.threshold(arr, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
                gray = Image.fromarray(arr)
            except Exception:
                pass

        tmpdir = Path(tempfile.mkdtemp(prefix="smartscan_pre_"))
        out = tmpdir / (path.stem + "_preprocessed.png")
        gray.save(str(out))
        return out

    def run_tesseract(self, img_path: Path) -> Tuple[str, List[str]]:
        warnings = []
        if pytesseract is None:
            return "", ["pytesseract is not installed. Tesseract OCR skipped."]
        try:
            resolved_cmd = _smartscan_configure_pytesseract(self.settings)
            if not resolved_cmd:
                warnings.append(
                    "Tesseract OCR could not locate tesseract.exe automatically. "
                    "Install Tesseract or set Settings > Tesseract command to "
                    r"C:\Program Files\Tesseract-OCR\tesseract.exe"
                )
            config = "--psm 6" if self.settings.mode == "Fast" else "--psm 6 --oem 3"
            txt = pytesseract.image_to_string(str(img_path), lang=self.settings.tesseract_lang or "eng", config=config)
            return normalize_space(txt), warnings
        except Exception as e:
            cmd_hint = getattr(getattr(pytesseract, "pytesseract", None), "tesseract_cmd", "")
            extra = f" (using: {cmd_hint})" if cmd_hint else ""
            warnings.append(f"Tesseract OCR failed{extra}: {e}")
            return "", warnings

    def run_paddleocr(self, img_path: Path) -> Tuple[str, List[str]]:
        warnings = []
        if PaddleOCR is None:
            return "", ["PaddleOCR is not installed. Paddle OCR skipped."]
        try:
            if self._paddle_instance is None:
                self._paddle_instance = PaddleOCR(use_angle_cls=True, lang=self.settings.paddle_lang or "en", show_log=False)
            result = self._paddle_instance.ocr(str(img_path), cls=True)
            lines: List[str] = []
            for page in result or []:
                for item in page or []:
                    try:
                        lines.append(str(item[1][0]))
                    except Exception:
                        pass
            return normalize_space("\n".join(lines)), warnings
        except Exception as e:
            warnings.append(f"PaddleOCR failed: {e}")
            return "", warnings

    def run_surya_cli(self, img_path: Path) -> Tuple[str, List[str]]:
        warnings = []
        cmd = self.settings.surya_command.strip() or "surya_ocr"
        # Surya packaging/CLI varies, so this is intentionally a best-effort hook.
        # If unavailable, it fails quietly and tells the user.
        try:
            proc = subprocess.run(
                [cmd, str(img_path)],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=120 if self.settings.mode == "Slow Careful" else 60,
            )
            if proc.returncode != 0:
                warnings.append(f"Surya CLI failed: {proc.stderr.strip()[:400]}")
                return "", warnings
            return normalize_space(proc.stdout), warnings
        except FileNotFoundError:
            warnings.append("Surya CLI command not found. Configure it in Settings or disable Surya.")
        except Exception as e:
            warnings.append(f"Surya CLI failed: {e}")
        return "", warnings

    def guess_tables_from_text(self, text: str) -> str:
        lines = split_lines(text)
        tableish = []
        in_material_table = False
        for ln in lines:
            if HEADER_QTY_RE.search(ln.lower()) and HEADER_DESC_RE.search(ln.lower()):
                in_material_table = True
                tableish.append(ln)
                continue
            if in_material_table:
                if re.match(r"^(subtotal|total|tax|terms|thank you|signature|prepared by|freight|shipping)", ln.strip(), re.I):
                    in_material_table = False
                    continue
                if is_definite_admin_noise(ln):
                    continue
                if has_table_spacing(ln) or looks_materialish(ln) or has_explicit_uom(ln):
                    tableish.append(ln)
                    continue
            elif looks_materialish(ln) and not is_definite_admin_noise(ln):
                tableish.append(ln)
        return normalize_space("\n".join(tableish))

    # --- Office document extraction (spreadsheets + Word) --------------------
    def extract_spreadsheet(self, path: Path) -> Tuple[str, str, List[str], List[List[str]]]:
        """Extract rows from xlsx/xlsm/xls/csv/tsv as aligned, column-separated text.

        Returns (native_text, table_text, warnings).  table_text uses a wide cell
        separator so MaterialParser's table/column logic sees clear columns.  Falls
        back to plain text when an optional library is unavailable, so a spreadsheet
        is never silently dropped.

        Returns (native_text, table_text, warnings, matrix) where matrix is the list
        of cleaned cell rows (empty list when unavailable).
        """
        warnings: List[str] = []
        ext = path.suffix.lower()
        matrix: List[List[Any]] = []
        try:
            if ext in {".csv", ".tsv"}:
                matrix = self._read_delimited_rows(path, warnings)
            elif ext in {".xlsx", ".xlsm"}:
                if openpyxl is None:
                    warnings.append(
                        "openpyxl not installed; reading spreadsheet as plain text. "
                        "Run 'pip install openpyxl' for accurate cell extraction."
                    )
                    txt = safe_read_text(path)
                    table = self.guess_tables_from_text(txt) if self.settings.enable_table_guess else ""
                    return txt, table, warnings, []
                matrix = self._read_xlsx_rows(path, warnings)
            elif ext == ".xls":
                if xlrd is None:
                    warnings.append(
                        "xlrd not installed; cannot read legacy .xls. "
                        "Run 'pip install xlrd' or resave the file as .xlsx."
                    )
                    return "", "", warnings, []
                matrix = self._read_xls_rows(path, warnings)
        except Exception as e:
            warnings.append(f"Spreadsheet extraction failed: {e}")
            return "", "", warnings, []

        table_lines: List[str] = []
        plain_lines: List[str] = []
        clean_matrix: List[List[str]] = []
        for row in matrix:
            cells = [normalize_space(str(c)) for c in row]
            while cells and cells[-1] == "":
                cells.pop()
            if not any(cells):
                continue
            clean_matrix.append(cells)
            table_lines.append("   ".join(cells))
            plain_lines.append(" ".join(c for c in cells if c))
        table_text = "\n".join(table_lines)
        native_text = "\n".join(plain_lines)
        if not table_text:
            warnings.append("Spreadsheet had no readable rows.")
        return native_text, table_text, warnings, clean_matrix

    def _read_delimited_rows(self, path: Path, warnings: List[str]) -> List[List[str]]:
        raw = safe_read_text(path)
        if not raw.strip():
            return []
        if path.suffix.lower() == ".tsv":
            delim = "\t"
        else:
            delim = ","
            try:
                dialect = csv.Sniffer().sniff(raw[:8192], delimiters=",;\t|")
                delim = dialect.delimiter
            except Exception:
                delim = ","
        try:
            return [list(r) for r in csv.reader(raw.splitlines(), delimiter=delim)]
        except Exception as e:
            warnings.append(f"CSV parse fell back to raw lines ({e}).")
            return [[ln] for ln in raw.splitlines()]

    def _read_xlsx_rows(self, path: Path, warnings: List[str]) -> List[List[Any]]:
        rows: List[List[Any]] = []
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        try:
            sheets = list(wb.worksheets)
            multi = len(sheets) > 1
            for ws in sheets:
                if multi:
                    rows.append([f"--- SHEET: {ws.title} ---"])
                count = 0
                for row in ws.iter_rows(values_only=True):
                    rows.append(["" if v is None else v for v in row])
                    count += 1
                    if count >= SPREADSHEET_MAX_ROWS_PER_SHEET:
                        warnings.append(f"Sheet '{ws.title}' truncated at {SPREADSHEET_MAX_ROWS_PER_SHEET} rows.")
                        break
        finally:
            try:
                wb.close()
            except Exception:
                pass
        return rows

    def _read_xls_rows(self, path: Path, warnings: List[str]) -> List[List[Any]]:
        rows: List[List[Any]] = []
        book = xlrd.open_workbook(str(path))
        sheets = book.sheets()
        multi = len(sheets) > 1
        for sheet in sheets:
            if multi:
                rows.append([f"--- SHEET: {sheet.name} ---"])
            nrows = min(sheet.nrows, SPREADSHEET_MAX_ROWS_PER_SHEET)
            for rx in range(nrows):
                rows.append([sheet.cell_value(rx, cx) for cx in range(sheet.ncols)])
            if sheet.nrows > SPREADSHEET_MAX_ROWS_PER_SHEET:
                warnings.append(f"Sheet '{sheet.name}' truncated at {SPREADSHEET_MAX_ROWS_PER_SHEET} rows.")
        return rows

    def extract_docx(self, path: Path) -> Tuple[str, str, List[str], List[List[List[str]]]]:
        """Extract paragraphs and table rows from a .docx file.

        Returns (native_text, table_text, warnings, table_matrices) where
        table_matrices is a list of per-table cell grids for the structured parser.
        """
        warnings: List[str] = []
        if docx is None:
            warnings.append(
                "python-docx not installed; cannot read .docx accurately. "
                "Run 'pip install python-docx'."
            )
            return "", "", warnings, []
        try:
            document = docx.Document(str(path))
        except Exception as e:
            warnings.append(f"DOCX open failed: {e}")
            return "", "", warnings, []
        paras: List[str] = []
        table_lines: List[str] = []
        table_matrices: List[List[List[str]]] = []
        try:
            for p in document.paragraphs:
                t = normalize_space(p.text)
                if t:
                    paras.append(t)
        except Exception:
            pass
        try:
            for table in document.tables:
                grid: List[List[str]] = []
                for row in table.rows:
                    cells = [normalize_space(c.text) for c in row.cells]
                    trimmed = list(cells)
                    while trimmed and trimmed[-1] == "":
                        trimmed.pop()
                    if any(trimmed):
                        table_lines.append("   ".join(trimmed))
                        grid.append(trimmed)
                if grid:
                    table_matrices.append(grid)
        except Exception:
            pass
        native_text = "\n".join(paras)
        table_text = "\n".join(table_lines)
        if not native_text and not table_text:
            warnings.append("DOCX had no readable text or tables.")
        return native_text, table_text, warnings, table_matrices


def has_table_spacing(line: str) -> bool:
    if re.search(r"\S\s{2,}\S", line):
        return True
    if line.count("|") >= 2 or line.count("\t") >= 2:
        return True
    return False


def contains_electrical_keyword(text: str) -> bool:
    low = (text or "").lower()
    for term in ELECTRICAL_KEYWORDS:
        t = term.lower()
        # Short abbreviations like CT/LB/GE cause false positives inside words such as
        # "electric" or "page", so require token boundaries for short terms.
        if len(t) <= 3 and t.isalnum():
            if re.search(r"(?<![a-z0-9])" + re.escape(t) + r"(?![a-z0-9])", low):
                return True
        elif t in low:
            return True
    return False


def looks_materialish(line: str) -> bool:
    low = line.lower().strip()
    if len(low) < 3:
        return False
    if contains_electrical_keyword(low):
        return True
    if QTY_LEADING_RE.match(line) and (SIZE_RE.search(line) or PART_NO_RE.search(line)):
        return True
    if re.search(r"\b\d+\s*(?:ea|ft|pcs?|boxes|rolls?|sticks?)\b", low):
        return True
    if SIZE_RE.search(low) and any(x in low for x in ["wire", "emt", "pvc", "conduit", "box", "connector", "breaker"]):
        return True
    return False





def has_explicit_uom(text: str) -> bool:
    """True when the line contains a real unit of measure token, not an address suffix."""
    if not text:
        return False
    # Do not let road abbreviations count as material UOM.  "rd" and "st" are intentionally absent.
    # "lot" is intentionally absent too (matches "parking lot", "a lot of", etc.).
    return bool(re.search(r"\b(?:ea|each|pc|pcs|ft|lf|feet|foot|box|boxes|roll|rolls|stick|sticks|set|sets|bag|bags|case|cases|reel|reels|spool|spools|coil|coils|bucket|buckets|drum|drums|pallet|pallets|gal|gallon|gallons|kit|kits|lb|lbs|pack|packs|pkg|pkgs|dozen|dozens|m|meter|meters|yd|yards)\b", text, re.I))


# Common material nouns (electrical + general hardware) used to keep all-caps
# material descriptions from being misread as company/customer names.
MATERIAL_NOUN_RE = re.compile(
    r"\b(?:tape|wire|cable|conduit|emt|pvc|rmc|imc|washer|washers|nut|nuts|bolt|bolts|"
    r"screw|screws|rod|rods|box|boxes|connector|connectors|coupling|couplings|"
    r"strap|straps|clamp|clamps|fitting|fittings|elbow|elbows|breaker|breakers|"
    r"lug|lugs|terminal|terminals|wirenut|anchor|anchors|hanger|hangers|bushing|"
    r"bushings|locknut|locknuts|nipple|nipples|plate|plates|cover|covers|gland|"
    r"raceway|strut|unistrut|fastener|fasteners|tie|ties|sleeve|sleeves|grommet|"
    r"receptacle|switch|switches|fuse|fuses|panel|panels|gasket|gaskets|"
    r"pipe|tubing|hub|hubs|reducer|reducers|knockout|standoff|spacer|spacers)\b",
    re.I,
)


def is_definite_admin_noise(line: str) -> bool:
    """Aggressively reject quote headers, addresses, contact info, and customer fields."""
    raw = normalize_space(line or "")
    if not raw:
        return True
    low = raw.lower()
    if EMAIL_RE.search(raw) or URL_RE.search(raw) or PHONE_RE.search(raw):
        return True
    if ADDRESS_LABEL_RE.search(raw) or DOC_ADMIN_RE.search(raw):
        return True
    # Addresses: allow only if the same line is clearly a material line with UOM/material words.
    if STREET_SUFFIX_RE.search(raw) or ZIP_STATE_RE.search(raw):
        # A street-suffix abbreviation often appears INSIDE a part number ("CT-ELEC"
        # contains "CT"=Court), which would wrongly drop a real item.  Keep the line
        # when it has UOM+electrical wording OR contains a concrete material noun
        # (tape, wire, conduit, ...) — genuine addresses do not contain those.
        if not (has_explicit_uom(raw) and contains_electrical_keyword(raw)) and not MATERIAL_NOUN_RE.search(raw):
            return True
    if re.search(r"^\s*(?:page\s+)?\d+\s+of\s+\d+\s*$", low):
        return True
    if re.fullmatch(r"[A-Z .,&'-]{4,}", raw) and not contains_electrical_keyword(raw):
        # All-caps with letters only *can* be a company/customer name — but many real
        # PO descriptions are also all-caps ("TAPE ELECTRICAL", "LOCK WASHER", "RHMS").
        # Only reject when the line carries no material signal whatsoever (no UOM word,
        # no size, no digits) AND it reads like a name/admin field.  Otherwise keep it
        # so legitimate all-caps material descriptions survive the OCR/text path.
        if has_explicit_uom(raw) or SIZE_RE.search(raw) or re.search(r"\d", raw):
            return False
        if MATERIAL_NOUN_RE.search(raw):
            return False
        return True
    return False


def material_signal_score(text: str, qty: str = "", unit: str = "", part: str = "") -> int:
    """Score whether a parsed row really looks like an electrical/material line."""
    t = normalize_space(text or "")
    score = 0
    if qty:
        score += 12
    if unit or has_explicit_uom(t):
        score += 22
    if contains_electrical_keyword(t):
        score += 30
    if SIZE_RE.search(t):
        score += 10
    if part and not re.fullmatch(r"\d+", part):
        score += 8
    # v2.8.6: catalog-style evidence.  A quantity plus a REAL catalog token is how
    # most distributor RFQ lines actually look ("15  Champion 40A-SW-82-P 45deg",
    # "50  Ipex 077436 Poly Plugs 5\""), yet these scored below the 38 keep-gate
    # whenever the description had no electrical keyword -- whole tables of brand +
    # part-number rows were being dropped.  Count a mixed letter/digit dashed token
    # (40A-SW-82-P) or a 5-7 digit catalog id (077436) as material evidence, but
    # only alongside a quantity so plain addresses/headers do not inflate; the
    # CUSTOMER_INFO / admin-noise penalties below still dominate for contact lines.
    if qty:
        catalogish = bool(
            (part and re.search(r"[A-Za-z]", part) and re.search(r"\d", part))
            or re.search(r"\b(?=[A-Za-z0-9-]{5,24}\b)(?=[^\s]*\d)(?=[^\s]*[A-Za-z])[A-Za-z0-9]+(?:-[A-Za-z0-9]+)+\b", t)
        )
        if catalogish:
            score += 16
        elif re.search(r"\b\d{5,7}\b", t) and re.search(r"[A-Za-z]{3,}", t):
            score += 12
    # Catalog-style lines with a Qty/UOM but few obvious keywords should still pass.
    if qty and (unit or has_explicit_uom(t)) and re.search(r"[A-Za-z]", t):
        score += 12
    if CUSTOMER_INFO_RE.search(t):
        score -= 50
    if is_definite_admin_noise(t):
        score -= 80
    return score


def should_keep_material_candidate(row: "MaterialLine") -> bool:
    src = row.source_line or row.description or ""
    if is_definite_admin_noise(src) or is_definite_admin_noise(row.description):
        return False
    # v2.8.6: rows from the header-aware table parser already passed that path's own
    # noise/contact/admin gates, and a row with a quantity under a recognized
    # QUANTITY/DESCRIPTION header is a line item regardless of keyword score --
    # re-applying the electrical-keyword score here was silently discarding the
    # table rows the header parser had just accepted (consumables, generic brands).
    if row.qty and "Header-aware Qty/Description extraction" in (row.notes or ""):
        return True
    desc = row.description or ""
    # Require at least some material evidence. This is the key noise reduction for
    # native-text PDFs where every header/address line can otherwise look table-ish.
    return material_signal_score(desc, row.qty, row.unit, row.part_number) >= 38



def parse_dash_quantity_source_line(raw: str) -> Optional["MaterialLine"]:
    """Lenient OCR audit parser for lines like '-(2) Safety Glasses'.

    This is intentionally used only by the extraction watchdog/final audit, not
    as the main parser.  It catches material-list lines that OCR clearly presents
    as quantity bullets even when the description lacks electrical keywords or a
    formal UOM.  These rows are flagged for review instead of being silently
    trusted.
    """
    line = (raw or "").strip()
    if not line:
        return None
    if is_definite_admin_noise(line):
        return None
    # Only treat explicit list quantity lines as candidates.  This avoids
    # ordinary sentences, addresses, and headers while catching:
    #   -(2) Safety Glasses
    #   -(25) Ear Plugs
    #   -(2 dozen) Work Gloves
    #   -(500pcs.) 3M Red/Yellow wire nuts
    #   -(1case/50pcs.) 4-11/16" x 5/8" cover plates
    m = re.match(
        r"^\s*(?:[-–—*•●▪]\s*)?[\(\[\{]\s*"
        r"(?P<qty>\d+(?:\.\d+)?)\s*"
        r"(?P<unit>[A-Za-z]+)?\.?"
        r"(?:\s*/\s*(?P<pack>\d+\s*(?:pc|pcs|piece|pieces|ea|each|case|cases|box|boxes|bag|bags)))?"
        r"\s*[\)\]\}]\s*(?P<desc>.+?)\s*$",
        line,
        re.I,
    )
    if not m:
        return None
    qty = normalize_qty(m.group("qty") or "")
    unit = normalize_unit(m.group("unit") or "")
    desc = clean_description(m.group("desc") or "")
    if not qty or not desc:
        return None
    if len(desc) < 2 or re.fullmatch(r"[-–—.,;:]+", desc):
        return None

    # If the unit was outside the parentheses as a plural word, recover it.
    if not unit:
        unit_words = r"ea|each|pcs?|pieces?|ft|feet|foot|lf|rolls?|boxes|box|bags?|cases?|coils?|buckets?|dozens?|packs?|pkgs?"
        unit_at_front = re.match(rf"^\s*(?P<unit>{unit_words})\b\s*(?P<rest>.*)$", desc, re.I)
        if unit_at_front:
            possible_unit = normalize_unit(unit_at_front.group("unit") or "")
            rest = clean_description(unit_at_front.group("rest") or "")
            if possible_unit and rest:
                unit = possible_unit
                desc = re.sub(r"^\s*of\s+", "", rest, flags=re.I).strip()

    part_number = ""
    part_matches = PART_NO_RE.findall(desc)
    for p in part_matches:
        p_clean = (p or "").strip().strip('"')
        pl = p_clean.lower().strip(".-_/ ")
        if SIZE_RE.fullmatch(p_clean) or re.fullmatch(r"\d+\s*-\s*\d+\s*/\s*\d+(?:\s*[\"'])?", p_clean):
            continue
        if re.fullmatch(r"[A-Za-z][A-Za-z-]{2,18}", p_clean) and not (p_clean.isupper() and "-" not in p_clean and len(p_clean) >= 4):
            continue
        if pl in ELECTRICAL_KEYWORDS or pl in ELECTRICAL_UNITS:
            continue
        if re.fullmatch(r"\d+", pl):
            continue
        part_number = p_clean
        break

    conf = 62
    if unit:
        conf += 10
    if contains_electrical_keyword(desc.lower()):
        conf += 12
    if SIZE_RE.search(desc):
        conf += 8
    if part_number:
        conf += 6
    return MaterialLine(
        qty=qty,
        unit=unit,
        description=desc,
        part_number=part_number,
        manufacturer=best_manufacturer_hint(desc),
        notes="OCR quantity-list audit candidate; please confirm",
        confidence=max(0, min(96, conf)),
        source_line=line,
        ai_status="warning",
        ai_message="Final Audit: OCR line has a clear -(quantity) pattern but was not otherwise accounted for.",
        watchdog_status="missing_candidate",
        watchdog_message="Clear OCR quantity-list line; verify qty/description.",
        highlight_color="yellow",
    )

def split_table_cells(line: str) -> List[str]:
    line = line.strip()
    if "|" in line:
        cells = [c.strip() for c in line.split("|")]
    elif "\t" in line:
        cells = [c.strip() for c in line.split("\t")]
    else:
        cells = [c.strip() for c in re.split(r"\s{2,}", line)]
    return [c for c in cells if c]



def is_integer_token(value: str) -> bool:
    return bool(re.fullmatch(r"\d{1,6}", (value or "").strip()))


def is_likely_line_number_token(value: str) -> bool:
    value = (value or "").strip()
    if not is_integer_token(value):
        return False
    try:
        n = int(value)
    except Exception:
        return False
    # RFQ table row/line numbers are usually small sequential integers.  A value like
    # 100 or 250 is much more likely to be a quantity, so only treat 1-99 as row ids.
    return 1 <= n <= 99


def is_price_cell(value: str) -> bool:
    v = (value or "").strip()
    if not v:
        return True
    if v in {"-", "--", "—", "–", "$", "$-"}:
        return True
    if re.fullmatch(r"\$?\s*\d{1,3}(?:,\d{3})*(?:\.\d{2})", v):
        return True
    if re.fullmatch(r"\$?\s*0+(?:\.0+)?", v):
        return True
    return False


def normalize_unit(unit: str) -> str:
    unit = (unit or "").strip().strip(". ,;:[]{}()")
    low = unit.lower()
    aliases = {
        "each": "EA", "ea": "EA", "pc": "EA", "pcs": "EA", "piece": "EA", "pieces": "EA",
        "foot": "FT", "feet": "FT", "ft": "FT", "lf": "FT",
        "box": "BOX", "boxes": "BOX", "roll": "ROLL", "rolls": "ROLL",
        "stick": "STICK", "sticks": "STICK", "set": "SET", "sets": "SET",
        "lot": "LOT", "lots": "LOT", "bag": "BAG", "bags": "BAG",
        "case": "CASE", "cases": "CASE", "reel": "REEL", "reels": "REEL",
        "spool": "SPOOL", "spools": "SPOOL", "coil": "COIL", "coils": "COIL",
        "bucket": "BUCKET", "buckets": "BUCKET", "dozen": "DOZEN", "dozens": "DOZEN", "lb": "LB", "lbs": "LB",
        "pack": "PACK", "packs": "PACK", "pkg": "PACK", "pkgs": "PACK",
        "drum": "DRUM", "drums": "DRUM", "pallet": "PALLET", "pallets": "PALLET",
        "gal": "GAL", "gallon": "GAL", "gallons": "GAL", "kit": "KIT", "kits": "KIT",
        "m": "M", "meter": "M", "meters": "M", "yd": "YD", "yard": "YD", "yards": "YD",
    }
    return aliases.get(low, unit.upper() if unit else "")


def strip_prices_from_text(text: str) -> str:
    # Remove common pricing columns that get glued onto descriptions by PDF/OCR text extraction.
    text = re.sub(r"(?:\$\s*)?\b\d{1,3}(?:,\d{3})*(?:\.\d{2})\b", " ", text or "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def best_description_from_cells(cells: List[str], skip_indexes: set) -> str:
    useful: List[str] = []
    for idx, cell in enumerate(cells):
        c = (cell or "").strip()
        if idx in skip_indexes or not c:
            continue
        if is_price_cell(c):
            continue
        if c in {"-", "--", "—", "–"}:
            continue
        useful.append(c)
    if not useful:
        return ""
    # Usually the last useful cell is the true Description column; earlier cells are
    # catalog/vendor numbers.  If the last cell is too weak, keep the joined text.
    last = clean_description(useful[-1])
    if contains_electrical_keyword(last.lower()) or SIZE_RE.search(last) or len(useful) > 2:
        return last
    return clean_description(" ".join(useful))

def split_qty_unit(value: str) -> Tuple[str, str]:
    value = (value or "").strip()
    m = re.match(r"^\(?\s*(\d+(?:\.\d+)?)\s*\)?\s*([A-Za-z]+)?\s*$", value)
    if not m:
        return normalize_qty(value), ""
    qty = normalize_qty(m.group(1) or "")
    unit_raw = (m.group(2) or "").strip()
    if unit_raw.lower() not in ELECTRICAL_UNITS:
        unit = ""
    else:
        unit = normalize_unit(unit_raw)
    return qty, unit


def normalize_qty(qty: str) -> str:
    qty = (qty or "").strip().strip("()[]{}")
    # Ignore phone-like fragments and ZIP codes accidentally captured as quantities.
    if re.fullmatch(r"\d{5}(?:-\d{4})?", qty):
        return ""
    return qty


def clean_part_number(part: str) -> str:
    part = (part or "").strip().strip(" |,;:")
    if not part:
        return ""
    matches = PART_NO_RE.findall(part)
    if not matches:
        return ""
    return matches[0]


def slice_by_header_positions(raw: str, positions: List[Tuple[int, str]]) -> Dict[str, str]:
    out: Dict[str, str] = {}
    if not positions:
        return out
    positions = sorted(positions)
    for idx, (start, label) in enumerate(positions):
        end = positions[idx + 1][0] if idx + 1 < len(positions) else len(raw)
        # Pad short rows so slices are safe.
        cell = raw[start:end].strip() if start < len(raw) else ""
        if cell:
            out[label] = cell
    # If the row is shorter or badly aligned, fall back to splitting and mapping in the same order.
    if (not out.get("qty") or not out.get("desc")) and has_table_spacing(raw):
        cells = split_table_cells(raw)
        for idx, (_, label) in enumerate(positions):
            if idx < len(cells) and cells[idx].strip():
                out.setdefault(label, cells[idx].strip())
    return out


def append_note(existing: str, note: str) -> str:
    existing = (existing or "").strip()
    if not existing:
        return note
    if note in existing:
        return existing
    return existing + "; " + note


def _normalize_audit_desc_key(text: str) -> str:
    """Normalize a material description for duplicate/missing accountability.

    This is deliberately stronger than display cleanup. It removes OCR/list
    punctuation, normalizes quotes/dashes, drops obvious packaging words at the
    beginning, and keeps size/catalog content so near-identical duplicate rows
    cluster together without merging different sizes.
    """
    t = (text or "").lower()
    t = t.replace("’", "'").replace("“", '"').replace("”", '"')
    t = re.sub(r"^[\s\-–—*•●▪]+", "", t)
    t = re.sub(r"^[\(\[\{]\s*\d+(?:\.\d+)?\s*(?:ea|each|pcs?|pieces?|ft|feet|foot|lf|rolls?|boxes|box|bags?|cases?|coils?|buckets?|dozens?|packs?|pkgs?)?\.?\s*[\)\]\}]\s*", "", t, flags=re.I)
    t = re.sub(r"^of\s+", "", t, flags=re.I)
    # common OCR variants that should not make a duplicate look different
    t = t.replace("—", "-").replace("–", "-")
    t = re.sub(r"\b(one\s*-?hole)\b", "one-hole", t)
    t = re.sub(r"\bdeep\s*[--]?\s*w\b", "deep w", t)
    t = re.sub(r"[^a-z0-9/\'\".#+-]+", " ", t)
    t = re.sub(r"\s+", " ", t).strip(" -–—|,;:")
    return t[:240]


def canonical_material_key(row: MaterialLine) -> Tuple[str, str, str, str]:
    desc = (row.description or "").lower()
    part = (row.part_number or "").lower()
    if part:
        desc = re.sub(r"(?<![a-z0-9])" + re.escape(part) + r"(?![a-z0-9])", " ", desc, flags=re.I)
    # Many general OCR lines append a part/catalog number to the description. Remove a
    # trailing catalog-looking token so the header-aware row and fallback row de-dupe
    # cleanly -- but ONLY when real descriptive words remain afterward. Otherwise a
    # description that IS entirely a part code (e.g. "12D1.5L" vs "16D1.5L") would be
    # stripped to empty and two different items would collapse into one false duplicate.
    stripped = re.sub(r"\b(?=[A-Z0-9._\-/]*\d)[A-Z0-9][A-Z0-9._\-/]{3,}\s*$", "", desc, flags=re.I)
    if re.search(r"[a-z]{3,}", stripped):
        desc = stripped
    desc = _normalize_audit_desc_key(desc)
    return ((row.qty or "").lower().strip(), normalize_unit(row.unit or "").lower(), desc, "")


# --- v2.8.4: part-number-aware duplicate suppression -------------------------------
# canonical_material_key() deliberately strips the catalog/part number out of the
# description so a column/header-aware row and a flattened-text row for the SAME item
# de-duplicate cleanly.  The side effect: two GENUINELY different parts whose
# descriptions clean to the same text (e.g. T&B "TY523M" vs "TY523MX") produced an
# identical key, so the canonical-key de-dup passes silently dropped one row.  The
# helpers below turn those de-dup sets into part-aware maps -- a key collision counts
# as a duplicate UNLESS the incoming row carries a real part number different from
# every real part number already recorded under that key.  Rows without a usable part
# number still collapse onto an existing key, preserving the original header/fallback
# de-duplication exactly.
def _real_part_sig(row: MaterialLine) -> str:
    """Normalized catalog part number, or "" when the row has no usable one."""
    p = (getattr(row, "part_number", "") or "").strip().lower()
    p = re.sub(r"[^a-z0-9]+", "", p)
    if len(p) >= 3 and re.search(r"\d", p):
        return p
    return ""


def key_seen_or_add(key: Tuple[str, str, str, str], row: MaterialLine, seen: dict) -> bool:
    """Part-aware de-dup gate for the canonical_material_key() passes.

    ``seen`` maps a canonical_material_key() tuple -> set of real part signatures
    already recorded for that key.  Returns True if ``row`` duplicates an entry
    already present (and the caller should skip it); otherwise records ``row`` and
    returns False.  A key collision is a duplicate unless ``row`` introduces a real
    part number distinct from every real part already stored for that key, so
    "TY523M" and "TY523MX" survive as separate lines while header-row vs.
    flattened-text copies of one item still collapse.
    """
    sig = _real_part_sig(row)
    prior = seen.get(key)
    if prior is None:
        seen[key] = {sig}
        return False
    real_prior = {p for p in prior if p}
    if sig and real_prior and sig not in real_prior:
        prior.add(sig)
        return False
    prior.add(sig)
    return True


def warn_key(row: MaterialLine) -> Tuple[Tuple[str, str, str, str], str]:
    """Part-aware key for duplicate *warnings* so two distinct part numbers that
    clean to the same description are not flagged as duplicates of each other."""
    return (canonical_material_key(row), _real_part_sig(row))


def audit_candidate_key(cand: "SourceMaterialCandidate") -> Tuple[str, str, str, str]:
    row = cand.parsed or MaterialLine(source_line=cand.source_line, description=cand.source_line)
    return canonical_material_key(row)


def loose_same_material(a: MaterialLine, b: MaterialLine, min_ratio: float = 0.86) -> bool:
    """Return True when two rows are almost certainly the same material line.

    Used only for audit/duplicate suppression, not for normal parsing.
    """
    if not a or not b:
        return False
    if (a.qty or "").strip() != (b.qty or "").strip():
        return False
    au = normalize_unit(a.unit or "")
    bu = normalize_unit(b.unit or "")
    if au and bu and au != bu:
        return False
    ad = _normalize_audit_desc_key(a.description or a.source_line)
    bd = _normalize_audit_desc_key(b.description or b.source_line)
    if not ad or not bd:
        return False
    if ad == bd or ad in bd or bd in ad:
        return True
    return SequenceMatcher(None, ad, bd).ratio() >= min_ratio


def is_probable_phone_or_address(text: str) -> bool:
    return is_probable_contact_line(text, strict_address=True)


def is_probable_contact_line(line: str, strict_address: bool = True) -> bool:
    raw = (line or "").strip()
    if not raw:
        return False
    low = raw.lower()
    material_hint = looks_materialish(raw)

    # Emails, URLs, phone/fax lines: exclude unless the same line is clearly material-heavy.
    if EMAIL_RE.search(raw) or URL_RE.search(raw):
        return not material_hint
    if PHONE_RE.search(raw):
        # A phone number should almost never be a material item. Exclude even if there are words.
        if not any(k in low for k in ["emt", "pvc", "conduit", "breaker", "wire", "cable"]):
            return True
    if ADDRESS_LABEL_RE.search(raw):
        return not material_hint
    if re.search(r"\b(?:phone|tel|fax|cell|mobile|email|e-mail|website|www)\b", low):
        return not material_hint

    if strict_address:
        # Street addresses normally start with a street number and contain a street suffix.
        if re.search(r"^\s*\d{1,6}\s+", raw) and STREET_SUFFIX_RE.search(raw):
            return not material_hint
        if ZIP_STATE_RE.search(raw) and not material_hint:
            return True
        if re.search(r"\b(?:suite|ste\.?|floor|fl\.?|unit|building|bldg\.?)\s*#?\s*\d", low) and not material_hint:
            return True
        if re.search(r"\bp\.?\s*o\.?\s*box\b", low):
            return True
    return False


def is_probably_contact_or_address(text: str) -> bool:
    """Compatibility wrapper used by the smart reviewer."""
    return is_probable_contact_line(text, strict_address=True)


def is_definite_noise_row(row: MaterialLine) -> bool:
    """Rows the smart scanner can safely hide before showing the user.

    We keep the logic conservative: only rows that the reviewer marked as definite
    contact/address/web noise, or that have no usable material fields after review,
    are removed from the user-facing output.
    """
    msg = (row.ai_message or "").lower()
    source = " ".join([row.source_line or "", row.description or "", row.notes or ""])
    if (row.ai_status or "").lower() == "error" and (
        "contact/address" in msg or "not a material line" in msg
    ):
        return True
    if is_definite_admin_noise(source):
        return True
    if is_probable_contact_line(source, strict_address=True) and not looks_materialish(source):
        return True
    if not (row.qty or row.unit or row.description or row.part_number).strip():
        return True
    return False


# ---------------------------------------------------------------------------
# Structured-table (spreadsheet / Word-table) extraction.
#
# When the source is a real grid, columns are explicit, so we map columns to
# fields and emit EVERY data row.  This deliberately bypasses the fuzzy,
# electrical-tuned "looks_materialish / material_signal_score" gates used for
# PDF/OCR text, which silently drop perfectly valid generic line items (the
# reason spreadsheets did not scan to 100%).
# ---------------------------------------------------------------------------

# Column-name synonyms.  Order matters only for reporting; matching is by set.
_STRUCT_HEADER_SYNONYMS = {
    "qty": [
        "qty", "quantity", "qnty", "qty.", "q'ty", "qty ordered", "order qty",
        "ordered", "count", "pcs", "pieces", "no. of units", "units", "ea",
        "amount", "amt", "req qty", "required", "req'd", "nbr", "number of",
    ],
    "unit": ["unit", "uom", "u/m", "u of m", "um", "unit of measure", "measure", "per"],
    "description": [
        "description", "desc", "desc.", "item description", "material", "materials",
        "item", "items", "product", "product description", "details", "detail",
        "name", "item name", "material description", "commodity", "line item",
        "service", "scope", "particulars", "goods", "article",
    ],
    "part_number": [
        "part", "part #", "part#", "part no", "part no.", "part number", "partno",
        "catalog", "catalog #", "catalog no", "cat", "cat #", "cat no", "cat.no",
        "catalog number", "sku", "model", "model #", "model no", "mpn", "mfg part",
        "manufacturer part", "item code", "item number", "stock #", "stock no",
        "ref", "ref #", "upc", "part/catalog", "cat. no", "catalog/part",
    ],
    "manufacturer": [
        "manufacturer", "mfg", "mfr", "mfg.", "make", "brand", "vendor", "supplier", "oem",
    ],
    "unit_price": [
        "unit price", "unit cost", "price", "price ea", "price/ea", "cost", "rate",
        "list price", "net price", "price each", "u/price", "unit $", "$/unit", "$ ea",
    ],
    "total_price": [
        "total", "total price", "ext price", "extended", "extended price", "ext.",
        "line total", "ext cost", "extended cost", "net amount", "ext amount", "line amount",
    ],
    "notes": ["notes", "note", "remarks", "remark", "comments", "comment"],
    "line_no": [
        "line", "line #", "line no", "line item", "item #", "item no", "item no.",
        "item", "no", "no.", "#", "s.no", "sr", "sr no", "seq", "lot #", "pos",
    ],
}

# Build a fast lookup from a normalized header token to its field name.  When the
# same token (e.g. "unit" or "amount") appears under multiple fields, the first
# field in this priority order wins for that ambiguous token.
_STRUCT_FIELD_PRIORITY = ["qty", "description", "part_number", "unit", "manufacturer",
                          "unit_price", "total_price", "notes", "line_no"]


def _struct_norm_header(text: str) -> str:
    return re.sub(r"\s+", " ", re.sub(r"[^a-z0-9#/. ]+", " ", str(text or "").lower())).strip()


# Tokens that must NOT be matched on their own because they are too ambiguous and
# routinely hijack the wrong column (handled only via multi-word synonyms instead).
_STRUCT_AMBIGUOUS_BARE = {"no", "no.", "#", "amount", "amt", "each", "ea", "unit", "rate",
                          "per", "make", "ref", "ref #", "model", "cat", "item", "line", "pos", "seq"}


def _struct_header_field(cell: str) -> str:
    """Return the field a header cell maps to, or '' if it is not a known header.

    Handles compound headers glued by punctuation, e.g. "STOCK NO./DESCRIPTION"
    (which contains 'description'), by matching on a slash/period-flattened form.
    A bare ambiguous token ("no", "amount", "unit", ...) is ignored on its own and
    only counts inside a longer, specific synonym.
    """
    t = _struct_norm_header(cell)
    if not t:
        return ""
    # Flatten slashes/periods to spaces so compound headers split into words.
    flat = re.sub(r"\s+", " ", re.sub(r"[/.]+", " ", t)).strip()
    tokens = set(flat.split())

    # 1) exact match on the whole normalized header (most specific).
    for field in _STRUCT_FIELD_PRIORITY:
        for syn in _STRUCT_HEADER_SYNONYMS.get(field, []):
            if t == syn or flat == re.sub(r"[/.]+", " ", syn).strip():
                return field
    # 2) whole-word / phrase match inside the flattened header. Longer synonyms first.
    for field in _STRUCT_FIELD_PRIORITY:
        for syn in sorted(_STRUCT_HEADER_SYNONYMS.get(field, []), key=len, reverse=True):
            syn_flat = re.sub(r"\s+", " ", re.sub(r"[/.]+", " ", syn)).strip()
            if not syn_flat:
                continue
            # skip ambiguous single tokens unless the header is exactly that token
            if syn_flat in _STRUCT_AMBIGUOUS_BARE and " " not in syn_flat and flat != syn_flat:
                continue
            if " " in syn_flat:
                if re.search(r"(?:^| )" + re.escape(syn_flat) + r"(?:$| )", flat):
                    return field
            else:
                if syn_flat in tokens:
                    return field
    return ""


_PRICE_VALUE_RE = re.compile(r"^\$?\s*-?\d{1,3}(?:,\d{3})*(?:\.\d+)?\s*$|^\$?\s*-?\d+(?:\.\d+)?\s*$")
_QTY_VALUE_RE = re.compile(r"^\d{1,6}(?:\.\d{1,3})?$")


def _struct_is_number(value: str) -> bool:
    v = str(value or "").strip().replace(",", "")
    if not v:
        return False
    v = v.lstrip("$").strip()
    try:
        float(v)
        return True
    except Exception:
        return False


def _struct_looks_like_total_row(cells: List[str]) -> bool:
    """True for subtotal/total/tax/shipping summary rows that are not line items."""
    joined = " ".join(c.lower() for c in cells if c).strip()
    if not joined:
        return True
    if re.match(r"^(sub\s*total|total|grand\s*total|tax|sales\s*tax|vat|freight|"
                r"shipping|handling|discount|balance|amount\s*due|terms|thank\s*you)\b",
                joined):
        # Only treat as a total row when there is no real description content beyond
        # the keyword + numbers (a real item could legitimately contain "total kit").
        non_numeric = [c for c in cells if c and not _struct_is_number(c)]
        if len(non_numeric) <= 1:
            return True
    return False


def detect_structured_header(matrix: List[List[str]], scan_rows: int = 15
                             ) -> Tuple[int, Dict[str, int]]:
    """Find the header row and a {field: column_index} map.

    Scans the first `scan_rows` non-empty rows, scoring each by how many of its
    cells map to known column names. The best-scoring row (>=2 mapped columns)
    is the header. Returns (header_row_index, column_map). header_row_index = -1
    when no usable header is found (caller then uses positional inference).
    """
    best_idx = -1
    best_map: Dict[str, int] = {}
    best_score = 0
    limit = min(len(matrix), scan_rows)
    for i in range(limit):
        cells = matrix[i]
        col_map: Dict[str, int] = {}
        mapped = 0
        numericish = 0
        for ci, cell in enumerate(cells):
            field = _struct_header_field(cell)
            if field and field not in col_map:
                col_map[field] = ci
                mapped += 1
            if _struct_is_number(cell):
                numericish += 1
        # A header row should be mostly labels, not numbers.
        score = mapped - numericish
        # Require at least a description-ish or qty anchor to call it a header.
        anchor = ("description" in col_map) or ("qty" in col_map) or ("part_number" in col_map)
        if anchor and mapped >= 2 and score > best_score:
            best_score = score
            best_idx = i
            best_map = col_map
    return best_idx, best_map


def infer_structured_columns(data_rows: List[List[str]]) -> Dict[str, int]:
    """Positional column inference when a spreadsheet has no header row.

    Uses per-column data types across the data rows:
      - the longest mostly-text column  -> description
      - a small-integer column          -> qty
      - decimal/$ columns               -> unit_price / total_price
    """
    if not data_rows:
        return {}
    width = max(len(r) for r in data_rows)
    cols: List[List[str]] = [[] for _ in range(width)]
    for r in data_rows:
        for ci in range(width):
            cols[ci].append(r[ci] if ci < len(r) else "")
    stats = []
    for ci in range(width):
        vals = [v for v in cols[ci] if str(v).strip()]
        n = max(1, len(vals))
        num = sum(1 for v in vals if _struct_is_number(v))
        qtyish = sum(1 for v in vals if _QTY_VALUE_RE.match(str(v).strip()))
        decimalish = sum(1 for v in vals if "." in str(v) or "$" in str(v))
        avg_len = sum(len(str(v)) for v in vals) / n
        alpha = sum(1 for v in vals if re.search(r"[A-Za-z]", str(v)))
        stats.append({
            "idx": ci, "count": len(vals), "num_ratio": num / n, "qtyish": qtyish,
            "decimal_ratio": decimalish / n, "avg_len": avg_len, "alpha_ratio": alpha / n,
        })
    col_map: Dict[str, int] = {}
    # description = highest avg_len among predominantly-text columns
    text_cols = [s for s in stats if s["alpha_ratio"] >= 0.5 and s["count"] > 0]
    if text_cols:
        col_map["description"] = max(text_cols, key=lambda s: s["avg_len"])["idx"]
    # qty = a mostly small-integer, low-decimal column (not the description col)
    qty_cands = [s for s in stats if s["idx"] != col_map.get("description")
                 and s["qtyish"] >= max(1, int(0.5 * s["count"])) and s["decimal_ratio"] < 0.3]
    if qty_cands:
        col_map["qty"] = min(qty_cands, key=lambda s: s["idx"])["idx"]
    # part/code = an alphanumeric column with codes (letters+digits or hyphens),
    # short-ish, not the description column. Detect before prices so a SKU like
    # "WGT-A" is not mistaken for anything else.
    code_cols = []
    for s in stats:
        if s["idx"] in col_map.values():
            continue
        vals = [str(v).strip() for v in cols[s["idx"]] if str(v).strip()]
        if not vals:
            continue
        codeish = sum(1 for v in vals if re.search(r"[A-Za-z]", v) and re.search(r"[0-9\-]", v) and " " not in v)
        if codeish >= max(1, int(0.6 * len(vals))) and s["avg_len"] <= 24:
            code_cols.append(s)
    if code_cols:
        col_map["part_number"] = min(code_cols, key=lambda s: s["idx"])["idx"]
    # price columns = numeric with decimals/$ (rightmost = total, earlier = unit)
    price_cands = [s for s in stats if s["idx"] not in col_map.values()
                   and s["num_ratio"] >= 0.6 and s["decimal_ratio"] >= 0.3]
    price_cands.sort(key=lambda s: s["idx"])
    if len(price_cands) >= 2:
        col_map["unit_price"] = price_cands[0]["idx"]
        col_map["total_price"] = price_cands[-1]["idx"]
    elif len(price_cands) == 1:
        col_map["unit_price"] = price_cands[0]["idx"]
    return col_map


class MaterialParser:
    def __init__(self, db: SmartScanDB, settings: Optional[ScanSettings] = None) -> None:
        self.db = db
        self.settings = settings or ScanSettings()

    def parse_flat_price_total_rows(self, text: str) -> List[MaterialLine]:
        """Parse clean selectable-text PDF PO lines where table columns are flattened.

        Example:
            300.00 1/2 threaded rod 0.92475 277.43 W21390,MISC,030 - MATERIA

        The text is already correct, but not separated by cell delimiters.  This
        deterministic pass extracts every quantity-led line before the AI/fuzzy
        parser can drop most rows.  It is intentionally PDF/text-oriented and does
        not affect the structured spreadsheet flow.
        """
        rows: List[MaterialLine] = []
        seen: set = set()
        price_num = r"\d{1,3}(?:,\d{3})*(?:\.\d{2,5})|\d+(?:\.\d{2,5})"
        qty_num = r"\d{1,6}(?:,\d{3})*(?:\.\d+)?"
        status_noise = re.compile(r"\b(?:subtotal|sales\s*tax|tax|total|grand\s+total|shipping|freight|authorized\s+signature|thank\s+you)\b", re.I)
        unit_words = r"ea|each|pc|pcs|piece|pieces|ft|lf|foot|feet|set|sets|box|boxes|roll|rolls|bag|bags|case|cases|kit|kits|lot|lots"

        for raw in split_lines_preserve(text or ""):
            line = normalize_space(raw)
            if not line:
                continue
            if line.startswith("---"):
                continue
            # Require a leading quantity, a description, then unit-cost + amount.
            # Keep this local to this parser so OCR/table-row cleanup patches do not
            # accidentally leave a selectable-text PDF row without a working value.
            line_work = line
            m = re.match(
                rf"^\s*(?P<qty>{qty_num})\s+(?P<body>.+?)\s+"
                rf"(?P<unit_price>{price_num})\s+(?P<amount>{price_num})"
                rf"(?:\s+(?P<job>[A-Za-z0-9][A-Za-z0-9,._/\- ]{{2,}}))?\s*$",
                line_work,
                re.I,
            )
            if not m:
                continue

            qty = normalize_qty((m.group("qty") or "").replace(",", ""))
            body = clean_description(m.group("body") or "")
            if not qty or not body:
                continue
            if status_noise.search(body):
                continue
            # Ignore obvious admin rows that happen to be numeric.
            if self.should_exclude_line(line) or is_probable_phone_or_address(body):
                continue

            unit = ""
            desc = body
            mu = re.match(rf"^(?P<unit>{unit_words})\b\s+(?P<desc>.+)$", body, re.I)
            if mu:
                unit = normalize_unit(mu.group("unit") or "")
                desc = clean_description(mu.group("desc") or "")
            desc = re.sub(r"\s+-\s*MATERIA[L]?\s*$", "", desc, flags=re.I).strip()
            desc = clean_description(desc)
            if not desc:
                continue

            part_number = ""
            for p in PART_NO_RE.findall(desc):
                p_clean = (p or "").strip().strip('"')
                pl = p_clean.lower().strip(".-_/ ")
                if SIZE_RE.fullmatch(p_clean) or pl in ELECTRICAL_KEYWORDS or pl in ELECTRICAL_UNITS:
                    continue
                if re.fullmatch(r"\d+", pl):
                    continue
                # Prefer true catalog-like tokens over ordinary words.
                if re.search(r"\d", p_clean) or (p_clean.isupper() and len(p_clean) >= 4):
                    part_number = p_clean
                    break

            notes = []
            up = (m.group("unit_price") or "").strip()
            amt = (m.group("amount") or "").strip()
            job = (m.group("job") or "").strip()
            if up:
                notes.append(f"Unit price: {up}")
            if amt:
                notes.append(f"Total: {amt}")
            if job:
                notes.append(f"Job: {job}")

            key = (qty, unit, desc.lower(), part_number.lower())
            if key in seen:
                continue
            seen.add(key)
            conf = 84
            if unit:
                conf += 4
            if contains_electrical_keyword(desc.lower()):
                conf += 5
            if SIZE_RE.search(desc):
                conf += 4
            if part_number:
                conf += 3
            rows.append(MaterialLine(
                qty=qty,
                unit=unit,
                description=desc,
                part_number=part_number,
                manufacturer=best_manufacturer_hint(desc),
                notes="; ".join(notes),
                confidence=max(40, min(98, conf)),
                source_line=line,
                source_view="text",
                ai_status="warning",
                ai_message="Parsed from flattened PDF price/amount table; please confirm.",
                highlight_color="yellow",
            ))
        return rows

    def parse(self, text: str, table_text: str = "") -> Tuple[List[MaterialLine], int]:
        # Preserve source/scan order. The review workflow depends on the user seeing
        # line 1, then line 2, etc. exactly as the document was scanned.
        raw_text = text or ""

        # v2.5.4: Image PDFs often produce a noisy TABLE / POSITIONAL section before
        # the clean OCR section.  When the OCR section itself contains a real ordered
        # material list, start parsing at "--- OCR TEXT ---" and ignore everything
        # above it.  This keeps the material list in the same order as the OCR view
        # and prevents early positional candidates from being pulled first.
        ocr_section = self._ocr_text_section(raw_text)
        use_ocr_section_only = False
        if ocr_section and ocr_section != raw_text:
            ordered_preview_count = 0
            for _raw in split_lines_preserve(ocr_section):
                if self.parse_ocr_ordered_list_line(_raw):
                    ordered_preview_count += 1
                    if ordered_preview_count >= 3:
                        break
            use_ocr_section_only = ordered_preview_count >= 3

        parse_text = ocr_section if use_ocr_section_only else raw_text
        combined = normalize_space(parse_text)
        # Preserve internal spacing for the table/header pass; use normalized lines for general parsing.
        # If we are trusting a clean OCR section, do not feed table_text/header candidates
        # into the parser because those are exactly what causes the out-of-order first pull.
        table_lines = split_lines_preserve(parse_text if use_ocr_section_only else (table_text or raw_text or ""))
        # General-pass lines come only from the primary text, NOT table_text, to avoid
        # re-parsing rows already captured by the header-aware table pass.
        lines = split_lines(combined)
        candidates: List[MaterialLine] = []
        seen_keys: dict = {}           # canonical description+qty+unit key
        seen_sources: set = set()        # exact source_line text (prevents same OCR line twice)
        seen_parts: set = set()          # distinct catalog part numbers already added

        def _norm_source(s: str) -> str:
            """Normalise a source line for dedup: lowercase, collapse whitespace/punct."""
            return re.sub(r"[^a-z0-9 ]+", " ", (s or "").lower()).strip()

        def add_row(m: Optional[MaterialLine]) -> None:
            if not m:
                return
            if self.should_exclude_line(m.source_line or m.description):
                return
            if not should_keep_material_candidate(m):
                return
            key = canonical_material_key(m)
            src_norm = _norm_source(m.source_line or m.description)
            part_key = (m.part_number or "").strip().lower()
            if key_seen_or_add(key, m, seen_keys):
                return
            if src_norm and src_norm in seen_sources:
                return
            # A distinct, real catalog part number (>=4 chars, contains a digit) should
            # only appear once.  This catches duplicates that slip past the key/source
            # checks because different parse passes cleaned qty/description differently
            # (e.g. one row kept the OCR-garbled "\2 EA ..." prefix and another stripped it).
            if part_key and len(part_key) >= 4 and re.search(r"\d", part_key) and part_key in seen_parts:
                return
            if src_norm:
                seen_sources.add(src_norm)
            if part_key and len(part_key) >= 4 and re.search(r"\d", part_key):
                seen_parts.add(part_key)
            m.scan_index = len(candidates) + 1
            candidates.append(m)

        # First pass: preserve clean OCR/email list order. This is especially important
        # for image PDFs: the OCR text can be perfect while the positional/table stream
        # can split columns into partial, out-of-order rows.
        ordered_list_count_before = len(candidates)
        for m in self.parse_ocr_ordered_list_rows(parse_text or ""):
            add_row(m)
        ordered_list_rows_found = len(candidates) - ordered_list_count_before

        # Native/selectable-text PO PDFs often flatten price tables into one line per row.
        # Parse those deterministically before the fuzzy/AI parser so clean rows are not dropped.
        flat_price_rows_found = 0
        for m in self.parse_flat_price_total_rows(parse_text or ""):
            add_row(m)
        flat_price_rows_found = len(candidates) - ordered_list_count_before - ordered_list_rows_found

        # Second pass: use explicit Qty/Quantity + Description/Item/Material headers when present.
        # If a clean OCR list produced several rows, trust that list and do NOT let the
        # noisier positional/table stream inject duplicate or out-of-order fragments.
        header_count_before = len(candidates)
        if getattr(self.settings, "enable_header_aware_extraction", True) and ordered_list_rows_found < 5:
            for m in self.parse_header_tables(table_lines):
                add_row(m)
        header_rows_found = len(candidates) - header_count_before

        # Third pass: general RFQ/material line parser for emails and informal lists.
        # If a real Qty/Description table or clean OCR ordered list was found, apply
        # STRICT gates so header/admin/address rows and positional fragments are not
        # added as false positives or duplicates.
        strict = header_rows_found >= 3 or ordered_list_rows_found >= 5 or flat_price_rows_found >= 3
        for raw in lines:
            if ordered_list_rows_found >= 5 and raw.startswith("--- TABLE / POSITIONAL CANDIDATES"):
                continue
            if getattr(self.settings, "exclude_contact_info", True) and self.should_exclude_line(raw):
                continue
            if strict:
                raw_score = material_signal_score(raw)
                # Raise gate: require very strong signal when a table already gave us rows
                if raw_score < 55 and not has_explicit_uom(raw):
                    continue
            for candidate in self.expand_possible_lines(raw):
                parsed = self.parse_line(candidate)
                if not parsed:
                    continue
                if strict:
                    # In table-heavy PDFs, fallback rows must have UOM + strong material signals.
                    if material_signal_score(parsed.description, parsed.qty, parsed.unit, parsed.part_number) < 55:
                        continue
                    if not has_explicit_uom(raw) and not parsed.unit:
                        continue
                add_row(parsed)

        # Parser-agnostic repair: fix any row where the visible line number landed in
        # Qty while the real qty+unit got pushed into the description (OCR column-gap
        # rendered as ')' / '|' / spaces).  Runs regardless of which parser produced
        # the row, so it heals scans that vary by OCR version/spacing.
        for _r in candidates:
            try:
                repair_line_number_as_qty(_r)
            except Exception:
                pass
        candidates = self.apply_correction_patterns(candidates)
        candidates = self.apply_learning_boost(candidates)
        filtered = [m for m in candidates if m.confidence >= 38]
        overall = self.overall_confidence(filtered, combined)
        return filtered, overall

    def parse_native_text_strategy(self, text: str, table_text: str = "") -> Tuple[List[MaterialLine], int]:
        """v2.5 native/text-PDF extraction path.

        The v2.6 OCR watchdog path is excellent for image PDFs, but it can be too
        aggressive when the original PDF already contains readable/selectable text.
        This method intentionally preserves the v2.5 order of operations for native
        text files/PDFs:
            1) header/table-aware extraction first
            2) general RFQ/material line parsing second
            3) no OCR-section-only override
            4) no watchdog-recovered OCR rows
        """
        combined = normalize_space(text or "")
        table_lines = split_lines_preserve(table_text or text or "")
        lines = split_lines(combined)
        candidates: List[MaterialLine] = []
        seen_keys: dict = {}
        seen_sources: set = set()
        seen_parts: set = set()

        def _norm_source(s: str) -> str:
            # v2.8.6: collapse whitespace runs too -- the same printed row arrives once
            # from the positional table lines (2-space column separators) and once from
            # the native text (single spaces); without collapsing, the two spellings
            # dodged this dedup and the item appeared twice.
            return re.sub(r"\s+", " ", re.sub(r"[^a-z0-9 ]+", " ", (s or "").lower())).strip()

        def add_row(m: Optional[MaterialLine]) -> None:
            if not m:
                return
            if self.should_exclude_line(m.source_line or m.description):
                return
            if not should_keep_material_candidate(m):
                return
            key = canonical_material_key(m)
            src_norm = _norm_source(m.source_line or m.description)
            part_key = (m.part_number or "").strip().lower()
            if key_seen_or_add(key, m, seen_keys):
                return
            if src_norm and src_norm in seen_sources:
                return
            if part_key and len(part_key) >= 4 and re.search(r"\d", part_key) and part_key in seen_parts:
                return
            if src_norm:
                seen_sources.add(src_norm)
            if part_key and len(part_key) >= 4 and re.search(r"\d", part_key):
                seen_parts.add(part_key)
            m.scan_index = len(candidates) + 1
            candidates.append(m)

        flat_price_count_before = len(candidates)
        for m in self.parse_flat_price_total_rows(text or ""):
            add_row(m)
        flat_price_rows_found = len(candidates) - flat_price_count_before

        header_count_before = len(candidates)
        if getattr(self.settings, "enable_header_aware_extraction", True):
            for m in self.parse_header_tables(table_lines):
                add_row(m)
        header_rows_found = len(candidates) - header_count_before

        strict = header_rows_found >= 3 or flat_price_rows_found >= 3
        for raw in lines:
            if getattr(self.settings, "exclude_contact_info", True) and self.should_exclude_line(raw):
                continue
            if strict:
                raw_score = material_signal_score(raw)
                if raw_score < 55 and not has_explicit_uom(raw):
                    continue
            for candidate in self.expand_possible_lines(raw):
                parsed = self.parse_line(candidate)
                if not parsed:
                    continue
                if strict:
                    if material_signal_score(parsed.description, parsed.qty, parsed.unit, parsed.part_number) < 55:
                        continue
                    if not has_explicit_uom(raw) and not parsed.unit:
                        continue
                add_row(parsed)

        # Parser-agnostic repair: fix any row where the visible line number landed in
        # Qty while the real qty+unit got pushed into the description (OCR column-gap
        # rendered as ')' / '|' / spaces).  Runs regardless of which parser produced
        # the row, so it heals scans that vary by OCR version/spacing.
        for _r in candidates:
            try:
                repair_line_number_as_qty(_r)
            except Exception:
                pass
        candidates = self.apply_correction_patterns(candidates)
        candidates = self.apply_learning_boost(candidates)
        filtered = [m for m in candidates if m.confidence >= 38]
        overall = self.overall_confidence(filtered, combined)
        return filtered, overall

    def parse_structured_rows(self, matrix: List[List[str]]) -> Tuple[List[MaterialLine], int]:
        """Parse an explicit cell grid (spreadsheet / Word table) by column mapping.

        Every data row under a detected (or inferred) column layout becomes a
        MaterialLine.  This path does NOT apply the electrical-keyword material
        gate, because being a row in a data column is itself the evidence; that is
        what lets generic spreadsheets scan to 100%.  Returns ([], 0) when the grid
        has no usable structure, so the caller can fall back to the text path.
        """
        # Drop fully-empty rows and sheet markers ("--- SHEET: X ---").
        rows: List[List[str]] = []
        for r in matrix or []:
            cells = [str(c).strip() for c in r]
            if not any(cells):
                continue
            if len(cells) == 1 and cells[0].startswith("--- SHEET:"):
                continue
            rows.append(cells)
        if len(rows) < 1:
            return [], 0

        header_idx, col_map = detect_structured_header(rows)
        if header_idx >= 0 and col_map:
            data_rows = rows[header_idx + 1:]
        else:
            # No header detected -> infer columns positionally from the data itself.
            col_map = infer_structured_columns(rows)
            data_rows = rows
        if not col_map or ("description" not in col_map and "qty" not in col_map and "part_number" not in col_map):
            return [], 0

        desc_i = col_map.get("description")
        qty_i = col_map.get("qty")
        unit_i = col_map.get("unit")
        part_i = col_map.get("part_number")
        mfg_i = col_map.get("manufacturer")
        uprice_i = col_map.get("unit_price")
        tprice_i = col_map.get("total_price")
        notes_i = col_map.get("notes")
        mapped_idx = {v for v in col_map.values() if isinstance(v, int)}

        def cell(row: List[str], idx: Optional[int]) -> str:
            if isinstance(idx, int) and 0 <= idx < len(row):
                return row[idx].strip()
            return ""

        results: List[MaterialLine] = []
        seen: set = set()
        scan_index = 0
        for row in data_rows:
            if _struct_looks_like_total_row(row):
                continue
            qty = cell(row, qty_i)
            unit = cell(row, unit_i)
            desc = cell(row, desc_i)
            part = cell(row, part_i)
            manufacturer = cell(row, mfg_i)
            uprice = cell(row, uprice_i)
            tprice = cell(row, tprice_i)
            note = cell(row, notes_i)

            # Spreadsheet status/availability flags such as N/S, BO, NLA are not
            # units of measure.  Some vendor Excel files put these flags in a
            # middle column that can be mistaken for UOM when headers are merged
            # or ambiguous.  Preserve the flag in notes, but keep Unit blank.
            status_tokens = {"N/S", "NS", "NO STOCK", "BO", "B/O", "OK TO BO", "DNS", "NLA"}
            if unit and unit.strip().upper() in status_tokens:
                note = append_note(note, unit.strip())
                unit = ""

            # If qty cell carries its unit ("10 EA"), split it.
            if qty:
                qsplit, usplit = split_qty_unit(qty)
                if qsplit:
                    qty = qsplit
                if usplit and not unit:
                    unit = usplit

            # When there is no description column, build description from the unmapped
            # text cells (skip qty/unit/price/line-number columns).
            if not desc:
                extras = [c for ci, c in enumerate(row)
                          if ci not in mapped_idx and c and not _struct_is_number(c)]
                desc = clean_description("  ".join(extras)) if extras else ""
            else:
                desc = clean_description(desc)

            # If still no description but we have a part number, use the part as the desc.
            if not desc and part:
                desc = part

            # Leading-quantity recovery: some POs put "QTY + DESCRIPTION" in ONE cell
            # (e.g. "100.0 512HD OR EQUAL", "50 3/4\" EMT MAYWEST").  Only split when:
            #   - there is no separate quantity column value, AND
            #   - the description starts with a clean standalone number (optionally a
            #     unit) followed by WHITESPACE, AND
            #   - a wordy remainder is left.
            # The whitespace requirement is critical: it prevents eating the "1" from
            # "12D1.5L" or the "3" from "3/4\"" (no space after the number there).
            if desc and not qty:
                mlead = re.match(
                    r'^\(?\s*(?P<qty>\d{1,5}(?:\.\d+)?)\s*\)?'        # standalone qty
                    r'(?:\s+(?P<unit>ea|each|pcs?|ft|feet|foot|lf|rolls?|boxes|box|'
                    r'bags?|cases?|lengths?|sticks?|sets?|pairs?|reels?|spools?|'
                    r'coils?|drums?|pallets?|gal|gallons?|kits?))?'    # optional unit word
                    r'\s+(?P<desc>\S.*)$',                             # REQUIRED space + desc
                    desc, re.I,
                )
                if mlead:
                    cand_qty = mlead.group("qty").strip()
                    cand_unit = (mlead.group("unit") or "").strip()
                    cand_desc = (mlead.group("desc") or "").strip()

                    # Spreadsheet/vendor cells sometimes append availability flags after
                    # the catalog number/description, e.g. "100.0 12D1.5L N/S".
                    # Treat that trailing flag as a note, not as UOM or part of the
                    # material description, and still split the leading quantity.
                    trailing_status = ""
                    try:
                        status_pat = r"(?:N/S|NS|NO STOCK|BO|B/O|OK TO BO|DNS|NLA)"
                        sm = re.search(r"\s+(?P<status>" + status_pat + r")\s*$", cand_desc, re.I)
                        if sm:
                            trailing_status = sm.group("status").strip()
                            cand_desc = cand_desc[:sm.start()].strip()
                    except Exception:
                        trailing_status = ""

                    # Require real material content to remain.  This now accepts both
                    # wordy descriptions ("3/4 EMT MAYWEST") and catalog-style codes
                    # ("12D1.5L", "512HD OR EQUAL") instead of only >=2-letter words.
                    has_catalogish_desc = bool(
                        re.search(r"[A-Za-z]{2,}", cand_desc)
                        or re.search(r"\d+[A-Za-z][A-Za-z0-9./-]*", cand_desc)
                        or re.search(r"[A-Za-z][A-Za-z0-9./-]*\d+", cand_desc)
                    )
                    if cand_qty and has_catalogish_desc:
                        qty = cand_qty
                        if cand_unit and not unit:
                            unit = cand_unit
                        desc = clean_description(cand_desc)
                        if trailing_status:
                            note = append_note(note, trailing_status)

            qty = normalize_qty(qty) if qty else ""
            unit = normalize_unit(unit) if unit else ""
            part = clean_part_number(part) if part else ""
            manufacturer = clean_description(manufacturer) if manufacturer else ""

            # A real line item needs at least a description or a part number.  A bare
            # quantity with nothing else is a stray cell, not an item.
            if not desc and not part:
                continue
            # Structured spreadsheet rows are high-trust: if a row sits under a
            # Qty + Description/Stock header, catalog-only descriptions like
            # "512HD OR EQUAL" or "16D1.5L" are still real material lines even
            # though generic OCR/PDF noise filters may not see electrical words.
            # Only drop unmistakable contact/address/web noise here.
            source_probe = " ".join(c for c in row if c)
            if is_probable_phone_or_address(source_probe) or EMAIL_RE.search(source_probe) or URL_RE.search(source_probe):
                continue
            if not qty and (is_definite_admin_noise(desc) or is_probable_phone_or_address(desc)):
                continue

            note_bits = []
            if uprice:
                note_bits.append(f"Unit price: {uprice}")
            if tprice:
                note_bits.append(f"Total: {tprice}")
            if note:
                note_bits.append(note)
            # Capture short PO flags (e.g. "N/S" no-stock, "BO", "OK TO BO") that sit in
            # an unmapped cell, so the signal is not lost.
            for ci, c in enumerate(row):
                if ci in mapped_idx or not c:
                    continue
                cu = c.strip().upper()
                if cu in status_tokens and cu not in " ".join(note_bits).upper():
                    note_bits.append(c.strip())
            notes = "; ".join(note_bits)

            # Confidence: structured rows are trustworthy.  Start high; nudge up for
            # richer rows, down a little for description-only rows.
            conf = 88
            if qty:
                conf += 4
            if part:
                conf += 4
            if unit:
                conf += 2
            if not qty and not part:
                conf -= 18
            conf = max(40, min(100, conf))

            source_line = "   ".join(c for c in row if c)
            key = (qty, unit, desc.lower(), part.lower())
            if key in seen:
                continue
            seen.add(key)
            scan_index += 1
            results.append(MaterialLine(
                qty=qty,
                unit=unit,
                description=desc,
                part_number=part,
                manufacturer=manufacturer,
                notes=notes,
                confidence=conf,
                source_line=source_line,
                scan_index=scan_index,
                source_view="table",
            ))

        if not results:
            return [], 0
        results = self.apply_correction_patterns(results)
        results = self.apply_learning_boost(results)
        overall = self.overall_confidence(results, "\n".join(r.source_line for r in results))
        # Structured extraction is high-trust; keep overall confidence from being
        # dragged down by the text-oriented scorer.
        overall = max(overall, 80)
        return results, overall

    def expand_possible_lines(self, raw: str) -> Iterable[str]:
        raw = BULLET_RE.sub("", raw.strip())
        if not raw:
            return []
        # Some emails paste multiple material entries separated by semicolons.
        parts = re.split(r"\s*;\s*", raw)
        out = []
        for p in parts:
            p = p.strip(" -–—\t")
            if p:
                out.append(p)
        return out

    def parse_ocr_ordered_list_rows(self, text: str) -> List[MaterialLine]:
        """Parse OCR/email list lines in the exact top-to-bottom text order.

        Image PDFs often OCR into very clean lines such as:
            -(3 boxes) TOPAZ 170's
            €400') 3/4" EMT
            -(20) 3/4" EMT to GF SS changeovers

        The positional/table stream for those same image PDFs can be noisier and
        may split columns out of order.  This parser intentionally trusts the
        OCR text order for parenthesized/bulleted material lists before the
        table/header fallback runs.
        """
        rows: List[MaterialLine] = []
        seen_keys: dict = {}
        seen_parts: set = set()
        for raw in split_lines_preserve(text or ""):
            parsed = self.parse_ocr_ordered_list_line(raw)
            if not parsed:
                continue
            key = canonical_material_key(parsed)
            part_key = (parsed.part_number or "").strip().lower()
            # A distinct, real catalog part number appearing twice is a duplicate row
            # (common when OCR/table streams both catch the same line).
            if key_seen_or_add(key, parsed, seen_keys):
                continue
            if part_key and len(part_key) >= 4 and part_key in seen_parts:
                continue
            if part_key and len(part_key) >= 4:
                seen_parts.add(part_key)
            rows.append(parsed)
        return rows

    def parse_ocr_ordered_list_line(self, raw: str) -> Optional[MaterialLine]:
        line = (raw or "").strip()
        if not line or self.is_header_or_noise(line) or self.should_exclude_line(line):
            return None
        # OCR sometimes inserts quotes/hash marks, vertical table separators, or
        # bracket fragments between the visible Line #, Qty, and Unit columns.
        # Normalize only for parsing; keep original source_line for review/highlight.
        # First: the thin vertical rule right after the LINE number is often misread as
        # stray punctuation glued to it ("7?", "5:", "13;", "8!", "4|").  Strip any such
        # punctuation that sits between the leading line number and the following space,
        # so the real quantity is not lost.  A real qty never looks like "7?", and
        # decimals/fractions ("2.5", "1/2") have a digit (not a space) after the mark.
        line_work = re.sub(r"^(\s*\d{1,3})[?!:;,)\].·•|*\"'`’”]+(?=\s)", r"\1", line)
        line_work = re.sub(r"^(\s*\d{1,3}[\.)]?\s+)[\'\"`´‘’“”#_]+(?=\d)", r"\1", line_work)
        # OCR can also insert stray quote/hash/underscore marks BETWEEN the real
        # quantity and the unit, e.g. `8 4 “SET tool...`.  Strip those only when
        # they sit after a numeric qty and immediately before a unit-looking word.
        # This keeps the visible line number from being mistaken for the qty.
        line_work = re.sub(r"(?<=\d)\s+[\'\"`´‘’“”#_]+(?=[A-Za-z])", " ", line_work)
        line_work = re.sub(r"[│¦]+", "|", line_work)
        # Some scanners render the vertical column rules as PERIODS or COMMAS, e.g.
        # "2 2,000 . EA | Deep kindorf ..." or "31 250 , FT . 10 THHN ...".
        # Collapse a period/comma that is used as a column gap — punctuation SURROUNDED
        # by spaces.  This is safe for real decimals ("2.5") and thousands separators
        # ("2,000") because those have no spaces around the dot/comma.
        line_work = re.sub(r"\s+[.,]\s+", " ", line_work)
        # A second pass catches back-to-back gaps like "10 , EA . 1 inch" -> "10 EA 1 inch".
        line_work = re.sub(r"\s+[.,]\s+", " ", line_work)
        # Strip a leading lone separator (dot/pipe/comma/degree) before the line number.
        line_work = re.sub(r"^\s*[.\|,°]\s*", "", line_work)
        # Strip trailing border-artifact characters OCR adds after the row
        # (e.g. a stray "7", "-", "|", ",", "." from the table's right rule).  A
        # space is required before "." so real decimals are never touched.
        line_work = re.sub(r"(?:\s+[-|,.]\s*)+$", "", line_work)
        # OCR garble cleanups commonly seen on scanned PO tables:
        #  - "20-12 EA"  : line number hyphen-glued to qty -> "20 12 EA"
        line_work = re.sub(r"^(\s*\d{1,3})\s*-\s*(?=\d)", r"\1 ", line_work)
        #  - garbled "EA" unit: "eal", "ea1", "eaj", "ea|", "100 eal B 12 ..." -> "EA"
        #    Only fix when a standalone EA-like token sits between a number and text.
        line_work = re.sub(r"(?<=\d)\s+ea[l1ij|]\b", " EA", line_work, flags=re.I)
        line_work = re.sub(r"\bea[l1ij]\b", "EA", line_work, flags=re.I)
        line_work = re.sub(r"\s*\|\s*", " | ", line_work)
        line_work = re.sub(r"\s+", " ", line_work).strip()

        # Ignore section labels added by the merger.
        if re.match(r"^---\s*(?:ocr text|native text|table|ocr page|page)\b", line, re.I):
            return None

        # Common OCR/bullet patterns:
        #   -(3 boxes) TOPAZ 170's
        #   -(1box) 3/4" EMT straps
        #   (20) 3/4" EMT connectors
        #   €400') 3/4" EMT        # OCR junk before a 400' quantity
        #   ((30') 3/4" Silver Strut
        #   {250') 1-1/4" EMT
        unit_words = (
            r"ea|each|pcs?|pieces?|ft|feet|foot|lf|rolls?|boxes|box|bags?|cases?|"
            r"coils?|buckets?|dozens?|lengths?|sticks?|sets?|lots?|pairs?|reels?|spools?|packs?|pkgs?"
        )

        # OCR of picture/PDF tables often includes the visible "Line" column before
        # the real quantity, e.g. "10 50 EA 3/4 EMT conduit" OR "10 | 120 FT 3/4 EMT".
        # The Line/Qty boundary may OCR as a space, a pipe, or a bracket, and that
        # separator is what previously broke detection on ~half the rows (qty became
        # the line number).  This pattern accepts ANY of those separators between the
        # line number and the quantity, and between the quantity and the unit.
        sep = r"(?:\s*[\|\]\)\[\(]\s*|\s+)"
        line_qty_match = re.match(
            rf"^\s*(?P<line_no>\d{{1,3}})[\.)]?{sep}"
            rf"[\'\"`´‘’“”#_]*"
            rf"(?P<qty>\d{{1,6}}(?:,\d{{3}})*(?:\.\d+)?){sep}"
            rf"[\'\"`´‘’“”#_]*"
            rf"(?P<unit>{unit_words}|['’])\b"
            rf"{sep}?"
            rf"(?P<desc>.+?)\s*$",
            line_work,
            re.I,
        )
        if line_qty_match:
            try:
                line_no_val = int((line_qty_match.group("line_no") or "0").strip())
            except Exception:
                line_no_val = 0
            # Conservative guard: row numbers are usually small visible line ids.
            # If the first number is huge, it is more likely a real quantity.
            if 0 < line_no_val <= 250:
                qty = normalize_qty(line_qty_match.group("qty") or "")
                raw_unit = (line_qty_match.group("unit") or "").strip()
                desc = clean_description(line_qty_match.group("desc") or "")
                desc = re.sub(r"^[\s\|\]\)_:;.,-]+", "", desc).strip()
                m = None
            else:
                m = None
        else:
            m = None

        qty_needs_review = False
        if not line_qty_match or (line_qty_match and 'qty' not in locals()):
            m = re.match(
                rf"^\s*(?:[-*•●▪]\s*)?(?:[^\w\s]{{0,4}}\s*)?[A-Za-z]?\s*[\(\[\{{]*\s*[A-Za-z]?\s*"
                rf"(?P<qty>\d+(?:\.\d+)?)\s*(?P<unit>{unit_words}|['’])?\s*"
                rf"[\)\]\}}'’]*\s*(?P<desc>.+?)\s*$",
                line,
                re.I,
            )
            if not m:
                # Garbled-quantity recovery: the leading Line/Qty column OCR'd as
                # non-digits (e.g. "S EA ...", "L? | EA ...", "\2 EA ...").  Rather than
                # drop the row (the user then can't see or fix it) or grab a wrong
                # number, capture Unit + Description (+ Part#) and flag the quantity
                # for review.  A real UOM token immediately before the description is
                # the anchor that tells us this is a line item with a damaged qty cell.
                mg = re.match(
                    rf"^\s*[^A-Za-z0-9]*[A-Za-z0-9?\\/]{{0,4}}{sep}?"
                    rf"(?P<unit>{unit_words})\b{sep}?(?P<desc>.+?)\s*$",
                    line_work, re.I,
                )
                if mg:
                    g_unit = normalize_unit(mg.group("unit") or "")
                    g_desc = clean_description(mg.group("desc") or "")
                    g_desc = re.sub(r"^[\s\|\]\)_:;.,-]+", "", g_desc).strip()
                    if g_unit and g_desc and (contains_electrical_keyword(g_desc.lower())
                                              or SIZE_RE.search(g_desc) or PART_NO_RE.search(g_desc)):
                        qty = ""
                        raw_unit = g_unit
                        desc = g_desc
                        qty_needs_review = True
                        m = None
                    else:
                        return None
                else:
                    return None
            else:
                qty = normalize_qty(m.group("qty") or "")
                raw_unit = (m.group("unit") or "").strip()
                desc = clean_description(m.group("desc") or "")
                # Guard against the line-number-as-qty trap when the REAL quantity was
                # OCR-garbled: e.g. "13 |S EA NEMA..." or "18) |\2 EA single device...".
                # Here the generic match grabbed the leading line number as qty and left
                # "<garble> EA <desc>" in the description.  Detect a UOM word sitting near
                # the front of that leftover description and, if found, treat the leading
                # number as the line id, blank the qty, and flag it for review.
                lead_no_m = re.match(r"^\s*(?P<no>\d{1,3})[\.)]?\s*[\|\]\)]?", line_work)
                desc_unit_m = re.match(
                    rf"^\s*[^A-Za-z0-9]*[A-Za-z0-9?\\/]{{0,4}}\s*[\|\]\)]?\s*"
                    rf"(?P<unit>{unit_words})\b\s*(?P<rest>.+)$",
                    desc, re.I,
                )
                if (lead_no_m and desc_unit_m and qty == lead_no_m.group("no")
                        and 0 < int(lead_no_m.group("no")) <= 250):
                    g_unit = normalize_unit(desc_unit_m.group("unit") or "")
                    g_rest = clean_description(desc_unit_m.group("rest") or "")
                    if g_unit and g_rest and (contains_electrical_keyword(g_rest.lower())
                                              or SIZE_RE.search(g_rest) or PART_NO_RE.search(g_rest)):
                        qty = ""
                        raw_unit = g_unit
                        desc = g_rest
                        qty_needs_review = True
        if not desc:
            return None
        if not qty and not qty_needs_review:
            return None

        if raw_unit in {"'", "’"}:
            unit = "FT"
        else:
            unit = normalize_unit(raw_unit)

        # Patterns like "-(3) Cases of Water" put the unit at the front of the
        # description instead of inside the parentheses. Recover it here.
        if not unit:
            unit_at_front = re.match(rf"^\s*(?P<unit>{unit_words})\b\s*(?P<rest>.*)$", desc, re.I)
            if unit_at_front:
                possible_unit = normalize_unit(unit_at_front.group("unit") or "")
                rest = (unit_at_front.group("rest") or "").strip()
                if possible_unit and rest:
                    unit = possible_unit
                    desc = rest

        # If the unit was written as "cases/boxes/bags of ...", keep the real
        # description clean instead of "of Water".
        desc = re.sub(r"^\s*of\s+", "", desc, flags=re.I).strip()

        # A quantity in feet is often written as "400' 3/4 EMT" and has no unit word.
        if not unit and re.search(r"\d+\s*['’]\s*", line):
            unit = "FT"

        # If the visible line number/qty/UOM leaked into the description, remove it
        # after qty/unit are known. This fixes image-PDF OCR rows like
        # "15 '100 FT 12 MC cable..." becoming Qty=100, Unit=FT, Desc=12 MC cable.
        desc = clean_ocr_table_description(desc, qty, unit)

        # Strong but not reckless scoring: parenthesized/bulleted OCR list items
        # are likely material rows, but still require material signal, a known
        # count/length unit, or a part/catalog-looking all-caps token.
        part_number = ""
        part_matches = PART_NO_RE.findall(desc)
        clean_parts = []
        for p in part_matches:
            pl = p.lower().strip(".-_/ ")
            p_clean = (p or "").strip().strip('"')
            if SIZE_RE.fullmatch(p_clean) or re.fullmatch(r"\d+\s*-\s*\d+\s*/\s*\d+(?:\s*[\"'])?", p_clean):
                continue
            if re.fullmatch(r"[A-Za-z][A-Za-z-]{2,18}", p_clean) and not (p_clean.isupper() and "-" not in p_clean and len(p_clean) >= 4):
                continue
            if pl in ELECTRICAL_KEYWORDS or pl in ELECTRICAL_UNITS:
                continue
            if re.fullmatch(r"\d+", pl):
                continue
            clean_parts.append(p)
        if clean_parts:
            part_number = clean_parts[0]

        conf = 54
        if unit:
            conf += 12
        if contains_electrical_keyword(desc.lower()):
            conf += 24
        if SIZE_RE.search(desc):
            conf += 10
        if part_number:
            conf += 6
        if not (looks_materialish(desc) or unit or part_number):
            conf -= 26

        row = MaterialLine(
            qty=qty,
            unit=unit,
            description=desc,
            part_number=part_number,
            manufacturer="",
            notes="OCR ordered list extraction",
            confidence=max(0, min(100, conf)),
            source_line=line,
        )
        if qty_needs_review:
            # The quantity cell was OCR-corrupted; surface the row for a quick human
            # fix rather than dropping it or inventing a number.
            row.ai_status = "warning"
            row.ai_message = "Quantity could not be read from this scanned line; please enter it."
            row.highlight_color = "yellow"
            row.notes = append_note(row.notes, "Qty needs review (OCR unreadable)")
            return row
        if not should_keep_material_candidate(row) and not (unit and qty and desc):
            return None
        return row

    def _ocr_text_section(self, text: str) -> str:
        """Return the portion of merged text that begins at the OCR TEXT section.

        The merged viewer is normally: TABLE / POSITIONAL, NATIVE TEXT, then OCR TEXT.
        For image-based PDFs, the OCR TEXT section is usually the cleanest source and
        should be treated as the start of extraction when it contains ordered list rows.
        """
        text = text or ""
        m = re.search(r"^---\s*OCR\s+TEXT\s*---\s*$", text, flags=re.I | re.M)
        if not m:
            return text
        return text[m.start():]

    def parse_header_tables(self, lines: List[str]) -> List[MaterialLine]:
        rows: List[MaterialLine] = []
        for i, header in enumerate(lines):
            if not self.is_qty_description_header(header):
                continue
            header_kind, schema = self.build_header_schema(header)
            if not schema or "qty" not in schema or "desc" not in schema:
                continue
            # Read down until a clear new section starts. Since blank lines are removed,
            # use a conservative window plus stop phrases.
            stop_after_noise = 0
            pending: List[Tuple[str, Optional[MaterialLine]]] = []
            for raw in lines[i + 1 : i + 75]:
                low = raw.lower().strip()
                if self.is_qty_description_header(raw):
                    break
                if re.match(r"^(subtotal|total|tax|terms|notes?|thank you|signature|prepared by|quote total)\b", low):
                    break
                if self.should_exclude_line(raw):
                    continue
                m = self.parse_structured_row(raw, header_kind, schema)
                pending.append((raw, m))
                if m:
                    stop_after_noise = 0
                else:
                    stop_after_noise += 1
                    # Once the table started, several non-material rows usually means the table ended.
                    if any(pm for _, pm in pending) and stop_after_noise >= 6:
                        break
            strict_hits = sum(1 for _, pm in pending if pm)
            # v2.8.6: relaxed second pass.  Three or more strict rows CONFIRM this is a
            # real line-item table -- so a sibling row under the same header with a
            # clean quantity and description is a line item even when its wording has
            # no electrical keyword or catalog token ("200  N 95 mask", "50  Buckets
            # Scrubs hand wipes").  The material-signal gate exists to reject pseudo-
            # tables in free text; once the table itself is proven, that gate was
            # dropping this business's normal safety/consumable lines.  The relaxed
            # pass still runs every noise/contact/admin check -- it only waives the
            # keyword score, exactly like the spreadsheet structured path does.
            for raw, m in pending:
                if m is None and strict_hits >= 3:
                    m = self.parse_structured_row(raw, header_kind, schema, relaxed=True)
                if m:
                    rows.append(m)
        return rows

    def is_qty_description_header(self, line: str) -> bool:
        low = line.lower()
        if not HEADER_QTY_RE.search(low):
            return False
        if not HEADER_DESC_RE.search(low):
            return False
        # Header-aware column extraction needs visible column separation; otherwise the
        # regular line parser handles "Quantity Description" style lists more safely.
        return has_table_spacing(line) or line.count("|") >= 1 or "\t" in line

    def build_header_schema(self, header: str) -> Tuple[str, Dict[str, Any]]:
        # Separator-mode: common in extracted PDFs: Qty  Description  Part #  Mfg
        if "|" in header or "\t" in header or re.search(r"\S\s{2,}\S", header):
            cells = split_table_cells(header)
            schema: Dict[str, Any] = {"cells": []}
            for idx, cell in enumerate(cells):
                label = self.header_label(cell)
                # v2.8.6: first column wins per label -- duplicate labels ("UOM" then
                # "UNIT PRICE", or two price columns) were overwriting the earlier,
                # correct column index with a rightward one.
                if label and label not in schema:
                    schema[label] = idx
                schema["cells"].append((label or "other", cell))
            return "cells", schema

        # Position-mode: use header word x/character positions and slice following lines by those ranges.
        labels: List[Tuple[int, str]] = []
        for label, rgx in (("qty", HEADER_QTY_RE), ("desc", HEADER_DESC_RE), ("unit", HEADER_UNIT_RE), ("part", re.compile(r"\b(?:part|part\s*#|part\s*no|catalog|cat\.?\s*no\.?)\b", re.I)), ("mfg", HEADER_MFG_RE)):
            m = rgx.search(header)
            if m:
                labels.append((m.start(), label))
        labels = sorted(set(labels))
        if len(labels) < 2:
            return "", {}
        schema = {"positions": labels, "header_len": len(header)}
        for pos, label in labels:
            schema[label] = pos
        return "positions", schema

    def header_label(self, text: str) -> str:
        t = text.lower().strip()
        if LINE_NO_HEADER_RE.search(t):
            return "line_no"
        # v2.8.6: price columns FIRST -- "UNIT PRICE" / "EXT PRICE" contain the word
        # "unit"/"amount" and were being labeled as the UOM/qty column, which put
        # "$0.85" in Unit once header-table extraction started working on native PDFs.
        if PRICE_HEADER_RE.search(t) and not re.fullmatch(r"(?:unit|uom|um|amount)\.?\s*", t):
            return "price"
        if HEADER_QTY_RE.search(t):
            return "qty"
        if HEADER_UNIT_RE.search(t):
            return "unit"
        if HEADER_MFG_RE.search(t):
            return "mfg"
        if re.search(r"\b(?:part|part\s*#|part\s*no|part\s*number|catalog|cat\.?\s*no\.?)\b", t, re.I):
            return "part"
        if HEADER_DESC_RE.search(t):
            return "desc"
        if PRICE_HEADER_RE.search(t):
            return "price"
        return ""

    def parse_structured_row(self, raw: str, header_kind: str, schema: Dict[str, Any], relaxed: bool = False) -> Optional[MaterialLine]:
        qty = unit = desc = part = manufacturer = ""
        if header_kind == "cells":
            cells = split_table_cells(raw)
            if len(cells) < 2:
                return None
            def get(label: str) -> str:
                idx = schema.get(label)
                if isinstance(idx, int) and idx < len(cells):
                    return cells[idx].strip()
                return ""
            qty = get("qty")
            unit = get("unit")
            desc = get("desc")
            part = get("part")
            manufacturer = get("mfg")
            # If there is no explicit unit column, allow QTY cell like "10 EA".
            qty, unit2 = split_qty_unit(qty)
            if unit2 and not unit:
                unit = unit2
            # v2.8.6: a mis-mapped column can hand a money value to Unit; a price is
            # never a unit of measure.
            if unit and is_price_cell(unit):
                unit = ""
            unit = normalize_unit(unit)
            unit = normalize_unit(unit)
            if not desc:
                # Combine non-price cells excluding row-number/qty/unit/price into a safe description.
                ignore = {schema.get("line_no"), schema.get("qty"), schema.get("unit"), schema.get("price")}
                desc = best_description_from_cells(cells, {x for x in ignore if isinstance(x, int)})
        elif header_kind == "positions":
            pieces = slice_by_header_positions(raw, schema.get("positions", []))
            qty = pieces.get("qty", "")
            unit = pieces.get("unit", "")
            desc = pieces.get("desc", "")
            part = pieces.get("part", "")
            manufacturer = pieces.get("mfg", "")
            qty, unit2 = split_qty_unit(qty)
            if unit2 and not unit:
                unit = unit2
        else:
            return None

        qty = normalize_qty(qty)
        desc = clean_description(desc)
        part = clean_part_number(part)
        manufacturer = clean_description(manufacturer)

        # When columns get shifted, recover from rows like "10 3/4 EMT".
        if not qty and desc:
            fallback = self.parse_line(desc)
            if fallback:
                fallback.source_line = raw
                fallback.confidence = min(100, fallback.confidence + 8)
                fallback.notes = append_note(fallback.notes, "Recovered from Qty/Description table row")
                return fallback

        if not qty or not desc:
            return None
        if self.should_exclude_line(raw) or is_definite_admin_noise(raw):
            return None
        if is_probable_phone_or_address(desc) or is_definite_admin_noise(desc):
            return None

        # A structured header match is strong evidence, but still require an actual material signal.
        conf = 62
        low = desc.lower()
        if contains_electrical_keyword(low):
            conf += 22
        if SIZE_RE.search(desc):
            conf += 12
        if part or PART_NO_RE.search(desc):
            conf += 8
        if unit.lower() in ELECTRICAL_UNITS or has_explicit_uom(raw):
            conf += 10
        sig = material_signal_score(desc, qty, unit, part)
        if relaxed:
            # v2.8.6: the table is already CONFIRMED (3+ strict sibling rows), the row
            # has a clean qty + description, and every noise/contact/admin check above
            # passed.  Waive only the keyword-score gate; a data row in a proven
            # line-item table is a line item regardless of its wording.
            if sig < 0 or len(re.sub(r"[^A-Za-z]", "", desc)) < 3:
                return None
            return MaterialLine(
                qty=qty,
                unit=unit,
                description=desc,
                part_number=part,
                manufacturer=manufacturer,
                notes="Header-aware Qty/Description extraction (table-confirmed row)",
                confidence=max(0, min(100, max(conf - 8, 46))),
                source_line=raw,
            )
        if sig < 38:
            return None
        if not looks_materialish(desc) and not part and sig < 50:
            conf -= 28
        if conf < 42:
            return None
        return MaterialLine(
            qty=qty,
            unit=unit,
            description=desc,
            part_number=part,
            manufacturer=manufacturer,
            notes="Header-aware Qty/Description extraction",
            confidence=max(0, min(100, conf)),
            source_line=raw,
        )

    def parse_numbered_table_row(self, line: str) -> Optional[MaterialLine]:
        """Handle PDF/OCR table rows where the first number is a row/line count.

        Example from RFQs/vendor quotes:
            7   25   EA   -   -   3/4 emt kindorf straps   $0.00  $0.00

        A normal quantity-first parser would incorrectly read qty=7. This method
        recognizes the row-number + quantity + unit pattern and returns qty=25,
        unit=EA, description=3/4 emt kindorf straps.
        """
        raw = (line or "").strip()
        if not raw:
            return None

        cells = split_table_cells(raw)
        if len(cells) >= 4 and is_likely_line_number_token(cells[0]):
            q, u = split_qty_unit(cells[1])
            if q and (u or (len(cells) > 2 and normalize_unit(cells[2]))):
                unit = u or normalize_unit(cells[2])
                if unit:
                    skip = {0, 1, 2}
                    desc = best_description_from_cells(cells, skip)
                    part = ""
                    # If there are catalog/part cells before the description, keep the cleanest one.
                    for c in cells[3:-1]:
                        pn = clean_part_number(c)
                        if pn and not is_price_cell(c):
                            part = pn
                            break
                    desc = clean_description(strip_prices_from_text(desc))
                    if desc and not is_probable_phone_or_address(desc):
                        conf = 66
                        if contains_electrical_keyword(desc.lower()):
                            conf += 22
                        if SIZE_RE.search(desc):
                            conf += 10
                        if part:
                            conf += 6
                        return MaterialLine(
                            qty=q,
                            unit=unit,
                            description=desc,
                            part_number=part,
                            manufacturer="",
                            notes="Detected row-number + Qty + UOM table pattern",
                            confidence=max(0, min(100, conf)),
                            source_line=raw,
                        )

        # Fallback for normalized lines where PDF spacing was lost but the row starts
        # with: row_number qty unit description...
        m = ROWNUM_QTY_UNIT_RE.match(raw)
        if m and is_likely_line_number_token(m.group("rownum")):
            qty = normalize_qty(m.group("qty"))
            unit = normalize_unit(m.group("unit"))
            rest = (m.group("rest") or "").strip()
            # Drop placeholder/catalog cells at the front when they are separated by spaces.
            tokens = rest.split()
            while tokens and tokens[0] in {"-", "--", "—", "–"}:
                tokens.pop(0)
            # If the next token is another placeholder or pure price, skip it too.
            while tokens and is_price_cell(tokens[0]):
                tokens.pop(0)
            desc = clean_description(strip_prices_from_text(" ".join(tokens)))
            if desc and not is_probable_phone_or_address(desc):
                conf = 62
                if contains_electrical_keyword(desc.lower()):
                    conf += 22
                if SIZE_RE.search(desc):
                    conf += 10
                if not looks_materialish(desc):
                    conf -= 15
                if conf >= 38:
                    return MaterialLine(
                        qty=qty,
                        unit=unit,
                        description=desc,
                        part_number="",
                        manufacturer="",
                        notes="Detected row-number + Qty + UOM text pattern",
                        confidence=max(0, min(100, conf)),
                        source_line=raw,
                    )
        return None

    def parse_line(self, line: str) -> Optional[MaterialLine]:
        original = line.strip()
        if not original or len(original) < 3:
            return None
        if self.is_header_or_noise(original) or self.should_exclude_line(original) or is_definite_admin_noise(original):
            return None

        numbered = self.parse_numbered_table_row(original)
        if numbered:
            return numbered

        qty = ""
        unit = ""
        desc = original
        notes = ""
        conf = 0

        m = QTY_LEADING_RE.match(original)
        if m:
            qty = normalize_qty((m.group("qty") or "").strip("()"))
            unit = normalize_unit(m.group("unit") or "")
            desc = (m.group("desc") or "").strip()
            conf += 20
        else:
            mt = QTY_TRAILING_RE.match(original)
            if mt and looks_materialish(original):
                qty = normalize_qty((mt.group("qty") or "").strip())
                unit = normalize_unit(mt.group("unit") or "")
                desc = (mt.group("desc") or "").strip()
                conf += 15

        # Explicit embedded labels, e.g. "Qty: 4 Description: 4S Deep Box".
        label_match = re.search(r"\b(?:qty|quantity)\s*[:#\-]?\s*(?P<qty>\d+(?:\.\d+)?)\s*(?P<unit>[A-Za-z]+)?\s+.*?\b(?:description|desc|material|item)\s*[:#\-]?\s*(?P<desc>.+)$", original, re.I)
        if label_match:
            qty = normalize_qty(label_match.group("qty") or qty)
            maybe_unit = (label_match.group("unit") or "").strip()
            if maybe_unit.lower() in ELECTRICAL_UNITS and not unit:
                unit = normalize_unit(maybe_unit)
            desc = (label_match.group("desc") or desc).strip()
            conf += 25
            notes = append_note(notes, "Read explicit Qty/Description labels")

        desc = clean_description(desc)
        if not desc or is_probable_phone_or_address(desc) or is_definite_admin_noise(desc):
            return None
        low = desc.lower()
        if contains_electrical_keyword(low):
            conf += 30
        if SIZE_RE.search(desc):
            conf += 15
        if qty:
            conf += 15
        if unit.lower() in ELECTRICAL_UNITS:
            conf += 8
        if has_table_spacing(original):
            conf += 5

        part_number = ""
        part_matches = PART_NO_RE.findall(desc)
        clean_parts = []
        for p in part_matches:
            pl = p.lower().strip(".-_/ ")
            # Do not treat ordinary words from the description as part numbers.
            # OCR/list scans were putting words like Water, Silver, Steel, nuts,
            # and washers into the Part # column. Keep true catalog-looking
            # tokens: mixed alpha/numeric, punctuation, or all-uppercase brands.
            p_clean = (p or "").strip().strip('"')
            if SIZE_RE.fullmatch(p_clean) or re.fullmatch(r"\d+\s*-\s*\d+\s*/\s*\d+(?:\s*[\"'])?", p_clean):
                continue
            if re.fullmatch(r"[A-Za-z][A-Za-z-]{2,18}", p_clean) and not (p_clean.isupper() and "-" not in p_clean and len(p_clean) >= 4):
                continue
            if pl in ELECTRICAL_KEYWORDS or pl in ELECTRICAL_UNITS:
                continue
            if pl in {"need", "quote", "price", "please", "thanks", "phone", "fax", "email", "address", "deep", "white", "black", "gray", "grey", "conduit", "boxes", "box", "receptacle", "connector", "coupling", "breaker", "switch", "cover", "plate", "wire", "cable"}:
                continue
            if re.fullmatch(r"\d+", pl):
                continue
            clean_parts.append(p)
        if clean_parts:
            part_number = clean_parts[0]
            conf += 8

        manufacturer = ""
        for mf in MANUFACTURER_HINTS:
            if mf in low:
                manufacturer = mf.title()
                conf += 8
                break

        if material_signal_score(desc, qty, unit, part_number) < 38:
            return None
        if conf < 42 and not looks_materialish(original):
            return None

        conf = max(0, min(100, conf))
        return MaterialLine(
            qty=qty,
            unit=unit,
            description=desc,
            part_number=part_number,
            manufacturer=manufacturer,
            notes=notes,
            confidence=conf,
            source_line=original,
        )

    def should_exclude_line(self, line: str) -> bool:
        if not getattr(self.settings, "exclude_contact_info", True):
            return False
        if is_definite_admin_noise(line):
            return True
        return is_probable_contact_line(line, strict_address=getattr(self.settings, "strict_address_filter", True))

    def is_header_or_noise(self, line: str) -> bool:
        low = line.lower().strip()
        noise = {
            "qty", "quantity", "description", "desc", "item", "material", "materials", "part", "part number",
            "unit", "uom", "price", "extension", "total", "subtotal", "tax", "quote", "rfq", "request for quote",
        }
        if low in noise:
            return True
        if self.is_qty_description_header(line):
            return True
        if len(low) > 180 and not contains_electrical_keyword(low):
            return True
        if re.match(r"^(from|to|subject|sent|date):", low):
            return True
        if low.startswith("http") or "@" in low and len(low.split()) <= 3:
            return True
        return False

    def apply_correction_patterns(self, rows: List[MaterialLine]) -> List[MaterialLine]:
        """Apply previously-confirmed parse patterns only on very similar source lines."""
        try:
            patterns = self.db.recent_correction_patterns(limit=400)
        except Exception:
            patterns = []
        if not patterns:
            return rows
        for row in rows:
            src = normalize_space(row.source_line or "")
            if not src:
                continue
            sig = normalize_source_signature(src)
            best = None
            best_score = 0.0
            for pat in patterns:
                psig = pat.get("source_signature") or ""
                pline = pat.get("source_line") or ""
                if not psig and not pline:
                    continue
                score = max(
                    SequenceMatcher(None, sig, psig).ratio() if psig else 0.0,
                    SequenceMatcher(None, src.lower(), pline.lower()).ratio() if pline else 0.0,
                )
                if score > best_score:
                    best_score = score
                    best = pat
            # Keep this high to prevent a learned correction from changing unrelated items.
            if not best or best_score < 0.94:
                continue
            changed = False
            if best.get("corrected_qty") and str(row.qty).strip() != str(best.get("corrected_qty")).strip():
                row.ai_suggest_qty = row.ai_suggest_qty or best.get("corrected_qty", "")
                if not row.qty or row.confidence < 65:
                    row.qty = best.get("corrected_qty", "")
                    changed = True
            if best.get("corrected_unit") and normalize_unit(row.unit) != normalize_unit(best.get("corrected_unit", "")):
                row.ai_suggest_unit = row.ai_suggest_unit or normalize_unit(best.get("corrected_unit", ""))
                if not row.unit or row.confidence < 65:
                    row.unit = normalize_unit(best.get("corrected_unit", ""))
                    changed = True
            if best.get("corrected_description"):
                learned_desc = best.get("corrected_description", "")
                sim = SequenceMatcher(None, clean_description(row.description).lower(), clean_description(learned_desc).lower()).ratio()
                if sim < 0.86 and row.confidence < 65:
                    row.ai_suggest_description = row.ai_suggest_description or learned_desc
                    row.description = learned_desc
                    changed = True
            if changed:
                row.confidence = min(100, max(row.confidence, 82))
                row.notes = append_note(row.notes, "Applied learned correction pattern")
                row.ai_message = append_note(row.ai_message, "Learned pattern adjusted this row from a prior confirmed correction.")
                row.ai_status = row.ai_status or "suggestion"
        return rows

    def apply_learning_boost(self, rows: List[MaterialLine]) -> List[MaterialLine]:
        learned = self.db.recent_corrections(limit=200)
        learned_terms = [clean_description(x.description).lower() for x in learned if x.description and x.review_status != "rejected"]
        for row in rows:
            rlow = row.description.lower()
            for term in learned_terms:
                if term and (term in rlow or rlow in term):
                    row.confidence = min(100, row.confidence + 8)
                    if not row.notes:
                        row.notes = "Matched learned correction pattern"
                    break
        return rows

    def overall_confidence(self, rows: List[MaterialLine], text: str) -> int:
        if not rows:
            return 0 if len(text.strip()) < 50 else 25
        avg = sum(r.confidence for r in rows) / max(1, len(rows))
        bonus = min(15, len(rows) * 2)
        qty_bonus = 10 if sum(1 for r in rows if r.qty) >= max(1, len(rows) // 2) else 0
        return int(max(0, min(100, avg + bonus + qty_bonus)))


def clean_description(desc: str) -> str:
    desc = (desc or "").strip()
    # OCR list lines often leave the closing parenthesis/bracket in front of the
    # description after parsing "-(20) 3/4 EMT" or "(3 boxes) TOPAZ".
    desc = re.sub(r"^[\)\]\}\>'’\s]+", "", desc).strip()
    # Remove pricing totals that frequently get glued to RFQ table descriptions.
    desc = strip_prices_from_text(desc) if "0.00" in desc or "$" in desc else desc
    # Vendor PDFs often append prices without a dollar sign, e.g. "112.907/C 112.91".
    desc = PRICE_TAIL_RE.sub("", desc)
    desc = re.sub(r"\s+\d+(?:\.\d{2,4})\s*/\s*(?:ea|each|c|m|ft|lf|pc|pcs|box|roll)\s*$", "", desc, flags=re.I)
    desc = re.sub(r"\s{2,}", " ", desc)
    desc = desc.strip(" |,;:-–—")
    # Remove leading placeholder cells from tables.
    desc = re.sub(r"^(?:[-–—]+\s+){1,4}", "", desc).strip()
    # Remove common quote request prefixes.
    desc = re.sub(r"^(please\s+)?(?:quote|price|need|looking for)\s+(?:on\s+)?", "", desc, flags=re.I)
    return desc.strip()


def clean_ocr_table_description(desc: str, qty: str = "", unit: str = "") -> str:
    """Clean OCR-table descriptions after Qty/UOM are already known.

    Image-only PDFs often OCR visible row numbers or quote/hash marks into the
    description, for example:
        '10 EA 5/8 x 8 hammer drill bit ...
        #50 EA 3/4 EMT compression connector ...
        "100 FT 12 MC cable ...
    Once the parser has the real qty/unit columns, these prefixes should not
    remain in Description.  Keep this narrow so normal Excel/table extraction is
    unaffected.
    """
    d = clean_description(desc or "")
    d = re.sub(r"^[\s'\"`´‘’“”#]+", "", d).strip()
    # Normalize stray column separators that some scans render between cells so they
    # don't survive inside the description (e.g. "1/2 lock washer ) LW-050").
    d = re.sub(r"\s*[\|\]]\s*", " ", d)
    # Collapse a " ) " used as a column gap (but keep a real "(...)" group intact).
    d = re.sub(r"\s+\)\s+", " ", d)
    q = re.escape(str(qty or "").strip())
    u = re.escape(str(unit or "").strip())
    sep = r"(?:\s*[\|\]\)]\s*|\s+)"
    if q and u:
        d = re.sub(rf"^(?:line\s*)?\d{{1,4}}[\.)]?\s+['\"#_]*{q}{sep}{u}\b\s*", "", d, flags=re.I).strip()
        d = re.sub(rf"^['\"#_]*{q}{sep}{u}\b\s*", "", d, flags=re.I).strip()
    elif q:
        d = re.sub(rf"^(?:line\s*)?\d{{1,4}}[\.)]?\s+['\"#]*{q}\b\s*", "", d, flags=re.I).strip()
        d = re.sub(rf"^['\"#]*{q}\b\s*", "", d, flags=re.I).strip()
    d = re.sub(r"^[\s'\"`´‘’“”#_.,;:-]+", "", d).strip()
    # OCR frequently glues a stray quote in front of a part number or word inside the
    # description (e.g. 'square washer "SQW-050', '“CC-075').  Remove a quote that sits
    # right before a LETTER — never one that sits after a digit, so inch marks like
    # 5/8" or 1/2" are preserved.
    d = re.sub(r"['\"`´‘’“”]+(?=[A-Za-z])", "", d)
    d = re.sub(r"(?<=\s)['\"`´‘’“”]+(?=\d)", "", d)
    # Remove a trailing table-border artifact OCR appends after the part number at the
    # very end of a row (a lone 1-2 char "7", "-", "|", quote).  Only strip a trailing
    # lone digit when it follows a part-number-like token, so real trailing values stay.
    d = re.sub(r"\s+(?:[-|,.]|['\"`´‘’“”])\s*$", "", d).strip()
    d = re.sub(r"([A-Za-z0-9]-\d{2,5})\s+\d{1,2}\s*$", r"\1", d).strip()
    return clean_description(d)


# Unit words used by the line-number-as-qty repair below.
_QF_UNIT_WORDS = (
    r"ea|each|pcs?|pieces?|ft|feet|foot|lf|rolls?|boxes|box|bags?|cases?|sets?|"
    r"coils?|buckets?|drums?|pallets?|gal|gallons?|kits?|lengths?|sticks?|pairs?|"
    r"reels?|spools?|packs?|pkgs?|dozens?|lbs?|m|meters?|yd|yards?"
)
_QF_UNIT_NORMSET = None


def repair_line_number_as_qty(row: "MaterialLine") -> "MaterialLine":
    """Fix the classic OCR table failure where the visible LINE number landed in the
    Qty field and the REAL quantity + unit got pushed to the front of the description.

    This runs on the finished row and is independent of which OCR separator was used
    (space, pipe, bracket, paren, etc.), so it repairs cases the line-level regexes
    miss because a vendor scan rendered the column gaps as ')' or other punctuation:

        Qty=5   Desc="500 ) EA ) 1/2 lock washer ) LW-050"
            ->  Qty=500  Unit=EA  Desc="1/2 lock washer"   (LW-050 -> part)

    It only fires when ALL of these hold, to stay safe:
      * the current Qty is a small integer that plausibly equals a line number, AND
      * the description begins with  <number> <sep> <UNIT> <sep> <rest...>, AND
      * a real descriptive remainder is left.
    """
    desc = (row.description or "").strip()
    if not desc:
        return row
    qty = str(row.qty or "").strip()
    # Current qty must look like a line number (small integer, <= 999) OR be empty.
    qty_is_linenum = bool(re.fullmatch(r"\d{1,3}", qty)) if qty else False
    if qty and not qty_is_linenum:
        # Real multi-hundred/comma quantities are not line numbers; leave them alone.
        if "," in qty or len(qty.replace(".", "")) >= 3:
            return row

    # v2.8.5 / MaINbox v3.8.05: do NOT rewrite a correct Qty just because the
    # description begins with a measurement like "10 ft length ...".  This was the
    # exact cause of line 26 in the 50-item image test becoming Qty=10 instead of
    # Qty=20: source was "26 20 EA 10 ft length ...".  Only run this repair when
    # the current qty actually equals the visible line number in the source row.
    src_line = normalize_space(getattr(row, "source_line", "") or "")
    src_line_no = ""
    m_src_no = re.match(r"^\s*\|?\s*(\d{1,4})\b", src_line)
    if m_src_no:
        src_line_no = m_src_no.group(1)
    if qty and src_line_no and qty != src_line_no:
        return row
    # Description must start with "<number> <sep> <unit> <sep> <rest>".  The separator
    # may be whitespace, a bracket/pipe, OR a space-surrounded comma/period used as a
    # column gap (e.g. "250 , FT . 10 THHN ...").  The quantity pattern consumes any
    # thousands comma / decimal first, so "2,000" and "2.5" are never split.
    _sep = r"(?:\s*[\|\]\)\(\[_]\s*|\s+[.,]\s+|\s+)"
    m = re.match(
        rf"^\(?\s*(?P<q>\d{{1,6}}(?:,\d{{3}})*(?:\.\d+)?)\s*\)?"
        rf"{_sep}"
        rf"(?P<u>{_QF_UNIT_WORDS})\b"
        rf"{_sep}"
        rf"(?P<rest>.+)$",
        desc, re.I,
    )
    if not m:
        return row
    real_qty = m.group("q").strip()
    real_unit = m.group("u").strip()
    rest = (m.group("rest") or "").strip()
    rest = re.sub(r"^[\s\|\]\)\(\[_:;.,-]+", "", rest).strip()
    if not rest or not re.search(r"[A-Za-z]{2,}", rest):
        return row
    # Guard: don't "repair" when the leading number is itself part of the description
    # measurement (e.g. a row that really is "2 1/2 inch coupling" with no unit word).
    # The required UOM word between the two tokens is what makes this safe.
    row.qty = normalize_qty(real_qty)
    if not row.unit:
        row.unit = normalize_unit(real_unit)
    # Strip a trailing catalog part token out of rest into the part field if empty.
    cleaned = clean_ocr_table_description(rest, row.qty, row.unit)
    if not row.part_number:
        for p in PART_NO_RE.findall(cleaned):
            pc = (p or "").strip().strip('"')
            pl = pc.lower().strip(".-_/ ")
            if SIZE_RE.fullmatch(pc) or pl in ELECTRICAL_KEYWORDS or pl in ELECTRICAL_UNITS:
                continue
            if re.fullmatch(r"\d+", pl):
                continue
            if re.search(r"\d", pc) or (pc.isupper() and len(pc) >= 4):
                row.part_number = pc
                break
    row.description = cleaned or rest
    return row


def order_materials_for_review(rows: List[MaterialLine]) -> List[MaterialLine]:
    """Return rows in a STABLE document/review order.

    The order is computed once (from page, then the scan sequence, then any known
    bbox-Y as a tie-break) and frozen into each row's ``review_order``.  Every later
    call simply sorts by that frozen value.  This is critical for multi-page scans:
    previously, clicking a row ran the OCR highlight locator which gave that one row a
    bbox, and because the sort key put bbox-rows in a different tier than bbox-less
    rows, the clicked row would jump to the top and scramble the whole list (the
    "row 26 shows the wrong item" bug).  Freezing the order removes that coupling.
    """
    # If every row already has a frozen order, just honor it — never reshuffle.
    if rows and all(getattr(r, "review_order", 0) for r in rows):
        ordered = sorted(rows, key=lambda r: r.review_order)
        for i, row in enumerate(ordered, start=1):
            row.scan_index = i
            if not row.review_status:
                row.review_status = "pending"
        return ordered

    def _bbox_xy(row: MaterialLine) -> Tuple[float, float]:
        try:
            x0, y0, _x1, _y1 = [float(v) for v in (row.source_bbox or "").split(",")]
            return x0, y0
        except Exception:
            return 10**9, 10**9

    def key(row: MaterialLine) -> Tuple[int, int, int, int]:
        # Primary axis is the PAGE.  Within a page, the row's position in the source
        # text (source_start offset) reflects its true top-to-bottom document order,
        # so use that FIRST — this keeps a row in its physical place even if a later
        # parse pass (e.g. the garbled-qty fallback) found it out of sequence and gave
        # it a larger scan_index.  scan_index breaks ties when offsets are unknown/equal.
        # A bbox-Y is only a final tie-break, so bbox presence never moves a row.
        page = row.source_page if row.source_page and row.source_page > 0 else 10**6
        idx = row.scan_index if row.scan_index and row.scan_index > 0 else 10**9
        offset = row.source_start if row.source_start is not None and row.source_start >= 0 else 10**12
        y_tie = 0
        if row.source_bbox:
            _x0, y0 = _bbox_xy(row)
            y_tie = int(round(y0 / 6.0)) if y0 < 10**8 else 0
        return (page, int(min(offset, 10**12)), idx, y_tie)

    ordered = sorted(rows, key=key)
    for i, row in enumerate(ordered, start=1):
        row.scan_index = i
        row.review_order = i  # freeze
        if not row.review_status:
            row.review_status = "pending"
    return ordered


class EvidenceMapper:
    """Best-effort mapper from extracted rows back to text offsets and PDF coordinates."""

    def __init__(self, file_path: Path, merged_text: str) -> None:
        self.file_path = Path(file_path)
        self.merged_text = merged_text or ""

    def annotate(self, rows: List[MaterialLine]) -> List[MaterialLine]:
        used_offsets: set[int] = set()
        for row in rows:
            start, end = self._find_text_span(row, used_offsets)
            row.source_start = start
            row.source_end = end
            # Preserve the structured-grid marker so the reviewer keeps trusting these
            # rows and the grid preview keeps highlighting them; only default to merged.
            if getattr(row, "source_view", "") != "table":
                row.source_view = "merged"
            if start >= 0:
                used_offsets.add(start)
                inferred_page = self._infer_page_from_offset(start)
                if inferred_page and not row.source_page:
                    row.source_page = inferred_page
            if self.file_path.suffix.lower() == ".pdf":
                page, bbox = self._find_pdf_bbox(row)
                if page:
                    row.source_page = page
                    row.source_bbox = bbox
                    row.source_bbox_kind = "pdf"
        return rows

    def _infer_page_from_offset(self, offset: int) -> int:
        if offset < 0:
            return 0
        prefix = self.merged_text[:offset]
        matches = list(re.finditer(r"---\s+(?:PAGE\s+)?(\d+)\s+(?:NATIVE TEXT|POSITIONAL LINES)|---\s+OCR\s+page_(\d+)", prefix, flags=re.I))
        if not matches:
            return 0
        m = matches[-1]
        val = m.group(1) or m.group(2)
        try:
            return int(val)
        except Exception:
            return 0

    def _find_text_span(self, row: MaterialLine, used_offsets: set[int]) -> Tuple[int, int]:
        # v2.5.4: Rows created from the OCR ordered-list parser should link back to
        # the OCR TEXT section, not to earlier TABLE / POSITIONAL CANDIDATES that may
        # contain similar fragments.  This keeps review order and text highlights clean.
        search_start = 0
        if "ocr ordered list" in (row.notes or "").lower():
            m = re.search(r"^---\s*OCR\s+TEXT\s*---\s*$", self.merged_text or "", flags=re.I | re.M)
            if m:
                search_start = m.start()

        needles = [row.source_line, row.description, row.part_number]
        for needle in needles:
            needle = (needle or "").strip()
            if not needle:
                continue
            pos = self.merged_text.find(needle, search_start)
            while pos in used_offsets and pos >= 0:
                pos = self.merged_text.find(needle, pos + 1)
            if pos >= 0:
                return pos, pos + len(needle)
        # Fallback: try a short phrase from the description.
        words = re.findall(r"[A-Za-z0-9#/\'\".-]+", row.description or "")
        if len(words) >= 3:
            phrase = " ".join(words[:5])
            pos = self.merged_text.lower().find(phrase.lower(), search_start)
            if pos >= 0:
                return pos, pos + len(phrase)
        return -1, -1

    def _find_pdf_bbox(self, row: MaterialLine) -> Tuple[int, str]:
        if fitz is None:
            return 0, ""
        candidates = self._pdf_candidates(row)
        try:
            doc = fitz.open(str(self.file_path))
            max_pages = min(len(doc), 200)
            preferred_pages: List[int] = []
            if row.source_page:
                preferred_pages.append(max(0, min(row.source_page - 1, max_pages - 1)))
            preferred_pages.extend([i for i in range(max_pages) if i not in preferred_pages])
            for page_i in preferred_pages:
                page = doc[page_i]
                for candidate in candidates:
                    rect = self._search_pdf_exact(page, candidate)
                    if rect is None:
                        rect = self._search_pdf_words(page, candidate)
                    if rect is not None:
                        doc.close()
                        return page_i + 1, f"{rect.x0:.2f},{rect.y0:.2f},{rect.x1:.2f},{rect.y1:.2f}"
            doc.close()
        except Exception:
            pass
        return 0, ""

    def _pdf_candidates(self, row: MaterialLine) -> List[str]:
        raw_candidates = [row.source_line, row.description, row.part_number]
        words = re.findall(r"[A-Za-z0-9#/\'\".-]+", row.description or "")
        if len(words) >= 3:
            raw_candidates.append(" ".join(words[:8]))
            raw_candidates.append(" ".join(words[:5]))
        if row.qty and row.description:
            raw_candidates.append(f"{row.qty} {row.description}")
        out: List[str] = []
        seen = set()
        for value in raw_candidates:
            value = normalize_space((value or "").strip())
            if len(value) < 3:
                continue
            value = value[:180]
            key = value.lower()
            if key not in seen:
                seen.add(key)
                out.append(value)
        return out

    def _search_pdf_exact(self, page: Any, candidate: str) -> Optional[Any]:
        try:
            rects = page.search_for(candidate)
        except Exception:
            rects = []
        if not rects:
            return None
        rect = rects[0]
        for r in rects[1:6]:
            rect |= r
        return rect

    def _search_pdf_words(self, page: Any, candidate: str) -> Optional[Any]:
        tokens = self._tokens(candidate)
        if len(tokens) < 2:
            return None
        # Avoid matching a qty-only prefix; description/part number tokens are more distinctive.
        tokens = [t for t in tokens if not re.fullmatch(r"\d+(?:\.\d+)?", t)]
        if len(tokens) < 2:
            return None
        tokens = tokens[:10]
        try:
            words = page.get_text("words", sort=True) or []
        except Exception:
            return None
        page_words = []
        for w in words:
            try:
                text = self._norm_token(str(w[4]))
                if text:
                    page_words.append((text, w))
            except Exception:
                continue
        if not page_words:
            return None
        best_rect = None
        best_score = 0
        for i, (tok, _) in enumerate(page_words):
            if tok != tokens[0]:
                continue
            match_words = [page_words[i][1]]
            token_idx = 1
            j = i + 1
            gaps = 0
            while j < len(page_words) and token_idx < len(tokens) and gaps <= 10:
                if page_words[j][0] == tokens[token_idx]:
                    match_words.append(page_words[j][1])
                    token_idx += 1
                    gaps = 0
                else:
                    gaps += 1
                j += 1
            score = token_idx
            if score > best_score and score >= min(3, len(tokens)):
                try:
                    rect = fitz.Rect(match_words[0][0], match_words[0][1], match_words[0][2], match_words[0][3])
                    for mw in match_words[1:]:
                        rect |= fitz.Rect(mw[0], mw[1], mw[2], mw[3])
                    best_rect = rect
                    best_score = score
                except Exception:
                    pass
            if best_score >= min(6, len(tokens)):
                break
        return best_rect

    def _tokens(self, text: str) -> List[str]:
        return [self._norm_token(t) for t in re.findall(r"[A-Za-z0-9#/\'\".-]+", text or "") if self._norm_token(t)]

    def _norm_token(self, text: str) -> str:
        return re.sub(r"[^a-z0-9]+", "", (text or "").lower())


class AIReviewIssue:
    pass


# === v2.8.5: catalog grounding =============================================
# Validate parsed rows against an optional product catalog
# (american_power_catalog.db). Non-destructive: it only fills the ai_suggest_*
# fields (when empty) and appends a short note, mirroring AIReviewer. If the DB
# is absent/unreadable the grounder is a silent no-op, so behaviour is never
# worse than before. Connection is read-only and shared (check_same_thread off).
import os as _os
import sys as _sys
import difflib as _difflib


def _norm_part(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", (s or "").lower())


def _resolve_catalog_db_path(settings) -> str:
    p = (getattr(settings, "catalog_db_path", "") or "").strip().strip('"')
    if p and _os.path.exists(p):
        return p
    env = (_os.environ.get("MAINBOX_CATALOG_DB", "") or "").strip().strip('"')
    if env and _os.path.exists(env):
        return env
    bases = []
    try:
        bases.append(_os.path.dirname(_os.path.abspath(_sys.argv[0] or ".")))
    except Exception:
        pass
    bases.append(_os.getcwd())
    for b in bases:
        cand = _os.path.join(b, "american_power_catalog.db")
        if _os.path.exists(cand):
            return cand
    return ""


class CatalogGrounder:
    def __init__(self, db_path: str, log=None) -> None:
        self.db_path = db_path
        self.log = log or (lambda m: None)
        self.ok = False
        self.by_part: Dict[str, Any] = {}
        self._keys: List[str] = []
        self.con = None
        self._load()

    def _load(self) -> None:
        if not self.db_path:
            return
        try:
            self.con = sqlite3.connect(self.db_path, check_same_thread=False)
            self.con.row_factory = sqlite3.Row
            cur = self.con.execute(
                "SELECT part_number, part_number_alt, description, "
                "COALESCE(brand, manufacturer_name) AS mfr "
                "FROM products WHERE part_number IS NOT NULL"
            )
            for r in cur:
                for key in (r["part_number"], r["part_number_alt"]):
                    if key:
                        k = _norm_part(str(key))
                        if k and k not in self.by_part:
                            self.by_part[k] = r
            self._keys = list(self.by_part.keys())
            self.ok = len(self.by_part) > 0
            if self.ok:
                self.log(f"Catalog grounding: loaded {len(self.by_part)} part keys "
                         f"from {_os.path.basename(self.db_path)}")
        except Exception as e:
            self.ok = False
            self.log(f"Catalog grounding disabled (load failed): {e}")

    def _lookup_part(self, part: str):
        k = _norm_part(part)
        if not k:
            return None
        rec = self.by_part.get(k)
        if rec is not None:
            return ("exact", rec)
        near = _difflib.get_close_matches(k, self._keys, n=1, cutoff=0.88)
        if near:
            return ("snapped", self.by_part[near[0]])
        return None

    def _lookup_desc(self, desc: str):
        if not desc or self.con is None:
            return None
        d = re.sub(r"[^a-z0-9 ]+", " ", desc.lower()).strip()
        toks = sorted({t for t in d.split() if len(t) >= 3}, key=len, reverse=True)[:3]
        if not toks:
            return None
        cand = {}
        for t in toks:
            try:
                rows = self.con.execute(
                    "SELECT part_number, description, COALESCE(brand, manufacturer_name) AS mfr "
                    "FROM products WHERE description_normalized LIKE ? LIMIT 60",
                    (f"%{t}%",)
                ).fetchall()
            except Exception:
                rows = []
            for r in rows:
                cand[(r["part_number"], r["description"])] = r
        best, score = None, 0.0
        for r in cand.values():
            s = SequenceMatcher(None, d, (r["description"] or "").lower()).ratio()
            if s > score:
                best, score = r, s
        if best is not None and score >= 0.62:
            return ("desc", best)
        return None

    def ground_row(self, row) -> None:
        if not self.ok:
            return
        part = (getattr(row, "part_number", "") or "").strip()
        hit = self._lookup_part(part)
        how = rec = None
        if hit:
            how, rec = hit
        elif not part:
            d = self._lookup_desc(getattr(row, "description", ""))
            if d:
                how, rec = d
        if rec is None:
            return
        cat_part = (rec["part_number"] or "").strip()
        cat_desc = (rec["description"] or "").strip()
        cat_mfr = (rec["mfr"] or "").strip()
        rowdesc = (getattr(row, "description", "") or "").strip()

        def _empty(attr):
            return not (getattr(row, attr, "") or "").strip()

        note = None
        if how == "exact":
            note = f"Catalog OK {cat_part}" + (f" - {cat_mfr}" if cat_mfr else "")
            if cat_mfr and _empty("manufacturer") and _empty("ai_suggest_manufacturer"):
                row.ai_suggest_manufacturer = cat_mfr
            if cat_desc and len(rowdesc) < 4 and _empty("ai_suggest_description"):
                row.ai_suggest_description = cat_desc
        elif how == "snapped":
            note = f"Catalog: {part or '?'} -> {cat_part}" + (f" - {cat_mfr}" if cat_mfr else "")
            if _empty("ai_suggest_part_number"):
                row.ai_suggest_part_number = cat_part
            if cat_mfr and _empty("manufacturer") and _empty("ai_suggest_manufacturer"):
                row.ai_suggest_manufacturer = cat_mfr
            if getattr(row, "ai_status", "") in ("", "ok"):
                row.ai_status = "suggestion"
        else:  # desc match (row had no part number)
            note = f"Catalog match -> {cat_part}" + (f" - {cat_mfr}" if cat_mfr else "")
            if cat_part and _empty("ai_suggest_part_number"):
                row.ai_suggest_part_number = cat_part
                if getattr(row, "ai_status", "") in ("", "ok"):
                    row.ai_status = "suggestion"
            if cat_mfr and _empty("manufacturer") and _empty("ai_suggest_manufacturer"):
                row.ai_suggest_manufacturer = cat_mfr

        if note:
            existing = getattr(row, "notes", "") or ""
            if "Catalog" not in existing:
                row.notes = (existing + " | " + note).strip(" |")
            aim = getattr(row, "ai_message", "") or ""
            if "Catalog" not in aim:
                row.ai_message = (aim + " | " + note).strip(" |") if aim else note

    def ground_rows(self, rows) -> None:
        if not self.ok:
            return
        for row in rows:
            try:
                self.ground_row(row)
            except Exception:
                pass


_CATALOG_GROUNDER_CACHE: Dict[str, "CatalogGrounder"] = {}


def _catalog_grounder_for(settings, log=None) -> "CatalogGrounder":
    path = _resolve_catalog_db_path(settings)
    g = _CATALOG_GROUNDER_CACHE.get(path)
    if g is None:
        g = CatalogGrounder(path, log=log)
        _CATALOG_GROUNDER_CACHE[path] = g
    return g


class AIReviewer:
    """Second-pass row inspector.

    The rule reviewer is fast, local, and deterministic. Optional Ollama support is
    deliberately conservative and only adds a short human-readable note unless the
    user accepts a suggested correction.
    """

    def __init__(self, settings: Optional[ScanSettings] = None, log=None) -> None:
        self.settings = settings or ScanSettings()
        self.log = log or (lambda msg: None)

    def review(self, rows: List[MaterialLine], merged_text: str = "") -> List[MaterialLine]:
        seen: Dict[Tuple[str, str, str, str], int] = {}
        for idx, row in enumerate(rows, start=1):
            self._clear_ai(row)
            issues: List[Tuple[str, str]] = []
            suggestions: Dict[str, str] = {}
            source = normalize_space(row.source_line or row.description or "")
            desc = normalize_space(row.description or "")

            # Structured-grid rows (spreadsheet / Word table) are high-trust: the row
            # came from an explicit data column, so the electrical-text heuristics below
            # (which assume messy OCR/PDF lines) must NOT flag them as "document noise",
            # invent a UOM from a description word, or demand electrical keywords. We
            # only check for genuinely actionable problems and otherwise mark OK.
            is_structured = (getattr(row, "source_view", "") == "table")
            if is_structured:
                if not desc or len(desc) < 2:
                    issues.append(("error", "Description is blank or too short."))
                else:
                    # Price/total accidentally glued into the description still worth flagging.
                    cleaned_desc = clean_description(desc)
                    no_money = MONEY_RE.sub("", cleaned_desc)
                    no_money = re.sub(r"\b(?:unit\s*price|extended?|extension|subtotal|total|tax)\b.*$", "", no_money, flags=re.I).strip(" -–—,;")
                    if no_money and normalize_space(no_money) != normalize_space(cleaned_desc):
                        issues.append(("suggestion", "Description appears to contain price/total text."))
                        suggestions["description"] = normalize_space(no_money)
                if str(row.qty).strip() and re.fullmatch(r"\d{7,}", str(row.qty).strip()):
                    issues.append(("warning", "Quantity looks like an ID, not a count."))
                key = warn_key(row)
                if key in seen:
                    issues.append(("warning", f"Possible duplicate of row #{seen[key]}."))
                else:
                    seen[key] = idx
                self._apply_issues(row, issues, suggestions)
                continue

            # Strong false positive filters.
            if is_definite_admin_noise(source) or is_definite_admin_noise(desc) or is_probably_contact_or_address(source) or is_probably_contact_or_address(desc):
                issues.append(("error", "Looks like contact/address/web information, not a material line."))
                suggestions["description"] = ""
            elif material_signal_score(desc, row.qty, row.unit, row.part_number) < 38:
                issues.append(("error", "Not enough material/quantity/UOM signal; likely document noise."))
                suggestions["description"] = ""

            # Classic table failure: line number used as quantity.
            mt = ROWNUM_QTY_UNIT_RE.match(source)
            if mt:
                rownum = mt.group("rownum")
                real_qty = mt.group("qty")
                real_unit = normalize_unit(mt.group("unit"))
                rest = clean_description(mt.group("rest"))
                # If parser chose the row number, or if the chosen qty equals the row's display index, fix it.
                if str(row.qty).strip() in {rownum, str(row.scan_index)} or str(row.qty).strip() != real_qty:
                    issues.append(("error", f"Likely used row/line number {rownum} as quantity; detected Qty {real_qty} {real_unit}."))
                    suggestions.update({"qty": real_qty, "unit": real_unit, "description": rest})

            # Header/table pattern: first token often line number, second token qty, third token UOM.
            tokens = source.split()
            if len(tokens) >= 4:
                if re.fullmatch(r"\d{1,4}", tokens[0] or "") and re.fullmatch(r"\d+(?:\.\d+)?", tokens[1] or "") and normalize_unit(tokens[2]):
                    if str(row.qty).strip() == tokens[0] or not row.unit:
                        issues.append(("error", "Detected row-number + Qty + UOM pattern."))
                        suggestions.update({"qty": tokens[1], "unit": normalize_unit(tokens[2]), "description": clean_description(" ".join(tokens[3:]))})

            if not str(row.qty).strip():
                issues.append(("warning", "Missing quantity."))
            elif re.fullmatch(r"\d{7,}", str(row.qty).strip()):
                issues.append(("error", "Quantity looks like an ID/phone/part number, not a count."))
            elif str(row.qty).strip() == str(row.scan_index) and source and re.search(r"\b(?:EA|FT|PC|PCS|BOX|ROLL|STICK|SET|PKG)\b", source, re.I):
                # If the source line explicitly shows LINE + QTY + UOM and the Qty
                # column really equals the line number (example: line 30, qty 30),
                # do not warn.  The earlier warning was useful for bad OCR, but it
                # created false review items on perfectly correct rows.
                explicit_qty_ok = False
                try:
                    mt2 = ROWNUM_QTY_UNIT_RE.match(source)
                    if mt2 and str(row.qty).strip() == str(mt2.group("qty") or "").strip() and normalize_unit(mt2.group("unit")) == normalize_unit(row.unit):
                        explicit_qty_ok = True
                    else:
                        # Handles pipe/table OCR forms like:
                        #   | 30 | 30 | EA . 1-1/4 in one-hole EMT strap |
                        mt_pipe = re.match(
                            r"^\s*\|?\s*(?P<rownum>\d{1,4})\s*\|?\s*(?P<qty>\d{1,6}(?:,\d{3})*(?:\.\d+)?)\s*\|?\s*(?P<unit>[A-Za-z']+)\b",
                            source, re.I,
                        )
                        if mt_pipe and mt_pipe.group("rownum") == str(row.scan_index) and mt_pipe.group("qty") == str(row.qty).strip() and normalize_unit(mt_pipe.group("unit")) == normalize_unit(row.unit):
                            explicit_qty_ok = True
                        else:
                            toks2 = [t.strip("|().") for t in source.split() if t.strip("|().")]
                            if len(toks2) >= 3 and toks2[0] == str(row.scan_index) and toks2[1] == str(row.qty).strip() and normalize_unit(toks2[2]) == normalize_unit(row.unit):
                                explicit_qty_ok = True
                except Exception:
                    explicit_qty_ok = False
                if not explicit_qty_ok:
                    issues.append(("warning", "Quantity equals row number; verify it is not the line count."))

            if not str(row.unit).strip():
                unit_guess = self._guess_unit_from_source(source)
                if unit_guess:
                    issues.append(("suggestion", f"Unit of measurement appears to be {unit_guess}."))
                    suggestions["unit"] = unit_guess
                else:
                    issues.append(("warning", "Missing unit/UOM."))

            if not desc or len(desc) < 3:
                issues.append(("error", "Description is blank or too short."))
            else:
                cleaned_desc = clean_description(desc)
                # Remove prices/totals that got glued to descriptions.
                no_money = MONEY_RE.sub("", cleaned_desc)
                no_money = re.sub(r"\b(?:unit\s*price|extended?|extension|subtotal|total|tax)\b.*$", "", no_money, flags=re.I).strip(" -–—,;")
                if no_money and normalize_space(no_money) != normalize_space(cleaned_desc):
                    issues.append(("suggestion", "Description appears to contain price/total text."))
                    suggestions["description"] = normalize_space(no_money)
                if not looks_materialish(desc) and not row.part_number:
                    issues.append(("warning", "Description does not strongly look like electrical material."))

            key = warn_key(row)
            if key in seen:
                issues.append(("warning", f"Possible duplicate of row #{seen[key]}."))
            else:
                seen[key] = idx

            self._apply_issues(row, issues, suggestions)

        if self.settings.enable_ollama_ai_review:
            self._ollama_review(rows)
        if getattr(self.settings, "enable_catalog_grounding", True):
            try:
                _catalog_grounder_for(self.settings, self.log).ground_rows(rows)
            except Exception as e:
                self.log(f"Catalog grounding skipped/failed: {e}")
        return rows

    def _clear_ai(self, row: MaterialLine) -> None:
        row.ai_status = ""
        row.ai_message = ""
        row.ai_suggest_qty = ""
        row.ai_suggest_unit = ""
        row.ai_suggest_description = ""
        row.ai_suggest_part_number = ""
        row.ai_suggest_manufacturer = ""

    def _apply_issues(self, row: MaterialLine, issues: List[Tuple[str, str]], suggestions: Dict[str, str]) -> None:
        severity_order = {"error": 3, "suggestion": 2, "warning": 1, "ok": 0}
        if not issues:
            row.ai_status = "ok"
            row.ai_message = "AI/rule review agrees"
        else:
            status = max((sev for sev, _ in issues), key=lambda x: severity_order.get(x, 0))
            row.ai_status = status
            row.ai_message = "; ".join(msg for _, msg in issues[:5])
        row.ai_suggest_qty = suggestions.get("qty", "")
        row.ai_suggest_unit = suggestions.get("unit", "")
        row.ai_suggest_description = suggestions.get("description", "")
        row.ai_suggest_part_number = suggestions.get("part_number", "")
        row.ai_suggest_manufacturer = suggestions.get("manufacturer", "")
        if row.ai_status in {"error", "warning", "suggestion"}:
            ai_note = f"AI Review: {row.ai_message}"
            if ai_note not in (row.notes or ""):
                row.notes = (row.notes + " | " + ai_note).strip(" |")

    def _guess_unit_from_source(self, source: str) -> str:
        for tok in re.findall(r"\b[A-Za-z]{1,10}\b", source or ""):
            u = normalize_unit(tok)
            if u:
                return u
        return ""

    def _ollama_review(self, rows: List[MaterialLine]) -> None:
        # Keep this optional and limited so the UI does not become dependent on Ollama.
        targets = []
        for r in rows:
            suspicious = r.ai_status in {"error", "warning", "suggestion"} or int(r.confidence or 0) < 65
            if self.settings.ai_review_mode == "Slow Careful" or (self.settings.ai_review_mode == "Normal" and suspicious):
                targets.append(r)
        if not targets:
            return
        for row in targets[:50]:
            try:
                prompt = (
                    "You are reviewing one extracted electrical RFQ material row. "
                    "Return one short sentence only. Do not output JSON. "
                    f"Row: qty={row.qty!r}, unit={row.unit!r}, desc={row.description!r}, part={row.part_number!r}, source={row.source_line!r}. "
                    "Say whether it looks correct or what likely error exists."
                )
                body = json.dumps({"model": self.settings.ai_review_ollama_model, "prompt": prompt, "stream": False}).encode("utf-8")
                req = urllib.request.Request(self.settings.ai_review_ollama_url, data=body, headers={"Content-Type": "application/json"})
                with urllib.request.urlopen(req, timeout=12) as resp:
                    data = json.loads(resp.read().decode("utf-8", errors="replace"))
                note = normalize_space(str(data.get("response", "")))[:220]
                if note:
                    row.ai_message = (row.ai_message + " | Ollama: " + note).strip(" |")
            except Exception as e:
                self.log(f"Ollama AI review skipped/failed: {e}")
                return



# ═══════════════════════════════════════════════════════════════════════════════
# v2.0  Anthropic real-time per-item reviewer
# ═══════════════════════════════════════════════════════════════════════════════
_ANTHROPIC_API_URL = "https://api.anthropic.com/v1/messages"
_ANTHROPIC_MODEL   = "claude-haiku-4-5-20251001"

_ANTHROPIC_SYSTEM = """You are a quality-control assistant for electrical-material RFQ scanning.
Given ONE extracted material row, answer strictly with a single JSON object — no markdown, no extra text.
Decide:
  status: "ok" | "warning" | "error"
    ok      = row looks like a real, plausible electrical material line
    warning = something is suspicious but might still be correct
    error   = almost certainly wrong, contact/address noise, or garbled OCR

Also fill these fields only when you have a concrete suggestion (leave "" otherwise):
  suggest_qty, suggest_unit, suggest_description, suggest_part_number

Example:
{"status":"ok","message":"Looks correct.","suggest_qty":"","suggest_unit":"","suggest_description":"","suggest_part_number":""}
"""

class AnthropicReviewer:
    """Calls claude-haiku once per material row immediately after parsing.

    Results are cached so re-running AI review is free.
    Falls back silently if the API key is missing or the call fails.
    """

    def __init__(self, api_key: str, log=None) -> None:
        self.api_key = (api_key or "").strip()
        self.log = log or (lambda msg: None)
        self._cache: Dict[str, Dict] = {}

    def review_row(self, row: "MaterialLine") -> Dict:
        """Return an AI verdict dict for this row.  Never raises."""
        if not self.api_key:
            return {}
        cache_key = f"{row.qty}|{row.unit}|{row.description}|{row.part_number}|{row.source_line}"
        if cache_key in self._cache:
            return self._cache[cache_key]
        user_msg = (
            f"qty={row.qty!r}  unit={row.unit!r}  description={row.description!r}  "
            f"part_number={row.part_number!r}  manufacturer={row.manufacturer!r}  "
            f"source_line={row.source_line!r}"
        )
        try:
            body = json.dumps({
                "model": _ANTHROPIC_MODEL,
                "max_tokens": 256,
                "system": _ANTHROPIC_SYSTEM,
                "messages": [{"role": "user", "content": user_msg}],
            }).encode("utf-8")
            req = urllib.request.Request(
                _ANTHROPIC_API_URL, data=body,
                headers={
                    "Content-Type": "application/json",
                    "x-api-key": self.api_key,
                    "anthropic-version": "2023-06-01",
                },
            )
            with urllib.request.urlopen(req, timeout=18) as resp:
                data = json.loads(resp.read().decode("utf-8", errors="replace"))
            text = "".join(b.get("text", "") for b in data.get("content", []) if b.get("type") == "text")
            # strip any accidental markdown fences
            text = re.sub(r"^```[a-z]*\n?|\n?```$", "", text.strip())
            result = json.loads(text)
            self._cache[cache_key] = result
            return result
        except Exception as exc:
            self.log(f"Anthropic reviewer: {exc}")
            return {}

    def apply(self, row: "MaterialLine") -> None:
        """Apply AI verdict to row in-place.  Sets ai_status, ai_message, suggestions,
        and highlight_color (green / yellow / red)."""
        verdict = self.review_row(row)
        if not verdict:
            return
        status = verdict.get("status", "")
        if status in {"ok", "warning", "error"}:
            # Only upgrade severity, never downgrade a rule-review verdict
            severity = {"ok": 0, "warning": 1, "suggestion": 1, "error": 2}
            if severity.get(status, 0) >= severity.get(row.ai_status, 0):
                row.ai_status = status
        msg = verdict.get("message", "")
        if msg:
            row.ai_message = (row.ai_message + " | Claude: " + msg).strip(" |") if row.ai_message else "Claude: " + msg
        for field_name in ("suggest_qty", "suggest_unit", "suggest_description", "suggest_part_number"):
            val = verdict.get(field_name, "")
            attr = "ai_" + field_name
            if val and not getattr(row, attr, ""):
                setattr(row, attr, val)
        # Assign highlight colour
        st = row.ai_status.lower()
        row.highlight_color = "green" if st == "ok" else ("red" if st == "error" else "yellow")


@dataclass
class SourceMaterialCandidate:
    index: int
    source_line: str
    source_start: int = -1
    source_end: int = -1
    source_page: int = 0
    parsed: Optional[MaterialLine] = None
    matched_row_index: int = -1
    issue: str = ""


class ExtractionWatchdog:
    """Source-vs-row audit layer.

    The normal extractor builds rows.  This watchdog builds a separate ledger of
    material-looking OCR/source lines, then checks whether each source line
    became exactly one material row in the same order.  It is intentionally
    conservative: it only recovers rows when the existing parser can parse the
    original source line, and it flags suspicious mismatches instead of silently
    rewriting user-visible data.
    """

    def __init__(self, parser: MaterialParser, log=None) -> None:
        self.parser = parser
        self.log = log or (lambda msg: None)

    def audit(self, rows: List[MaterialLine], merged_text: str) -> Tuple[List[MaterialLine], List[str]]:
        ledger = self.build_source_ledger(merged_text or "")
        if not ledger:
            return rows, []

        rows = list(rows or [])
        warnings: List[str] = []
        used_rows: set[int] = set()
        recovered_count = 0
        mismatch_count = 0
        duplicate_count = 0

        # First pass: match each ledger source line to the best existing row.
        for cand in ledger:
            best_idx, best_score = self.best_row_match(cand, rows, used_rows)
            if best_idx >= 0 and best_score >= 0.58:
                cand.matched_row_index = best_idx
                used_rows.add(best_idx)
                row = rows[best_idx]
                row.source_ledger_index = cand.index
                if not row.source_line:
                    row.source_line = cand.source_line
                if row.source_start is None or row.source_start < 0:
                    row.source_start = cand.source_start
                    row.source_end = cand.source_end
                if cand.source_page and not row.source_page:
                    row.source_page = cand.source_page
                self.audit_matched_fields(cand, row)
                if row.watchdog_status in {"warning", "suggestion"}:
                    mismatch_count += 1
                continue

            # No match: recover a pending row directly from the source line.
            if cand.parsed and should_keep_material_candidate(cand.parsed):
                # v2.8.3: before adding a recovered row, check the existing list one
                # more time using a stricter duplicate/accountability comparison.
                # This prevents the watchdog from creating a second row when the
                # source line was already extracted but did not clear the fuzzy
                # match threshold because OCR punctuation differed.
                dup_idx = self.find_existing_material_match(cand.parsed, rows)
                if dup_idx >= 0:
                    cand.matched_row_index = dup_idx
                    used_rows.add(dup_idx)
                    row = rows[dup_idx]
                    row.source_ledger_index = row.source_ledger_index or cand.index
                    if not row.source_line:
                        row.source_line = cand.source_line
                    if row.source_start is None or row.source_start < 0:
                        row.source_start = cand.source_start
                        row.source_end = cand.source_end
                    row.watchdog_status = row.watchdog_status or "covered"
                    row.watchdog_message = row.watchdog_message or "Source line accounted for by duplicate-tightener match."
                    continue

                recovered = cand.parsed
                recovered.source_ledger_index = cand.index
                recovered.source_start = cand.source_start
                recovered.source_end = cand.source_end
                recovered.source_page = cand.source_page
                recovered.review_status = "pending"
                recovered.ai_status = "warning"
                recovered.ai_message = (
                    recovered.ai_message
                    or "AI Watchdog recovered this source line because it was missing from the material list."
                )
                recovered.watchdog_status = "recovered"
                recovered.watchdog_message = (
                    recovered.watchdog_message
                    or "Recovered from OCR/source ledger; please confirm."
                )
                recovered.notes = append_note(recovered.notes, "AI Watchdog: recovered missing OCR/source line")
                recovered.highlight_color = "yellow"
                recovered.snapshot_original()
                rows.append(recovered)
                recovered_count += 1

        # Second pass: flag likely duplicate rows that map to the same source meaning.
        seen: Dict[Tuple[str, str, str, str], int] = {}
        for i, row in enumerate(rows):
            if (row.review_status or "").lower() == "rejected":
                continue
            key = warn_key(row)
            if key in seen:
                duplicate_count += 1
                row.ai_status = row.ai_status or "warning"
                row.ai_message = append_note(row.ai_message, f"AI Watchdog: possible duplicate of row #{seen[key] + 1}.")
                row.watchdog_status = "warning"
                row.watchdog_message = append_note(row.watchdog_message, "Possible duplicate created during extraction.")
                row.notes = append_note(row.notes, "AI Watchdog: possible duplicate row")
                row.highlight_color = "yellow"
            else:
                # If the exact key is new but a very close same-qty/same-desc row
                # already exists, cluster it as a duplicate as well.
                near = -1
                for j, prev in enumerate(rows[:i]):
                    if (prev.review_status or "").lower() == "rejected":
                        continue
                    if loose_same_material(prev, row, min_ratio=0.90):
                        near = j
                        break
                if near >= 0:
                    duplicate_count += 1
                    row.ai_status = row.ai_status or "warning"
                    row.ai_message = append_note(row.ai_message, f"AI Watchdog: possible duplicate of row #{near + 1}.")
                    row.watchdog_status = "warning"
                    row.watchdog_message = append_note(row.watchdog_message, "Possible near-duplicate created during extraction.")
                    row.notes = append_note(row.notes, "AI Watchdog: possible duplicate row")
                    row.highlight_color = "yellow"
                else:
                    seen[key] = i

        rows = order_materials_for_review(rows)
        # Preserve ledger order where possible when source_start is known.
        for i, row in enumerate(rows, start=1):
            row.scan_index = i

        if recovered_count:
            warnings.append(f"AI Watchdog recovered {recovered_count} material line(s) from the OCR/source ledger.")
        if mismatch_count:
            warnings.append(f"AI Watchdog flagged {mismatch_count} row(s) whose qty/unit/description may not match the source line.")
        if duplicate_count:
            warnings.append(f"AI Watchdog flagged {duplicate_count} possible duplicate row(s).")
        if ledger:
            matched = sum(1 for c in ledger if c.matched_row_index >= 0)
            warnings.append(f"AI Watchdog source coverage: {matched + recovered_count}/{len(ledger)} material-looking source line(s) accounted for.")
        return rows, warnings

    def build_source_ledger(self, merged_text: str) -> List[SourceMaterialCandidate]:
        text = merged_text or ""
        # Prefer the clean OCR section when present; otherwise audit the whole merged text.
        audit_text = self.parser._ocr_text_section(text)
        base_offset = 0 if audit_text == text else text.find(audit_text)
        if base_offset < 0:
            base_offset = 0

        ledger: List[SourceMaterialCandidate] = []
        offset = 0
        page = 0
        for raw in split_lines_preserve(audit_text):
            line_start = base_offset + offset
            line_end = line_start + len(raw)
            offset += len(raw) + 1

            m_page = re.search(r"---\s+OCR\s+page_(\d+)", raw, flags=re.I)
            if m_page:
                try:
                    page = int(m_page.group(1))
                except Exception:
                    pass
                continue
            if re.match(r"^---\s*(?:ocr text|native text|table|page)\b", raw.strip(), re.I):
                continue

            parsed = self.parser.parse_ocr_ordered_list_line(raw)
            if not parsed:
                # v2.8.1: separate audit pass for clean OCR bullet quantity lines.
                # These are the source lines the user sees as "-(quantity) description".
                # Even if they lack electrical keywords/UOM (safety glasses, ear plugs,
                # work gloves, etc.), they should be flagged as possible missing rows
                # instead of being silently ignored.
                parsed = parse_dash_quantity_source_line(raw)
            if not parsed:
                # Conservative fallback: only consider strong material-looking rows that contain
                # qty/UOM-like content so ordinary email text is not added to the ledger.
                if material_signal_score(raw) < 70 or not (has_explicit_uom(raw) or re.search(r"\(\s*\d+", raw)):
                    continue
                parsed = self.parser.parse_line(raw)
            if not parsed:
                continue
            parsed.source_line = parsed.source_line or raw.strip()
            parsed.source_start = line_start
            parsed.source_end = line_end
            parsed.source_page = page
            parsed.source_ledger_index = len(ledger) + 1
            ledger.append(SourceMaterialCandidate(
                index=len(ledger) + 1,
                source_line=raw.strip(),
                source_start=line_start,
                source_end=line_end,
                source_page=page,
                parsed=parsed,
            ))
        return ledger

    def best_row_match(self, cand: SourceMaterialCandidate, rows: List[MaterialLine], used_rows: set[int]) -> Tuple[int, float]:
        best_idx = -1
        best_score = 0.0
        cand_text = self._norm(cand.source_line)
        cand_desc = self._norm(cand.parsed.description if cand.parsed else cand.source_line)
        for i, row in enumerate(rows):
            if i in used_rows:
                continue
            row_text = self._norm(row.source_line or row.description)
            row_desc = self._norm(row.description or row.source_line)
            if not row_text and not row_desc:
                continue
            score = max(
                SequenceMatcher(None, cand_text, row_text).ratio(),
                SequenceMatcher(None, cand_desc, row_desc).ratio(),
            )
            # Strong boost when qty/unit agree and the descriptions overlap.
            if cand.parsed:
                if str(cand.parsed.qty).strip() and str(cand.parsed.qty).strip() == str(row.qty).strip():
                    score += 0.12
                if normalize_unit(cand.parsed.unit) and normalize_unit(cand.parsed.unit) == normalize_unit(row.unit):
                    score += 0.10
                if cand_desc and row_desc and (cand_desc in row_desc or row_desc in cand_desc):
                    score += 0.15
                # v2.8.3: exact/near duplicate material keys are strong evidence
                # that the source line is accounted for even when the raw OCR
                # text differs by punctuation, wrapping, or quote characters.
                if canonical_material_key(cand.parsed) == canonical_material_key(row):
                    score += 0.25
                elif loose_same_material(cand.parsed, row, min_ratio=0.88):
                    score += 0.18
            score = min(score, 1.0)
            if score > best_score:
                best_score = score
                best_idx = i
        return best_idx, best_score

    def find_existing_material_match(self, parsed: MaterialLine, rows: List[MaterialLine]) -> int:
        """Find an already-extracted row that accounts for a parsed source line.

        This is stricter than best_row_match in one way (qty must agree) but
        more forgiving about OCR punctuation.  It is used to avoid adding a
        duplicate recovered row when the row already exists.
        """
        if not parsed:
            return -1
        target_key = canonical_material_key(parsed)
        for i, row in enumerate(rows):
            if (row.review_status or "").lower() == "rejected":
                continue
            if canonical_material_key(row) == target_key:
                return i
            if loose_same_material(parsed, row, min_ratio=0.88):
                return i
        return -1

    def audit_matched_fields(self, cand: SourceMaterialCandidate, row: MaterialLine) -> None:
        expected = cand.parsed
        if not expected:
            return
        issues: List[str] = []
        suggestions: Dict[str, str] = {}
        if expected.qty and str(expected.qty).strip() != str(row.qty).strip():
            issues.append(f"qty should likely be {expected.qty}")
            suggestions["qty"] = expected.qty
        if expected.unit and normalize_unit(expected.unit) != normalize_unit(row.unit):
            issues.append(f"unit should likely be {normalize_unit(expected.unit)}")
            suggestions["unit"] = normalize_unit(expected.unit)
        expected_desc = self._norm(expected.description)
        row_desc = self._norm(row.description)
        if expected_desc and row_desc:
            sim = SequenceMatcher(None, expected_desc, row_desc).ratio()
            if sim < 0.62 and not (expected_desc in row_desc or row_desc in expected_desc):
                issues.append("description may not match the source line")
                suggestions["description"] = expected.description
        if not issues:
            row.watchdog_status = row.watchdog_status or "covered"
            row.watchdog_message = row.watchdog_message or "Source line accounted for."
            return
        row.watchdog_status = "suggestion"
        row.watchdog_message = "; ".join(issues)
        row.ai_status = "suggestion" if (row.ai_status or "").lower() not in {"error", "warning"} else row.ai_status
        row.ai_message = append_note(row.ai_message, "AI Watchdog: " + row.watchdog_message)
        row.ai_suggest_qty = row.ai_suggest_qty or suggestions.get("qty", "")
        row.ai_suggest_unit = row.ai_suggest_unit or suggestions.get("unit", "")
        row.ai_suggest_description = row.ai_suggest_description or suggestions.get("description", "")
        row.notes = append_note(row.notes, "AI Watchdog: source-vs-row mismatch")
        row.highlight_color = "yellow"

    def _norm(self, text: str) -> str:
        text = (text or "").lower()
        text = re.sub(r"[^a-z0-9/\"'\- ]+", " ", text)
        text = re.sub(r"\s+", " ", text).strip()
        return text


class SourceCoverageAuditor:
    """Non-destructive final audit used by the UI and the native-text route.

    It proves that every material-looking source line is either matched to a row,
    available as a possible missed row, or intentionally ignored/rejected by the user.
    """

    def __init__(self, parser: MaterialParser, log=None) -> None:
        self.parser = parser
        self.log = log or (lambda msg: None)
        self.watchdog = ExtractionWatchdog(parser, log=log)

    def audit(self, rows: List[MaterialLine], merged_text: str, mutate_rows: bool = True) -> Dict[str, Any]:
        rows = list(rows or [])
        ledger = self.watchdog.build_source_ledger(merged_text or "")
        used: set[int] = set()
        matched = 0
        mismatches: List[Tuple[int, str]] = []
        missing: List[SourceMaterialCandidate] = []
        for cand in ledger:
            idx, score = self.watchdog.best_row_match(cand, rows, used)
            if idx >= 0 and score >= 0.58:
                matched += 1
                used.add(idx)
                row = rows[idx]
                before = row.watchdog_message or ""
                if mutate_rows:
                    self.watchdog.audit_matched_fields(cand, row)
                msg = row.watchdog_message or before
                if (row.watchdog_status or "").lower() in {"warning", "suggestion"} or (msg and msg != "Source line accounted for."):
                    mismatches.append((idx, msg or "Source fields may not match."))
            else:
                # v2.8.3: Before reporting a source line as missing, give it a
                # final strict duplicate/accountability check.  This catches OCR
                # lines like '-(2) Safety Glasses' when the table already contains
                # Qty=2 / Description=Safety Glasses but the exact source text
                # span/highlight was not linked.
                dup_idx = self.watchdog.find_existing_material_match(cand.parsed, rows) if cand.parsed else -1
                if dup_idx >= 0:
                    matched += 1
                    used.add(dup_idx)
                    if mutate_rows:
                        row = rows[dup_idx]
                        row.source_ledger_index = row.source_ledger_index or cand.index
                        if not row.source_line:
                            row.source_line = cand.source_line
                        if row.source_start is None or row.source_start < 0:
                            row.source_start = cand.source_start
                            row.source_end = cand.source_end
                        row.watchdog_status = row.watchdog_status or "covered"
                        row.watchdog_message = row.watchdog_message or "Source line accounted for by final audit duplicate-tightener."
                else:
                    missing.append(cand)

        duplicate_rows: List[int] = []
        seen: Dict[Tuple[str, str, str, str], int] = {}
        for i, row in enumerate(rows):
            if (row.review_status or "").lower() == "rejected":
                continue
            key = warn_key(row)
            dup_of = -1
            if key in seen:
                dup_of = seen[key]
            else:
                for j, prev in enumerate(rows[:i]):
                    if (prev.review_status or "").lower() == "rejected":
                        continue
                    if loose_same_material(prev, row, min_ratio=0.90):
                        dup_of = j
                        break
            if dup_of >= 0:
                duplicate_rows.append(i)
                if mutate_rows:
                    row.ai_status = row.ai_status or "warning"
                    row.ai_message = append_note(row.ai_message, f"Final Audit: possible duplicate of row #{dup_of + 1}.")
                    row.notes = append_note(row.notes, "Final Audit: possible duplicate")
                    row.highlight_color = "yellow"
            else:
                seen[key] = i

        pending = sum(1 for r in rows if (r.review_status or "pending") == "pending")
        confirmed = sum(1 for r in rows if (r.review_status or "") in {"confirmed", "corrected"})
        rejected = sum(1 for r in rows if (r.review_status or "") == "rejected")
        issues = len(missing) + len(mismatches) + len(duplicate_rows) + pending
        score = 100
        if ledger:
            score = int(round(100 * matched / max(1, len(ledger))))
        if issues:
            score = max(0, min(score, 100 - min(60, issues * 4)))

        reason_lines = []
        for i, row in enumerate(rows, start=1):
            reasons = []
            if row.source_line:
                reasons.append("source linked")
            if row.qty:
                reasons.append("qty found")
            else:
                reasons.append("qty missing")
            if row.unit:
                reasons.append("unit found")
            if row.description:
                reasons.append("description found")
            if (row.ai_status or "").lower() in {"warning", "error", "suggestion"}:
                reasons.append("review flagged")
            reason_lines.append(f"Row {i}: " + ", ".join(reasons))

        return {
            "ledger_count": len(ledger),
            "row_count": len(rows),
            "matched_count": matched,
            "confirmed_count": confirmed,
            "pending_count": pending,
            "rejected_count": rejected,
            "missing": missing,
            "missing_count": len(missing),
            "mismatches": mismatches,
            "mismatch_count": len(mismatches),
            "duplicates": duplicate_rows,
            "duplicate_count": len(duplicate_rows),
            "score": score,
            "issues": issues,
            "reason_lines": reason_lines,
        }


class ScanEngine:
    def __init__(self, db: SmartScanDB, settings: ScanSettings, log=None,
                 on_item_reviewed=None) -> None:
        self.db = db
        self.settings = settings
        self.log = log or (lambda msg: None)
        self.on_item_reviewed = on_item_reviewed  # callback(row) called after each item is AI-reviewed
        self.extractor = FileExtractor(settings, log=self.log)
        self.parser = MaterialParser(db, settings)
        # Build Anthropic reviewer (only active when api key is provided)
        api_key = getattr(settings, "anthropic_api_key", "") or os.environ.get("ANTHROPIC_API_KEY", "")
        self._anthropic = AnthropicReviewer(api_key, log=self.log) if (api_key and getattr(settings, "enable_anthropic_review", False)) else None

    def scan_file(self, path: Path, force_rescan: bool = False) -> ScanResult:
        start = time.time()
        path = Path(path)
        file_hash = sha256_file(path)
        # Spreadsheets are fast/deterministic and are actively improved by the
        # SmartScan structured parser.  Do not reuse an older cached spreadsheet
        # result here, otherwise fixes to Excel parsing can appear to "do nothing"
        # and stale 5-line/duplicate results keep coming back until the user manually
        # Force Rescans.  PDFs/images still use the normal cache behavior.
        spreadsheet_source = (path.suffix.lower() in SUPPORTED_SPREADSHEET_EXTS)
        if self.settings.enable_cache and not force_rescan and not spreadsheet_source:
            cached = self.db.latest_cached_scan(file_hash)
            if cached:
                cached.engine_summary = "cache + " + (cached.engine_summary or "")
                return cached

        warnings: List[str] = []
        native_text, ocr_text, table_text, warn, summary = self.extractor.extract(path)
        warnings.extend(warn)
        merged = self.merge_sources(native_text, ocr_text, table_text)

        # Structured-grid fast path: spreadsheets and Word tables have explicit
        # columns, so map columns -> fields and emit every data row.  This avoids
        # the fuzzy, electrical-tuned text heuristics that were dropping valid
        # generic rows (the cause of spreadsheets not scanning to 100%).  It is also
        # faster: the OCR watchdog and source-coverage auditor are skipped because
        # there is nothing to "recover" from an exact grid.
        structured = getattr(self.extractor, "last_structured", None)
        materials = None
        conf = 0
        if structured and structured.get("matrix"):
            try:
                s_rows, s_conf = self.parser.parse_structured_rows(structured["matrix"])
            except Exception as e:
                s_rows, s_conf = [], 0
                warnings.append(f"Structured table parse failed ({e}); using text path.")
            if s_rows:
                materials, conf = s_rows, s_conf
                kind = structured.get("kind", "table")
                warnings.append(
                    f"Smart route: structured {kind} detected; used direct column "
                    f"mapping and captured {len(s_rows)} row(s) (no fuzzy text filtering)."
                )

        if materials is None:
            # Ultimate Smart Scanner routing:
            # - If the original file already has meaningful readable/selectable text,
            #   use the proven v2.5 native-text strategy.
            # - If the file is image-based and needs OCR, use the v2.6 OCR/watchdog
            #   strategy that audits OCR source lines and recovers missing rows.
            use_native_text_strategy = self._should_use_native_text_strategy(native_text, table_text)
            if use_native_text_strategy:
                materials, conf = self.parser.parse_native_text_strategy(merged, table_text)
                # v2.8: use the watchdog as a silent verifier only.  Native/selectable-text
                # PDFs stay on the v2.5 extraction path, but the secondary audit can still
                # warn about possible source lines that need review.
                audit = SourceCoverageAuditor(self.parser, log=self.log).audit(materials, merged, mutate_rows=True)
                if audit.get("missing_count") or audit.get("mismatch_count") or audit.get("duplicate_count"):
                    warnings.append(
                        "Final Audit verifier: "
                        f"{audit.get('missing_count', 0)} possible missed line(s), "
                        f"{audit.get('mismatch_count', 0)} field mismatch(es), "
                        f"{audit.get('duplicate_count', 0)} duplicate(s)."
                    )
                warnings.append("Smart route: readable original text detected; using v2.5 native-text extraction path.")
            else:
                materials, conf = self.parser.parse(merged, table_text)

                # v2.6: AI Watchdog / Extraction Auditor.  It builds a source-line
                # ledger from the OCR/source text and checks that every material-looking
                # source line became exactly one row before the UI sees the result.
                watchdog = ExtractionWatchdog(self.parser, log=self.log)
                materials, audit_warnings = watchdog.audit(materials, merged)
                warnings.extend(audit_warnings)
                warnings.append("Smart route: image/OCR-first document detected; using v2.6 OCR watchdog extraction path.")

        materials = EvidenceMapper(path, merged).annotate(materials)
        materials = order_materials_for_review(materials)

        # Snapshot original scanner guesses before any correction
        for r in materials:
            r.snapshot_original()

        if self.settings.smart_review_during_scan and self.settings.enable_rule_review:
            reviewer = AIReviewer(self.settings, log=self.log)
            materials = reviewer.review(materials, merged)
            before_count = len(materials)
            if self.settings.hide_definite_noise_rows:
                # Do not let OCR/PDF-oriented noise hiding remove structured Excel
                # rows. Spreadsheet rows came from explicit cells/headers, so
                # catalog-only descriptions are valid and should stay visible for
                # one-at-a-time review/highlighting.
                materials = [r for r in materials if (getattr(r, "source_view", "") == "table") or not is_definite_noise_row(r)]
                hidden_count = before_count - len(materials)
                if hidden_count:
                    warnings.append(f"Smart review hid {hidden_count} definite contact/address/noise row(s).")
            materials = order_materials_for_review(materials)
            conf = self.parser.overall_confidence(materials, merged)

        # v2.0: per-item Anthropic Claude review (runs in scan thread)
        if self._anthropic:
            for r in materials:
                self._anthropic.apply(r)
                if self.on_item_reviewed:
                    self.on_item_reviewed(r)  # stream updates to UI

        result = ScanResult(
            file_path=str(path),
            file_hash=file_hash,
            engine_summary=summary,
            native_text=native_text,
            ocr_text=ocr_text,
            table_text=table_text,
            merged_text=merged,
            materials=materials,
            confidence=conf,
            warnings=warnings,
            elapsed_sec=round(time.time() - start, 2),
        )
        self.db.save_scan(result)
        return result

    def _should_use_native_text_strategy(self, native_text: str, table_text: str = "") -> bool:
        """Return True when the original file has enough usable native text.

        This is intentionally conservative: a scanned/image PDF can still produce a
        tiny amount of embedded metadata or stray text.  We only choose the v2.5
        native path when the native stream has real words plus material/table
        signals.  Otherwise the v2.6 OCR watchdog path remains the default.
        """
        native = normalize_space(native_text or "")
        if not native.strip():
            return False

        printable = sum(1 for ch in native if ch.isprintable() or ch in "\r\n\t")
        printable_ratio = printable / max(1, len(native))
        if printable_ratio < 0.82:
            return False

        lines = [ln.strip() for ln in split_lines_preserve(native) if ln.strip()]
        nontrivial_lines = [ln for ln in lines if len(re.sub(r"\W+", "", ln)) >= 3]
        if len(nontrivial_lines) < 4 and len(native) < 160:
            return False

        header_hits = sum(1 for ln in nontrivial_lines if self.parser.is_qty_description_header(ln))
        material_hits = 0
        table_like_hits = 0
        for ln in nontrivial_lines[:250]:
            if has_table_spacing(ln) or "|" in ln or "\t" in ln:
                table_like_hits += 1
            if material_signal_score(ln) >= 45 or has_explicit_uom(ln):
                material_hits += 1
            if material_hits >= 3:
                break

        if header_hits and (material_hits >= 1 or table_like_hits >= 2):
            return True
        if material_hits >= 3:
            return True
        if len(native) >= 500 and material_hits >= 2:
            return True
        return False

    def merge_sources(self, native: str, ocr: str, table: str) -> str:
        parts = []
        if table.strip():
            parts.append("--- TABLE / POSITIONAL CANDIDATES ---\n" + table)
        if native.strip():
            parts.append("--- NATIVE TEXT ---\n" + native)
        if ocr.strip():
            parts.append("--- OCR TEXT ---\n" + ocr)
        return normalize_space("\n\n".join(parts))


class SettingsWindow(tk.Toplevel):
    def __init__(self, master: tk.Tk, settings: ScanSettings, on_save) -> None:
        super().__init__(master)
        self.title("SmartScan Settings")
        self.geometry("760x680")
        self.minsize(680, 560)
        self.settings = settings
        self.on_save = on_save
        self.vars: Dict[str, Any] = {}
        self._build()
        try:
            _apply_mainbox_theme_to_tk_widget(self, _smartscan_mainbox_theme_config())
        except Exception:
            pass
        self.transient(master)
        self.grab_set()

    def _build(self) -> None:
        root = ttk.Frame(self, padding=12)
        root.pack(fill="both", expand=True)
        ttk.Label(root, text="OCR / Scanner Settings", font=("Segoe UI", 14, "bold")).pack(anchor="w")
        ttk.Label(root, text="Optional engines are used only if installed. The app will still run without them.").pack(anchor="w", pady=(0, 8))

        canvas = tk.Canvas(root, highlightthickness=0)
        scroll = ttk.Scrollbar(root, orient="vertical", command=canvas.yview)
        body = ttk.Frame(canvas)
        body.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=body, anchor="nw")
        canvas.configure(yscrollcommand=scroll.set)
        canvas.pack(side="left", fill="both", expand=True)
        scroll.pack(side="right", fill="y")

        self._combo(body, "mode", "Default scan mode", ["Fast", "Normal", "Slow Careful"])
        self._check(body, "enable_native_pdf", "Enable native PDF extraction with PyMuPDF")
        self._check(body, "enable_text_files", "Enable TXT/CSV/HTML/EML text extraction")
        self._check(body, "enable_preprocess", "Preprocess images before OCR")
        self._check(body, "enable_tesseract", "Enable Tesseract OCR")
        self._check(body, "enable_paddleocr", "Enable PaddleOCR if installed")
        self._check(body, "enable_surya_cli", "Enable Surya CLI hook if installed")
        self._check(body, "enable_table_guess", "Enable layout/table guessing")
        self._check(body, "enable_header_aware_extraction", "Prefer Qty/Description header-aware extraction")
        self._check(body, "exclude_contact_info", "Exclude phone numbers, emails, websites, and contact lines")
        self._check(body, "strict_address_filter", "Exclude address/job-site lines unless clearly material")
        self._check(body, "enable_cache", "Cache scan results by file hash")
        self._check(body, "learn_from_corrections", "Learn from saved corrections")
        self._check(body, "auto_deep_scan_low_confidence", "Auto-suggest deep scan when confidence is low")
        self._spin(body, "confidence_threshold", "Low confidence threshold", 0, 100)

        ttk.Separator(body).pack(fill="x", pady=8)
        ttk.Label(body, text="AI Review / Learning", font=("Segoe UI", 11, "bold")).pack(anchor="w")
        self._check(body, "auto_ai_review_after_scan", "Run AI review after each scan")
        self._check(body, "smart_review_during_scan", "Smart-review each item during extraction before showing results")
        self._check(body, "hide_definite_noise_rows", "Hide definite phone/address/contact noise rows from output")
        self._check(body, "show_all_review_highlights", "Show green/yellow/red highlights on the file preview")
        self._check(body, "enable_rule_review", "Enable fast rule-based review")
        self._check(body, "enable_ollama_ai_review", "Enable optional Ollama AI review")
        self._combo(body, "ai_review_mode", "AI review mode", ["Fast", "Normal", "Slow Careful"])
        self._check(body, "ai_review_only_suspicious", "AI-review only suspicious rows in Normal mode")
        self._entry(body, "ai_review_ollama_model", "Ollama review model")
        self._entry(body, "ai_review_ollama_url", "Ollama API URL")

        ttk.Separator(body).pack(fill="x", pady=8)
        ttk.Label(body, text="Anthropic Claude (v2.0 real-time per-item review)", font=("Segoe UI", 11, "bold")).pack(anchor="w")
        ttk.Label(body, text="API key is saved in the DB. Leave blank to use ANTHROPIC_API_KEY env var.").pack(anchor="w", pady=(0, 2))
        self._entry_pw(body, "anthropic_api_key", "Anthropic API key")
        self._check(body, "enable_anthropic_review", "Enable Anthropic Claude per-item review (streams during scan)")

        ttk.Separator(body).pack(fill="x", pady=8)
        ttk.Label(body, text="Performance", font=("Segoe UI", 11, "bold")).pack(anchor="w")
        self._spin(body, "max_pages_fast", "Fast mode max PDF pages", 1, 500)
        self._spin(body, "max_pages_normal", "Normal mode max PDF pages", 1, 500)
        self._spin(body, "max_pages_careful", "Slow Careful max PDF pages", 1, 1000)
        self._spin(body, "render_dpi_fast", "Fast render DPI", 72, 600)
        self._spin(body, "render_dpi_normal", "Normal render DPI", 72, 600)
        self._spin(body, "render_dpi_careful", "Slow Careful render DPI", 72, 600)

        ttk.Separator(body).pack(fill="x", pady=8)
        ttk.Label(body, text="Engine Paths / Language", font=("Segoe UI", 11, "bold")).pack(anchor="w")
        self._entry(body, "tesseract_cmd", "Tesseract executable path (optional)")
        self._entry(body, "tesseract_lang", "Tesseract language")
        self._entry(body, "paddle_lang", "PaddleOCR language")
        self._entry(body, "surya_command", "Surya CLI command")
        self._entry(body, "output_dir", "Default export folder")

        btns = ttk.Frame(root)
        btns.pack(fill="x", pady=(8, 0))
        ttk.Button(btns, text="Save Settings", command=self._save).pack(side="right")
        ttk.Button(btns, text="Cancel", command=self.destroy).pack(side="right", padx=(0, 8))

    def _row(self, parent, label: str):
        frame = ttk.Frame(parent)
        frame.pack(fill="x", pady=3)
        ttk.Label(frame, text=label, width=34).pack(side="left")
        return frame

    def _check(self, parent, attr: str, label: str) -> None:
        var = tk.BooleanVar(value=bool(getattr(self.settings, attr)))
        self.vars[attr] = var
        ttk.Checkbutton(parent, text=label, variable=var).pack(anchor="w", pady=2)

    def _entry(self, parent, attr: str, label: str) -> None:
        frame = self._row(parent, label)
        var = tk.StringVar(value=str(getattr(self.settings, attr)))
        self.vars[attr] = var
        ttk.Entry(frame, textvariable=var).pack(side="left", fill="x", expand=True)

    def _entry_pw(self, parent, attr: str, label: str) -> None:
        """Password-masked entry for API keys."""
        frame = self._row(parent, label)
        var = tk.StringVar(value=str(getattr(self.settings, attr)))
        self.vars[attr] = var
        entry = ttk.Entry(frame, textvariable=var, show="*")
        entry.pack(side="left", fill="x", expand=True)
        # show/hide toggle
        show_var = tk.BooleanVar(value=False)
        def toggle():
            entry.config(show="" if show_var.get() else "*")
        ttk.Checkbutton(frame, text="Show", variable=show_var, command=toggle).pack(side="left")

    def _spin(self, parent, attr: str, label: str, frm: int, to: int) -> None:
        frame = self._row(parent, label)
        var = tk.IntVar(value=int(getattr(self.settings, attr)))
        self.vars[attr] = var
        ttk.Spinbox(frame, from_=frm, to=to, textvariable=var, width=8).pack(side="left")

    def _combo(self, parent, attr: str, label: str, values: List[str]) -> None:
        frame = self._row(parent, label)
        var = tk.StringVar(value=str(getattr(self.settings, attr)))
        self.vars[attr] = var
        ttk.Combobox(frame, textvariable=var, values=values, state="readonly", width=20).pack(side="left")

    def _save(self) -> None:
        for attr, var in self.vars.items():
            try:
                setattr(self.settings, attr, var.get())
            except Exception:
                pass
        self.on_save(self.settings)
        self.destroy()


class SmartScanApp:
    def __init__(self, mainbox_files: Optional[List[str]] = None, mainbox_json_out: str = "", mainbox_review_mode: bool = False) -> None:
        self.db = SmartScanDB()
        self.settings = self.db.load_settings()
        self.results: List[ScanResult] = []
        self.current_result: Optional[ScanResult] = None
        self.job_q: "queue.Queue[Tuple[str, bool, ScanSettings]]" = queue.Queue()
        self.ui_q: "queue.Queue[Tuple[str, Any]]" = queue.Queue()
        self.worker_thread: Optional[threading.Thread] = None
        self.busy = False
        self.mainbox_review_mode = bool(mainbox_review_mode)
        self.mainbox_files = [Path(p) for p in (mainbox_files or []) if str(p or "").strip()]
        self.mainbox_json_out = str(mainbox_json_out or "")

        if TkinterDnD is not None:
            self.root = TkinterDnD.Tk()
        else:
            self.root = tk.Tk()
        self.root.title(("SmartScan Review for MaINbox" if self.mainbox_review_mode else f"{APP_NAME} v{APP_VERSION}"))
        self.root.geometry("1440x860")
        self.root.minsize(1080, 680)

        self.tree_item_map: Dict[str, MaterialLine] = {}
        self.current_text_source = "merged"
        self.preview_photo = None
        self.preview_pil_image = None
        self.preview_page = 0
        self.preview_total_pages = 0
        self.preview_scale = 1.0
        self.preview_original_size = (0, 0)
        self.preview_display_size = (0, 0)
        self.hl_span_var = tk.StringVar(value="span: —")  # initialised before _build_ui
        # v2.2 bbox draw state
        self._bbox_draw_start: Optional[Tuple[int, int]] = None   # canvas pixel coords
        self._bbox_rect_id: Optional[int] = None                  # temp canvas item id
        # v2.4 scroll preservation
        self._render_scroll_target: Optional[Tuple[float, float]] = None
        self._render_saved_scroll: Tuple[float, float] = (0.0, 0.0)
        self._render_keep_scroll: bool = False
        self._style()
        self._build_ui()
        try:
            _apply_mainbox_theme_to_tk_widget(self.root, getattr(self, "_mainbox_theme_cfg", _smartscan_mainbox_theme_config()))
        except Exception:
            pass
        self.root.after(150, self._pump_ui_queue)
        self._start_worker()
        if self.mainbox_review_mode and self.mainbox_files:
            self.root.after(400, self._mainbox_bootstrap_review)

    def _style(self) -> None:
        self._mainbox_theme_cfg = _smartscan_mainbox_theme_config()
        cfg = self._mainbox_theme_cfg
        try:
            self.root.configure(bg=cfg.get("bg", "#1e2229"))
            self.root.option_add("*Font", (cfg.get("font_family", "Segoe UI"), int(cfg.get("font_size", 9)), cfg.get("font_weight", "normal")))
        except Exception:
            pass
        style = ttk.Style(self.root)
        try:
            style.theme_use("clam")
        except Exception:
            pass
        base_font = (cfg.get("font_family", "Segoe UI"), int(cfg.get("font_size", 9)), cfg.get("font_weight", "normal"))
        title_font = (cfg.get("font_family", "Segoe UI"), max(12, int(cfg.get("font_size", 9)) + 7), "bold")
        style.configure(".", background=cfg.get("bg", "#1e2229"), foreground=cfg.get("fg", "#dde6f0"), fieldbackground=cfg.get("entry_bg", "#252b34"), font=base_font, bordercolor=cfg.get("bg3", "#2c3340"), troughcolor=cfg.get("bg2", "#252b34"), selectbackground=cfg.get("sel_bg", "#3a5a8a"), selectforeground=cfg.get("sel_fg", "#ffffff"))
        style.configure("TFrame", background=cfg.get("bg", "#1e2229"))
        style.configure("TLabelframe", background=cfg.get("bg2", "#252b34"), foreground=cfg.get("accent", "#4a9eff"))
        style.configure("TLabelframe.Label", background=cfg.get("bg2", "#252b34"), foreground=cfg.get("accent", "#4a9eff"), font=base_font)
        style.configure("TLabel", background=cfg.get("bg", "#1e2229"), foreground=cfg.get("fg", "#dde6f0"), font=base_font)
        style.configure("TButton", background=cfg.get("button_bg", "#2e3a4e"), foreground=cfg.get("button_fg", "#d6e4f7"), font=base_font, padding=5)
        style.map("TButton", background=[("active", cfg.get("accent2", "#3a7fd5")), ("disabled", cfg.get("bg2", "#252b34"))], foreground=[("disabled", cfg.get("fg_dim", "#8fa0b8"))])
        style.configure("TEntry", fieldbackground=cfg.get("entry_bg", "#252b34"), foreground=cfg.get("fg", "#dde6f0"), insertcolor=cfg.get("fg", "#dde6f0"), font=base_font)
        style.configure("TCombobox", fieldbackground=cfg.get("entry_bg", "#252b34"), background=cfg.get("button_bg", "#2e3a4e"), foreground=cfg.get("fg", "#dde6f0"), arrowcolor=cfg.get("accent", "#4a9eff"), font=base_font)
        style.configure("Treeview", background=cfg.get("bg3", "#2c3340"), foreground=cfg.get("fg", "#dde6f0"), fieldbackground=cfg.get("bg3", "#2c3340"), rowheight=max(22, int(cfg.get("font_size", 9)) + 13), font=base_font, borderwidth=0)
        style.configure("Treeview.Heading", background=cfg.get("bg2", "#252b34"), foreground=cfg.get("accent", "#4a9eff"), font=(cfg.get("font_family", "Segoe UI"), int(cfg.get("font_size", 9)), "bold"))
        style.map("Treeview", background=[("selected", cfg.get("sel_bg", "#3a5a8a"))], foreground=[("selected", cfg.get("sel_fg", "#ffffff"))])
        style.configure("TNotebook", background=cfg.get("bg", "#1e2229"), borderwidth=0)
        style.configure("TNotebook.Tab", background=cfg.get("bg2", "#252b34"), foreground=cfg.get("fg", "#dde6f0"), padding=(8, 4), font=base_font)
        style.map("TNotebook.Tab", background=[("selected", cfg.get("bg3", "#2c3340")), ("active", cfg.get("button_bg", "#2e3a4e"))], foreground=[("selected", cfg.get("accent", "#4a9eff"))])
        style.configure("Title.TLabel", background=cfg.get("bg", "#1e2229"), foreground=cfg.get("accent", "#4a9eff"), font=title_font)
        style.configure("Good.TLabel", background=cfg.get("bg", "#1e2229"), foreground="#42c76b")
        style.configure("Warn.TLabel", background=cfg.get("bg", "#1e2229"), foreground="#f0b24a")
        style.configure("Bad.TLabel", background=cfg.get("bg", "#1e2229"), foreground="#ff6b7a")

    def _build_ui(self) -> None:
        root = ttk.Frame(self.root, padding=10)
        root.pack(fill="both", expand=True)

        header = ttk.Frame(root)
        header.pack(fill="x")
        ttk.Label(header, text="SmartScan Extractor", style="Title.TLabel").pack(side="left")
        ttk.Label(header, text=("  Correct the list, save learning, then click Done to return to MaINbox" if self.mainbox_review_mode else "  RFQ / material data extraction engine")).pack(side="left", padx=(8, 0))

        top_btns = ttk.Frame(root)
        top_btns.pack(fill="x", pady=(8, 8))
        ttk.Button(top_btns, text="Open Files", command=self.open_files).pack(side="left")
        ttk.Button(top_btns, text="Fast Scan", command=lambda: self.scan_selected_mode("Fast", False)).pack(side="left", padx=4)
        ttk.Button(top_btns, text="Normal Scan", command=lambda: self.scan_selected_mode("Normal", False)).pack(side="left", padx=4)
        ttk.Button(top_btns, text="Slow Careful Scan", command=lambda: self.scan_selected_mode("Slow Careful", True)).pack(side="left", padx=4)
        ttk.Button(top_btns, text="Force Rescan", command=lambda: self.scan_selected_mode(self.settings.mode, True)).pack(side="left", padx=4)
        ttk.Button(top_btns, text="Run AI Review", command=self.run_ai_review).pack(side="left", padx=(12, 4))
        ttk.Button(top_btns, text="Final Audit", command=self.show_final_audit).pack(side="left", padx=4)
        ttk.Button(top_btns, text="Missed Lines", command=self.show_possible_missed_lines).pack(side="left", padx=4)
        ttk.Button(top_btns, text="Save Corrections", command=self.save_corrections).pack(side="left", padx=(16, 4))
        ttk.Button(top_btns, text="Export CSV", command=self.export_csv).pack(side="left", padx=4)
        ttk.Button(top_btns, text="Export JSON", command=self.export_json).pack(side="left", padx=4)
        if self.mainbox_review_mode:
            ttk.Button(top_btns, text="Done → Send Back to MaINbox", command=self.mainbox_done_send_back).pack(side="right", padx=(10, 0))
            ttk.Button(top_btns, text="Cancel", command=self.mainbox_cancel_review).pack(side="right", padx=4)
        ttk.Button(top_btns, text="Settings", command=self.open_settings).pack(side="right")

        self.status_var = tk.StringVar(value="Open or drag files to begin.")
        self.conf_var = tk.StringVar(value="Confidence: —")
        status_bar = ttk.Frame(root)
        status_bar.pack(fill="x", pady=(0, 6))
        ttk.Label(status_bar, textvariable=self.status_var).pack(side="left")
        ttk.Label(status_bar, textvariable=self.conf_var).pack(side="right")

        panes = ttk.Panedwindow(root, orient="horizontal")
        panes.pack(fill="both", expand=True)

        left = ttk.Frame(panes, padding=(0, 0, 8, 0))
        mid = ttk.Frame(panes, padding=(8, 0, 8, 0))
        right = ttk.Frame(panes, padding=(8, 0, 0, 0))
        panes.add(left, weight=1)
        panes.add(mid, weight=2)
        panes.add(right, weight=3)
        # Start the Files pane visibly open. Users can still drag the divider closed/open.
        try:
            root.after(250, lambda: panes.sashpos(0, max(185, min(285, int(root.winfo_width() * 0.16)))))
        except Exception:
            pass

        ttk.Label(left, text="Files").pack(anchor="w")
        self.file_list = tk.Listbox(left, selectmode="extended", height=12)
        self.file_list.pack(fill="both", expand=True)
        self.file_list.bind("<<ListboxSelect>>", self.on_file_select)
        if DND_FILES is not None:
            try:
                self.file_list.drop_target_register(DND_FILES)
                self.file_list.dnd_bind("<<Drop>>", self.on_drop)
            except Exception:
                pass
        ttk.Button(left, text="Remove Selected", command=self.remove_selected_files).pack(fill="x", pady=(6, 0))
        ttk.Button(left, text="Clear Files", command=self.clear_files).pack(fill="x", pady=(4, 0))

        ttk.Label(mid, text="Evidence Viewer").pack(anchor="w")
        self.evidence_tabs = ttk.Notebook(mid)
        self.evidence_tabs.pack(fill="both", expand=True)

        text_tab = ttk.Frame(self.evidence_tabs)
        preview_tab = ttk.Frame(self.evidence_tabs)
        self.evidence_tabs.add(text_tab, text="Extracted / Merged Text")
        self.evidence_tabs.add(preview_tab, text="File Preview")

        self.text_box = tk.Text(text_tab, wrap="word", undo=True)
        self.text_box.pack(fill="both", expand=True)
        self.text_box.tag_configure("evidence_current", background="#5f5130", foreground="#ffffff")
        self.text_box.tag_configure("evidence_soft", background="#2f425c", foreground="#dbeafe")
        self.text_box.tag_configure("evidence_ok",         background="#244a32", foreground="#d7f7df")  # green
        self.text_box.tag_configure("evidence_warning",    background="#5a4a22", foreground="#fff4cc")  # yellow
        self.text_box.tag_configure("evidence_error",      background="#5a2a2a", foreground="#ffdede")  # red
        self.text_box.tag_configure("evidence_suggestion", background="#2d4260", foreground="#dcecff")  # blue
        # v2.5.7: selected material row guide - a brighter yellow full-pane band
        # so the blue table row always corresponds to a very visible source-text row.
        self.text_box.tag_configure("evidence_row_select", background="#6b5a1a", foreground="#ffffff")
        # v2.5 manual cross-check: bright orange temporary confirmation highlight
        self.text_box.tag_configure("evidence_confirm", background="#7a4d16", foreground="#ffffff")
        # v2.0 persistent background highlights (all rows at once)
        self.text_box.tag_configure("hl_green",  background="#244a32", foreground="#d7f7df", elide=False)
        self.text_box.tag_configure("hl_yellow", background="#5a4a22", foreground="#fff4cc", elide=False)
        self.text_box.tag_configure("hl_red",    background="#5a2a2a", foreground="#ffdede", elide=False)
        self.text_box.bind("<ButtonRelease-1>", self._on_text_box_click_select_row)

        preview_tools = ttk.Frame(preview_tab)
        preview_tools.pack(fill="x", pady=(0, 4))
        ttk.Button(preview_tools, text="◀ Prev Page", command=self.prev_preview_page).pack(side="left")
        ttk.Button(preview_tools, text="Next Page ▶", command=self.next_preview_page).pack(side="left", padx=4)
        ttk.Label(preview_tools, text="Page").pack(side="left", padx=(12, 2))
        self.preview_page_var = tk.StringVar(value="1")
        self.preview_page_entry = ttk.Entry(preview_tools, textvariable=self.preview_page_var, width=5)
        self.preview_page_entry.pack(side="left")
        self.preview_page_entry.bind("<Return>", lambda e: self.go_preview_page())
        ttk.Button(preview_tools, text="Go", command=self.go_preview_page).pack(side="left", padx=4)
        self.preview_total_var = tk.StringVar(value="/ 1")
        ttk.Label(preview_tools, textvariable=self.preview_total_var).pack(side="left")
        ttk.Button(preview_tools, text="Find Highlight", command=self.find_highlight_on_current_page).pack(side="right")
        ttk.Label(preview_tools, text="  ■ Green=OK  ■ Yellow=Possibly wrong  ■ Red=Definite error").pack(side="right", padx=(0, 10))

        # v2.2 — Bbox draw toolbar: lets user drag a new highlight box on the preview
        bbox_tools = ttk.Frame(preview_tab)
        bbox_tools.pack(fill="x", pady=(0, 2))
        self.bbox_draw_var = tk.BooleanVar(value=False)
        self.bbox_draw_btn = ttk.Checkbutton(
            bbox_tools, text="✏ Draw Highlight Box",
            variable=self.bbox_draw_var,
            command=self._on_bbox_draw_toggle,
        )
        self.bbox_draw_btn.pack(side="left", padx=(4, 8))
        ttk.Label(bbox_tools, text="Drag on preview to set the item's location on the page. Saves with corrections.").pack(side="left")
        self.bbox_coords_var = tk.StringVar(value="")
        ttk.Label(bbox_tools, textvariable=self.bbox_coords_var, foreground="#555555").pack(side="right", padx=4)
        ttk.Button(bbox_tools, text="Clear Box", command=self._clear_preview_bbox).pack(side="right", padx=4)

        preview_frame = ttk.Frame(preview_tab)
        preview_frame.pack(fill="both", expand=True)
        self.preview_canvas = tk.Canvas(preview_frame, background="#f5f5f5")
        self.preview_canvas.pack(side="left", fill="both", expand=True)
        # v2.2 mouse bindings for bbox draw mode
        self.preview_canvas.bind("<ButtonPress-1>",   self._bbox_mouse_press)
        self.preview_canvas.bind("<B1-Motion>",        self._bbox_mouse_drag)
        self.preview_canvas.bind("<ButtonRelease-1>",  self._bbox_mouse_release)
        preview_y = ttk.Scrollbar(preview_frame, orient="vertical", command=self.preview_canvas.yview)
        preview_y.pack(side="right", fill="y")
        preview_x = ttk.Scrollbar(preview_tab, orient="horizontal", command=self.preview_canvas.xview)
        preview_x.pack(side="bottom", fill="x")
        self.preview_canvas.configure(yscrollcommand=preview_y.set, xscrollcommand=preview_x.set)

        mid_tools = ttk.Frame(mid)
        mid_tools.pack(fill="x", pady=(6, 0))
        ttk.Button(mid_tools, text="Copy Text", command=self.copy_text).pack(side="left")
        ttk.Button(mid_tools, text="Show Native", command=lambda: self.show_text_source("native")).pack(side="left", padx=4)
        ttk.Button(mid_tools, text="Show OCR", command=lambda: self.show_text_source("ocr")).pack(side="left", padx=4)
        ttk.Button(mid_tools, text="Show Merged", command=lambda: self.show_text_source("merged")).pack(side="left", padx=4)
        ttk.Button(mid_tools, text="Show Preview", command=self.show_preview_tab).pack(side="left", padx=(12, 4))

        # v2.3 — Simplified highlight editor: select text, click Draw Highlight.
        hl_tools = ttk.Frame(mid)
        hl_tools.pack(fill="x", pady=(2, 0))
        ttk.Label(hl_tools, text="Text highlight:").pack(side="left")
        ttk.Button(hl_tools, text="✏ Draw Highlight",
                   command=self._use_text_selection_as_highlight).pack(side="left", padx=(6, 4))
        ttk.Button(hl_tools, text="↺ Reset",
                   command=self._reset_highlight).pack(side="left", padx=(0, 8))
        ttk.Label(hl_tools, text="← Select text above, then click Draw Highlight",
                  foreground="#666666").pack(side="left")
        ttk.Label(hl_tools, textvariable=self.hl_span_var,
                  foreground="#333333").pack(side="right", padx=4)

        ttk.Label(right, text="Extracted Material Lines - review in source order").pack(anchor="w")
        cols = ("order", "status", "ai", "qty", "unit", "description", "part", "manufacturer", "confidence", "notes")
        self.tree = ttk.Treeview(right, columns=cols, show="headings", selectmode="extended")
        headings = {
            "order": "#", "status": "Status", "ai": "AI Review", "qty": "Qty", "unit": "Unit", "description": "Description", "part": "Part #",
            "manufacturer": "Manufacturer", "confidence": "Conf.", "notes": "Notes",
        }
        widths = {"order": 45, "status": 85, "ai": 105, "qty": 60, "unit": 60, "description": 300, "part": 105, "manufacturer": 110, "confidence": 55, "notes": 210}
        for c in cols:
            self.tree.heading(c, text=headings[c])
            self.tree.column(c, width=widths[c], anchor="w")
        self.tree.pack(fill="both", expand=True)
        self.tree.tag_configure("pending", background="#4a4225", foreground="#fff4cc")
        self.tree.tag_configure("confirmed", background="#243f2c", foreground="#d7f7df")
        self.tree.tag_configure("corrected", background="#2d4260", foreground="#dcecff")
        self.tree.tag_configure("rejected", background="#4a2525", foreground="#ffdede")
        self.tree.tag_configure("ai_ok", background="#243f2c", foreground="#d7f7df")
        self.tree.tag_configure("ai_warning", background="#5a4a22", foreground="#fff4cc")
        self.tree.tag_configure("ai_suggestion", background="#2d4260", foreground="#dcecff")
        self.tree.tag_configure("ai_error", background="#5a2a2a", foreground="#ffdede")
        self.tree.bind("<Double-1>", self.edit_tree_cell)
        self.tree.bind("<<TreeviewSelect>>", self.on_material_select)

        material_btns = ttk.Frame(right)
        material_btns.pack(fill="x", pady=(6, 0))
        ttk.Button(material_btns, text="Confirm ✓", command=self.confirm_selected_rows).pack(side="left")
        ttk.Button(material_btns, text="Reject ✕", command=self.reject_selected_rows).pack(side="left", padx=4)
        ttk.Button(material_btns, text="Accept AI Fix", command=self.accept_ai_fix_selected).pack(side="left", padx=4)
        ttk.Button(material_btns, text="Keep Original", command=self.keep_original_selected).pack(side="left", padx=4)
        ttk.Button(material_btns, text="Next Issue", command=self.select_next_ai_issue).pack(side="left", padx=4)
        ttk.Button(material_btns, text="Next Pending", command=self.select_next_pending).pack(side="left", padx=4)
        ttk.Button(material_btns, text="Add Row", command=self.add_material_row).pack(side="left", padx=(16, 0))
        ttk.Button(material_btns, text="Delete Row", command=self.delete_material_rows).pack(side="left", padx=4)
        ttk.Button(material_btns, text="Copy Materials", command=self.copy_materials).pack(side="left", padx=4)

        bottom = ttk.Frame(root)
        bottom.pack(fill="x", pady=(8, 0))
        ttk.Label(bottom, text="Log / Warnings").pack(anchor="w")
        self.log_box = tk.Text(bottom, height=5, wrap="word")
        self.log_box.pack(fill="x")

        self.files: List[Path] = []

    def _mainbox_bootstrap_review(self) -> None:
        """MaINbox mode: pre-load saved RFQ attachments and start a full SmartScan review."""
        try:
            self.add_files([p for p in self.mainbox_files if p.exists() and p.is_file()])
            if self.files:
                # Select everything so the normal SmartScan worker scans the full RFQ packet.
                try:
                    self.file_list.selection_set(0, "end")
                except Exception:
                    pass
                self.status_var.set("MaINbox RFQ review: scanning attachments in SmartScan...")
                self.scan_selected_mode("Normal", True)
        except Exception as exc:
            self.log(f"MaINbox review bootstrap failed: {exc}")

    def _mainbox_collect_all_rows(self) -> List[MaterialLine]:
        """Return the user-reviewed rows across every SmartScan result.

        The currently visible tree may contain manual edits not yet copied back into
        current_result, so first sync the visible file's tree rows. Then combine all
        non-rejected rows from every scanned file in source order.
        """
        try:
            if self.current_result:
                self.current_result.materials = self.tree_rows()
        except Exception:
            pass
        rows: List[MaterialLine] = []
        for result in self.results or []:
            for r in (result.materials or []):
                if (getattr(r, "review_status", "") or "").lower() == "rejected":
                    continue
                if not ((getattr(r, "description", "") or "").strip() or (getattr(r, "part_number", "") or "").strip()):
                    continue
                rows.append(r)
        if not rows and self.current_result:
            for r in self.tree_rows():
                if (getattr(r, "review_status", "") or "").lower() != "rejected":
                    rows.append(r)
        return rows

    def _mainbox_payload(self, ok: bool, cancelled: bool = False, error: str = "") -> Dict[str, Any]:
        rows = self._mainbox_collect_all_rows() if ok else []
        row_dicts = []
        for res in (self.results or []):
            for r in (res.materials or []):
                if r in rows:
                    d = r.as_dict() if hasattr(r, "as_dict") else dataclasses.asdict(r)
                    d["source_file_path"] = getattr(res, "file_path", "") or ""
                    d["source_file_name"] = Path(getattr(res, "file_path", "") or "").name
                    row_dicts.append(d)
        if not row_dicts:
            for r in rows:
                row_dicts.append(r.as_dict() if hasattr(r, "as_dict") else dataclasses.asdict(r))
        return {
            "ok": bool(ok),
            "cancelled": bool(cancelled),
            "error": error or "",
            "rows": row_dicts,
            "results": [res.as_dict() for res in (self.results or [])],
            "merged_text": "\n\n".join((getattr(res, "merged_text", "") or "") for res in (self.results or []))[:50000],
            "warnings": [w for res in (self.results or []) for w in (getattr(res, "warnings", []) or [])],
            "engine_summary": "; ".join((getattr(res, "engine_summary", "") or "") for res in (self.results or []) if getattr(res, "engine_summary", "")),
        }

    def _write_mainbox_payload(self, payload: Dict[str, Any]) -> None:
        if not self.mainbox_json_out:
            return
        Path(self.mainbox_json_out).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    def mainbox_done_send_back(self) -> None:
        """Save SmartScan learning and return reviewed material rows to MaINbox."""
        if self.busy or not self.job_q.empty():
            messagebox.showinfo("SmartScan", "SmartScan is still scanning. Please wait until the scan is Ready, then click Done again.")
            return
        try:
            if self.current_result:
                self.current_result.materials = self.tree_rows()
            # Save corrections/learning for every scanned file, not just the visible one.
            for result in self.results or []:
                rows = result.materials or []
                self.db.save_corrections(result.file_hash, rows)
                self.db.save_correction_patterns(result.file_hash, rows)
                self.db.save_bbox_corrections(result.file_hash, rows)
                result.confidence = MaterialParser(self.db).overall_confidence(rows, result.merged_text)
                self.db.save_scan(result)
            payload = self._mainbox_payload(True)
            self._write_mainbox_payload(payload)
            self.status_var.set(f"Sent {len(payload.get('rows', []))} reviewed row(s) back to MaINbox.")
            self.root.destroy()
        except Exception as exc:
            payload = self._mainbox_payload(False, error=str(exc))
            payload["traceback"] = traceback.format_exc()
            try:
                self._write_mainbox_payload(payload)
            except Exception:
                pass
            messagebox.showerror("SmartScan", f"Could not send rows back to MaINbox:\n{exc}")

    def mainbox_cancel_review(self) -> None:
        try:
            self._write_mainbox_payload({"ok": False, "cancelled": True, "rows": [], "error": "User cancelled SmartScan review."})
        except Exception:
            pass
        self.root.destroy()

    def _start_worker(self) -> None:
        self.worker_thread = threading.Thread(target=self._worker_loop, daemon=True)
        self.worker_thread.start()

    def _worker_loop(self) -> None:
        while True:
            try:
                file_path, force_rescan, settings = self.job_q.get()
                self.ui_q.put(("status", f"Scanning {Path(file_path).name}..."))
                def _item_reviewed(row, q=self.ui_q):
                    q.put(("item_reviewed", row))
                engine = ScanEngine(
                    self.db, settings,
                    log=lambda m: self.ui_q.put(("log", m)),
                    on_item_reviewed=_item_reviewed,
                )
                result = engine.scan_file(Path(file_path), force_rescan=force_rescan)
                self.ui_q.put(("result", result))
            except Exception:
                self.ui_q.put(("error", traceback.format_exc()))
            finally:
                self.job_q.task_done()
                self.ui_q.put(("job_done", None))

    def _pump_ui_queue(self) -> None:
        try:
            while True:
                typ, payload = self.ui_q.get_nowait()
                if typ == "status":
                    self.status_var.set(payload)
                elif typ == "log":
                    self.log(payload)
                elif typ == "error":
                    self.log(payload)
                    messagebox.showerror("SmartScan Error", payload[:1200])
                elif typ == "result":
                    self.add_result(payload)
                elif typ == "item_reviewed":
                    self._on_item_reviewed(payload)
                elif typ == "job_done":
                    if self.job_q.empty():
                        self.busy = False
                        self.status_var.set("Ready.")
                self.ui_q.task_done()
        except queue.Empty:
            pass
        self.root.after(150, self._pump_ui_queue)

    def log(self, msg: str) -> None:
        self.log_box.insert("end", f"[{datetime.now().strftime('%H:%M:%S')}] {msg}\n")
        self.log_box.see("end")

    def open_files(self) -> None:
        paths = filedialog.askopenfilenames(
            title="Select files to scan",
            filetypes=[
                ("Supported files", "*.pdf *.png *.jpg *.jpeg *.bmp *.tif *.tiff *.webp *.txt *.csv *.json *.eml *.html *.htm"),
                ("All files", "*.*"),
            ],
        )
        self.add_files([Path(p) for p in paths])

    def on_drop(self, event) -> None:
        raw = event.data
        # Tcl list parsing handles paths with spaces better.
        try:
            paths = [Path(p) for p in self.root.tk.splitlist(raw)]
        except Exception:
            paths = [Path(raw)]
        self.add_files(paths)

    def add_files(self, paths: List[Path]) -> None:
        added = 0
        for p in paths:
            if not p.exists() or not p.is_file():
                continue
            if p.suffix.lower() not in SUPPORTED_DOC_EXTS:
                self.log(f"Unsupported file skipped: {p.name}")
                continue
            if p not in self.files:
                self.files.append(p)
                self.file_list.insert("end", str(p))
                added += 1
        if added:
            self.status_var.set(f"Added {added} file(s). Choose a scan mode.")

    def selected_paths(self) -> List[Path]:
        idxs = list(self.file_list.curselection())
        if not idxs:
            return self.files[:]
        return [self.files[i] for i in idxs]

    def scan_selected_mode(self, mode: str, force_rescan: bool) -> None:
        paths = self.selected_paths()
        if not paths:
            messagebox.showinfo("SmartScan", "Please open one or more files first.")
            return
        settings = ScanSettings.from_dict(self.settings.as_dict())
        settings.mode = mode
        self.busy = True
        self.status_var.set(f"Queued {len(paths)} file(s) for {mode} scan...")
        for p in paths:
            self.job_q.put((str(p), force_rescan, settings))

    def add_result(self, result: ScanResult) -> None:
        # replace existing result for same path/hash in current session
        self.results = [r for r in self.results if not (r.file_path == result.file_path and r.file_hash == result.file_hash)]
        self.results.append(result)
        self.current_result = result
        self._apply_stored_bboxes(result)  # v2.2: pre-fill user-corrected bbox locations
        self.show_result(result)
        warn = f" Warnings: {len(result.warnings)}." if result.warnings else ""
        self.log(f"Scanned {Path(result.file_path).name}: {len(result.materials)} material line(s), confidence {result.confidence}%, {result.elapsed_sec}s. Engines: {result.engine_summary}.{warn}")
        for w in result.warnings:
            self.log("Warning: " + w)
        if self.settings.auto_deep_scan_low_confidence and result.confidence < self.settings.confidence_threshold:
            self.log("Low confidence. Try Slow Careful Scan or enable more OCR engines in Settings.")
        if self.settings.auto_ai_review_after_scan and self.settings.enable_rule_review:
            # v1.8 usually reviews during scan. Only run a UI review pass if the
            # cached/legacy result has no AI status yet.
            if not any((m.ai_status or "").strip() for m in result.materials):
                self.run_ai_review(auto=True)

    def _apply_stored_bboxes(self, result: "ScanResult") -> None:
        """Load saved bbox corrections from DB and apply to matching rows.
        This pre-fills the correct location for rows the user already corrected
        on a previous scan of the same file.
        """
        stored = self.db.load_bbox_corrections(result.file_hash)
        if not stored:
            return
        for entry in stored:
            desc_lower = (entry.get("description") or "").lower().strip()
            src_lower  = (entry.get("source_line") or "").lower().strip()
            bbox       = entry.get("source_bbox", "")
            kind       = entry.get("source_bbox_kind", "image")
            page       = int(entry.get("source_page") or 0)
            if not bbox:
                continue
            for row in result.materials:
                if row.source_bbox:   # already has a box, don't overwrite
                    continue
                r_desc = (row.description or "").lower().strip()
                r_src  = (row.source_line  or "").lower().strip()
                if (desc_lower and desc_lower == r_desc) or (src_lower and src_lower == r_src):
                    row.source_bbox      = bbox
                    row.source_bbox_kind = kind
                    if page:
                        row.source_page  = page
                    break   # apply to first match only

    def _on_item_reviewed(self, row: "MaterialLine") -> None:
        """Called from the UI pump after each Anthropic-reviewed item during a scan.
        Refreshes the tree row and updates the source highlight in real time.
        """
        # Find the tree item for this row by scan_index or object identity
        for item, r in self.tree_item_map.items():
            if r is row or (r.scan_index == row.scan_index and r.source_line == row.source_line):
                self.refresh_tree_row(item)
                break
        # Re-paint all source highlights so the new colour shows immediately
        if self.current_result:
            self.apply_all_source_highlights()

    def show_result(self, result: ScanResult) -> None:
        self.current_text_source = "merged"
        self.text_box.delete("1.0", "end")
        self.text_box.insert("1.0", result.merged_text or "")
        self.populate_tree(result.materials)
        first = self.tree.get_children()
        if first:
            self.tree.selection_set(first[0])
            self.tree.focus(first[0])
        self.render_preview(page=1, highlight=self.selected_material_line())
        self.conf_var.set(f"Confidence: {result.confidence}% | Lines: {len(result.materials)} | {Path(result.file_path).name}")
        self.update_review_status_bar()
        self.apply_all_source_highlights()

    def populate_tree(self, rows: List[MaterialLine]) -> None:
        self.tree_item_map.clear()
        for item in self.tree.get_children():
            self.tree.delete(item)
        rows = order_materials_for_review(rows)
        for r in rows:
            tag = self.row_tag(r)
            status = r.review_status if r.review_status in {"pending", "confirmed", "corrected", "rejected"} else "pending"
            ai_text = self.ai_label(r)
            item = self.tree.insert("", "end", values=(r.scan_index, status, ai_text, r.qty, r.unit, r.description, r.part_number, r.manufacturer, r.confidence, r.notes), tags=(tag,))
            self.tree_item_map[item] = r

    def on_file_select(self, event=None) -> None:
        # In MaINbox review mode, preserve edits from the currently visible file
        # before switching to another attachment.
        try:
            if self.current_result:
                self.current_result.materials = self.tree_rows()
        except Exception:
            pass
        idxs = list(self.file_list.curselection())
        if not idxs:
            return
        p = self.files[idxs[0]]
        for r in reversed(self.results):
            if Path(r.file_path) == p:
                self.current_result = r
                self.show_result(r)
                return

    def show_text_source(self, source: str) -> None:
        if not self.current_result:
            return
        self.current_text_source = source
        text = {
            "native": self.current_result.native_text,
            "ocr": self.current_result.ocr_text,
            "merged": self.current_result.merged_text,
        }.get(source, self.current_result.merged_text)
        self.text_box.delete("1.0", "end")
        self.text_box.insert("1.0", text or "")
        self.highlight_selected_evidence()
        if source == "merged":
            self.apply_all_source_highlights()  # v2.0
        self.evidence_tabs.select(0)

    def remove_selected_files(self) -> None:
        for idx in reversed(self.file_list.curselection()):
            self.file_list.delete(idx)
            try:
                del self.files[idx]
            except Exception:
                pass

    def clear_files(self) -> None:
        self.files.clear()
        self.file_list.delete(0, "end")

    def open_settings(self) -> None:
        SettingsWindow(self.root, self.settings, self.save_settings)

    def save_settings(self, settings: ScanSettings) -> None:
        self.settings = settings
        self.db.save_settings(settings)
        self.status_var.set("Settings saved.")

    def ai_label(self, row: MaterialLine) -> str:
        st = (row.ai_status or "").lower()
        if st == "ok":
            return "✓ OK"
        if st == "error":
            return "✕ Error"
        if st == "warning":
            return "⚠ Review"
        if st == "suggestion":
            return "↻ Fix?"
        return ""

    def row_tag(self, row: MaterialLine) -> str:
        st = (row.ai_status or "").lower()
        if st == "error":
            return "ai_error"
        if st == "warning":
            return "ai_warning"
        if st == "suggestion":
            return "ai_suggestion"
        status = row.review_status if row.review_status in {"pending", "confirmed", "corrected", "rejected"} else "pending"
        if st == "ok" and status == "pending":
            return "ai_ok"
        return status

    def refresh_tree_row(self, item: str) -> None:
        row = self.tree_item_map.get(item)
        if not row:
            return
        status = row.review_status if row.review_status in {"pending", "confirmed", "corrected", "rejected"} else "pending"
        self.tree.item(item, values=(row.scan_index, status, self.ai_label(row), row.qty, row.unit, row.description, row.part_number, row.manufacturer, row.confidence, row.notes), tags=(self.row_tag(row),))

    def run_ai_review(self, auto: bool = False) -> None:
        if not self.current_result:
            if not auto:
                messagebox.showinfo("SmartScan", "Scan a file first, then run AI review.")
            return
        rows = self.tree_rows()
        if not rows:
            if not auto:
                messagebox.showinfo("SmartScan", "There are no extracted material rows to review yet.")
            return

        if not auto:
            self.status_var.set("Running AI/rule review...")
            self.root.update_idletasks()

        # v2.6: re-run the source-coverage watchdog on demand too, so cached
        # results or manually edited rows can still be checked against the OCR/source ledger.
        try:
            watchdog = ExtractionWatchdog(MaterialParser(self.db, self.settings), log=self.log)
            rows, audit_warnings = watchdog.audit(rows, self.current_result.merged_text)
            for w in audit_warnings:
                self.log(w)
        except Exception as exc:
            self.log(f"AI Watchdog audit skipped: {exc}")

        reviewer = AIReviewer(self.settings, log=self.log)
        rows = reviewer.review(rows, self.current_result.merged_text)
        hidden_count = 0
        if getattr(self.settings, "hide_definite_noise_rows", True):
            before_count = len(rows)
            rows = [r for r in rows if not is_definite_noise_row(r)]
            hidden_count = before_count - len(rows)
        self.current_result.materials = rows
        self.populate_tree(rows)
        self.update_review_status_bar()
        self.apply_all_source_highlights()  # v2.0 re-paint green/yellow/red

        errors = sum(1 for r in rows if (r.ai_status or "").lower() == "error")
        warnings = sum(1 for r in rows if (r.ai_status or "").lower() == "warning")
        suggestions = sum(1 for r in rows if (r.ai_status or "").lower() == "suggestion")
        ok = sum(1 for r in rows if (r.ai_status or "").lower() == "ok")
        issues = errors + warnings + suggestions

        if issues:
            self.select_next_ai_issue(show_message=False)
        else:
            children = self.tree.get_children()
            if children:
                self.tree.selection_set(children[0])
                self.tree.focus(children[0])
                self.tree.see(children[0])

        review_type = "Rule review"
        if self.settings.enable_ollama_ai_review:
            review_type += " + Ollama notes"
        noise_msg = f" Hidden noise rows: {hidden_count}." if hidden_count else ""
        msg = f"{review_type} complete: {issues} issue(s) found ({errors} error, {warnings} warning, {suggestions} suggested fix, {ok} OK).{noise_msg}"
        self.status_var.set(msg)
        self.log(msg)

        if not auto:
            if issues:
                messagebox.showinfo(
                    "AI Review Complete",
                    f"{msg}\n\nThe first issue was selected automatically. Red rows are likely errors, yellow rows need review, and blue rows have a suggested fix. Use Accept AI Fix, Keep Original, Confirm, Reject, or Next Issue."
                )
            else:
                messagebox.showinfo(
                    "AI Review Complete",
                    f"{msg}\n\nNothing was flagged. Rows should show ✓ OK in the AI Review column. If you want a slower model-based review, enable Ollama AI Review in Settings and make sure Ollama is running."
                )

    def accept_ai_fix_selected(self) -> None:
        sel = self.tree.selection()
        if not sel:
            messagebox.showinfo("SmartScan", "Select a row with an AI suggestion first.")
            return
        changed = 0
        for item in sel:
            row = self.tree_item_map.get(item)
            if not row:
                continue
            if row.ai_suggest_qty:
                row.qty = row.ai_suggest_qty
            if row.ai_suggest_unit:
                row.unit = row.ai_suggest_unit
            if row.ai_suggest_description:
                row.description = row.ai_suggest_description
            if row.ai_suggest_part_number:
                row.part_number = row.ai_suggest_part_number
            if row.ai_suggest_manufacturer:
                row.manufacturer = row.ai_suggest_manufacturer
            if any([row.ai_suggest_qty, row.ai_suggest_unit, row.ai_suggest_description, row.ai_suggest_part_number, row.ai_suggest_manufacturer]):
                row.review_status = "corrected"
                row.user_decision = "accepted_ai"
                row.ai_status = "ok"
                row.ai_message = "User accepted AI/rule suggestion"
                row.highlight_color = "green"  # v2.0
                row.notes = (row.notes + " | Accepted AI fix").strip(" |")
                changed += 1
            self.refresh_tree_row(item)
        self.update_review_status_bar()
        self.apply_all_source_highlights()  # v2.0 re-paint
        self.status_var.set(f"Accepted AI fix for {changed} row(s). Save Corrections to teach the scanner.")
        self.select_next_ai_issue(show_message=False)

    def keep_original_selected(self) -> None:
        sel = self.tree.selection()
        if not sel:
            messagebox.showinfo("SmartScan", "Select one or more rows first.")
            return
        for item in sel:
            row = self.tree_item_map.get(item)
            if not row:
                continue
            row.review_status = "confirmed"
            row.user_decision = "kept_original"
            if row.ai_status in {"error", "warning", "suggestion"}:
                row.notes = (row.notes + " | User kept original despite AI warning").strip(" |")
            row.ai_status = "ok"
            row.ai_message = "User confirmed original"
            row.highlight_color = "green"  # v2.0
            self.refresh_tree_row(item)
        self.update_review_status_bar()
        self.apply_all_source_highlights()  # v2.0 re-paint
        self.status_var.set("Original row(s) kept and confirmed. Save Corrections to teach the scanner.")
        self.select_next_ai_issue(show_message=False)

    def select_next_ai_issue(self, show_message: bool = True) -> None:
        children = list(self.tree.get_children())
        if not children:
            return
        current = self.tree.selection()[0] if self.tree.selection() else None
        start = children.index(current) + 1 if current in children else 0
        ordered = children[start:] + children[:start]
        for item in ordered:
            row = self.tree_item_map.get(item)
            if row and (row.ai_status or "").lower() in {"error", "warning", "suggestion"} and row.review_status != "rejected":
                self.tree.selection_set(item)
                self.tree.focus(item)
                self.tree.see(item)
                self.on_material_select()
                return
        if show_message:
            messagebox.showinfo("SmartScan", "No AI review issues left.")

    def tree_rows(self) -> List[MaterialLine]:
        rows: List[MaterialLine] = []
        for item in self.tree.get_children():
            vals = self.tree.item(item, "values")
            vals = list(vals) + [""] * 10
            try:
                scan_index = int(vals[0])
            except Exception:
                scan_index = len(rows) + 1
            status = str(vals[1] or "pending").lower()
            try:
                conf = int(vals[8])
            except Exception:
                conf = 0
            existing = self.tree_item_map.get(item, MaterialLine())
            existing.scan_index = scan_index
            existing.review_status = status if status in {"pending", "confirmed", "corrected", "rejected"} else "pending"
            existing.qty = str(vals[3])
            existing.unit = str(vals[4])
            existing.description = str(vals[5])
            existing.part_number = str(vals[6])
            existing.manufacturer = str(vals[7])
            existing.confidence = conf
            existing.notes = str(vals[9])
            if not existing.source_line:
                existing.source_line = existing.description
            rows.append(existing)
        return rows

    def _set_review_status_for_items(self, items: Iterable[str], status: str) -> None:
        if status not in {"pending", "confirmed", "corrected", "rejected"}:
            return
        for item in items:
            row = self.tree_item_map.get(item)
            if not row:
                continue
            row.review_status = status
            vals = list(self.tree.item(item, "values"))
            while len(vals) < 10:
                vals.append("")
            vals[1] = status
            self.tree.item(item, values=vals, tags=(self.row_tag(row),))
        self.update_review_status_bar()

    def confirm_selected_rows(self) -> None:
        sel = self.tree.selection()
        if not sel:
            messagebox.showinfo("SmartScan", "Select one or more extracted material rows first.")
            return
        self._set_review_status_for_items(sel, "confirmed")
        for item in sel:
            row = self.tree_item_map.get(item)
            if row:
                row.user_decision = "confirmed"
                if not row.ai_status:
                    row.ai_status = "ok"
                self.refresh_tree_row(item)
        self.status_var.set("Selected row(s) confirmed. Save Corrections to teach the scanner.")
        self.select_next_pending()

    def reject_selected_rows(self) -> None:
        sel = self.tree.selection()
        if not sel:
            messagebox.showinfo("SmartScan", "Select one or more extracted material rows first.")
            return
        self._set_review_status_for_items(sel, "rejected")
        for item in sel:
            row = self.tree_item_map.get(item)
            if row:
                row.user_decision = "rejected"
                self.refresh_tree_row(item)
        self.status_var.set("Selected row(s) rejected. Save Corrections to teach the scanner what to ignore.")
        self.select_next_pending()

    def select_next_pending(self) -> None:
        children = list(self.tree.get_children())
        if not children:
            return
        current = self.tree.selection()[0] if self.tree.selection() else None
        start = children.index(current) + 1 if current in children else 0
        ordered = children[start:] + children[:start]
        for item in ordered:
            row = self.tree_item_map.get(item)
            if row and row.review_status == "pending":
                self.tree.selection_set(item)
                self.tree.focus(item)
                self.tree.see(item)
                self.on_material_select()
                return
        self.status_var.set("All visible rows are reviewed. Save Corrections to store the learning.")

    def update_review_status_bar(self) -> None:
        rows = self.tree_rows()
        total = len(rows)
        pending = sum(1 for r in rows if r.review_status == "pending")
        confirmed = sum(1 for r in rows if r.review_status == "confirmed")
        corrected = sum(1 for r in rows if r.review_status == "corrected")
        rejected = sum(1 for r in rows if r.review_status == "rejected")
        ai_errors = sum(1 for r in rows if (r.ai_status or "").lower() == "error")
        ai_warn = sum(1 for r in rows if (r.ai_status or "").lower() in {"warning", "suggestion"})
        audit_bits = ""
        try:
            audit = self._current_final_audit(mutate_rows=False) if self.current_result else None
            if audit:
                audit_bits = f" | Audit: {audit.get('score', 0)}% / {audit.get('issues', 0)} open"
        except Exception:
            audit_bits = ""
        self.conf_var.set(f"Review: {confirmed} confirmed | {corrected} corrected | {rejected} rejected | {pending} pending | AI issues: {ai_errors} errors/{ai_warn} warnings | Lines: {total}{audit_bits}")

    def _current_final_audit(self, mutate_rows: bool = True) -> Optional[Dict[str, Any]]:
        if not self.current_result:
            return None
        parser = MaterialParser(self.db, self.settings)
        return SourceCoverageAuditor(parser, log=self.log).audit(
            self.tree_rows(), self.current_result.merged_text or "", mutate_rows=mutate_rows
        )

    def _format_final_audit_report(self, audit: Dict[str, Any]) -> str:
        lines = [
            "SmartScan Final Audit",
            "",
            f"Audit score: {audit.get('score', 0)}%",
            f"Source material-looking lines: {audit.get('ledger_count', 0)}",
            f"Extracted material rows: {audit.get('row_count', 0)}",
            f"Source lines accounted for: {audit.get('matched_count', 0)} / {audit.get('ledger_count', 0)}",
            f"Confirmed/corrected rows: {audit.get('confirmed_count', 0)}",
            f"Pending rows: {audit.get('pending_count', 0)}",
            f"Rejected rows: {audit.get('rejected_count', 0)}",
            "",
            "Open items:",
            f"- Possible missed source lines: {audit.get('missing_count', 0)}",
            f"- Qty/unit/description mismatches: {audit.get('mismatch_count', 0)}",
            f"- Possible duplicate rows: {audit.get('duplicate_count', 0)}",
        ]
        missing = audit.get("missing") or []
        if missing:
            lines.append("")
            lines.append("Possible missed source lines:")
            for cand in missing[:20]:
                lines.append(f"- {cand.source_line}")
            if len(missing) > 20:
                lines.append(f"- ... plus {len(missing) - 20} more")
        mismatches = audit.get("mismatches") or []
        if mismatches:
            lines.append("")
            lines.append("Rows with possible field mismatches:")
            for idx, msg in mismatches[:20]:
                lines.append(f"- Row {idx + 1}: {msg}")
            if len(mismatches) > 20:
                lines.append(f"- ... plus {len(mismatches) - 20} more")
        duplicates = audit.get("duplicates") or []
        if duplicates:
            lines.append("")
            lines.append("Possible duplicate rows:")
            for idx in duplicates[:20]:
                lines.append(f"- Row {idx + 1}")
        return "\n".join(lines)

    def show_final_audit(self) -> None:
        if not self.current_result:
            messagebox.showinfo("SmartScan", "No scan result selected.")
            return
        audit = self._current_final_audit(mutate_rows=True)
        if not audit:
            return
        self.populate_tree(self.tree_rows())
        self.update_review_status_bar()
        report = self._format_final_audit_report(audit)
        win = tk.Toplevel(self.root)
        win.title("SmartScan Final Audit")
        win.geometry("720x560")
        ttk.Label(win, text="Final Audit", font=("Segoe UI", 14, "bold")).pack(anchor="w", padx=12, pady=(10, 4))
        txt = tk.Text(win, wrap="word")
        txt.pack(fill="both", expand=True, padx=12, pady=(0, 8))
        txt.insert("1.0", report)
        txt.configure(state="disabled")
        btns = ttk.Frame(win)
        btns.pack(fill="x", padx=12, pady=(0, 12))
        ttk.Button(btns, text="Review Missed Lines", command=lambda: (win.destroy(), self.show_possible_missed_lines())).pack(side="left")
        ttk.Button(btns, text="Next Issue", command=self.select_next_ai_issue).pack(side="left", padx=6)
        ttk.Button(btns, text="Close", command=win.destroy).pack(side="right")
        self.status_var.set(f"Final Audit: {audit.get('score', 0)}% | {audit.get('issues', 0)} open item(s).")

    def show_possible_missed_lines(self) -> None:
        if not self.current_result:
            messagebox.showinfo("SmartScan", "No scan result selected.")
            return
        audit = self._current_final_audit(mutate_rows=False)
        if not audit:
            return
        missing: List[SourceMaterialCandidate] = list(audit.get("missing") or [])
        if not missing:
            messagebox.showinfo("SmartScan", "No possible missed source lines were found.")
            return
        win = tk.Toplevel(self.root)
        win.title("Possible Missed Lines")
        win.geometry("860x460")
        ttk.Label(win, text="Possible Missed Source Lines", font=("Segoe UI", 13, "bold")).pack(anchor="w", padx=10, pady=(10, 2))
        ttk.Label(win, text="Select lines that should become material rows. Use Ignore/Close for non-material text.").pack(anchor="w", padx=10, pady=(0, 6))
        lb = tk.Listbox(win, selectmode="extended")
        lb.pack(fill="both", expand=True, padx=10, pady=(0, 8))
        for cand in missing:
            parsed = cand.parsed
            parsed_text = f"  → {parsed.qty} {parsed.unit} {parsed.description}" if parsed else ""
            lb.insert("end", f"{cand.index}: {cand.source_line}{parsed_text}")

        def add_selected() -> None:
            sel = list(lb.curselection())
            if not sel:
                messagebox.showinfo("SmartScan", "Select one or more missed lines first.")
                return
            added = 0
            for pos in sel:
                cand = missing[pos]
                if not cand.parsed:
                    continue
                row = cand.parsed
                row.source_ledger_index = cand.index
                row.source_start = cand.source_start
                row.source_end = cand.source_end
                row.source_page = cand.source_page
                row.review_status = "pending"
                row.ai_status = "warning"
                row.ai_message = "Final Audit: added from possible missed source line; please confirm."
                row.notes = append_note(row.notes, "Final Audit: added from missed-line queue")
                row.highlight_color = "yellow"
                row.snapshot_original()
                if self.current_result:
                    self.current_result.materials.append(row)
                added += 1
            if added:
                self.populate_tree(order_materials_for_review(self.current_result.materials if self.current_result else self.tree_rows()))
                self.update_review_status_bar()
                self.status_var.set(f"Added {added} missed line(s). Please confirm them before saving.")
                win.destroy()

        btns = ttk.Frame(win)
        btns.pack(fill="x", padx=10, pady=(0, 10))
        ttk.Button(btns, text="Add Selected as Row(s)", command=add_selected).pack(side="left")
        ttk.Button(btns, text="Close / Ignore", command=win.destroy).pack(side="right")

    def save_corrections(self) -> None:
        if not self.current_result:
            messagebox.showinfo("SmartScan", "No scan result selected.")
            return
        rows = self.tree_rows()
        audit = self._current_final_audit(mutate_rows=True)
        if audit and audit.get("issues", 0) > 0:
            proceed = messagebox.askyesno(
                "Final Audit Has Open Items",
                self._format_final_audit_report(audit) + "\n\nSave anyway?",
            )
            if not proceed:
                self.populate_tree(rows)
                self.update_review_status_bar()
                return
        self.db.save_corrections(self.current_result.file_hash, rows)
        self.db.save_correction_patterns(self.current_result.file_hash, rows)
        # v2.2: also persist any user-drawn bbox corrections
        self.db.save_bbox_corrections(self.current_result.file_hash, rows)
        self.current_result.materials = rows
        self.current_result.confidence = MaterialParser(self.db).overall_confidence(rows, self.current_result.merged_text)
        self.db.save_scan(self.current_result)
        self.update_review_status_bar()
        bbox_count = sum(1 for r in rows if r.source_bbox)
        self.status_var.set(f"Corrections saved ({bbox_count} bbox location(s) stored for future learning).")
        self.log(f"Saved {len(rows)} reviewed/corrected material line(s), {bbox_count} with bbox.")

    def add_material_row(self) -> None:
        """Insert a blank row immediately after the currently-selected item (or at end)."""
        sel = self.tree.selection()
        after_item = sel[-1] if sel else ""   # insert after last selected; "" means end
        # Determine the new scan_index: one more than the selected row's index
        if after_item and after_item in self.tree_item_map:
            ref_row = self.tree_item_map[after_item]
            next_idx = (ref_row.scan_index or 0) + 1
        else:
            next_idx = len(self.tree.get_children()) + 1
        row = MaterialLine(
            confidence=80, notes="manual", source_line="",
            scan_index=next_idx, review_status="corrected",
            ai_status="ok", ai_message="Manual row",
            highlight_color="green",
        )
        # tkinter Treeview.insert: index="" means end; use after_item when we have one
        if after_item:
            # Find the positional index so we can insert right after it
            children = self.tree.get_children()
            try:
                pos = list(children).index(after_item) + 1
            except ValueError:
                pos = "end"
        else:
            pos = "end"
        item = self.tree.insert("", pos, values=(next_idx, "corrected", "✓ OK", "", "", "", "", "", 80, "manual"), tags=("corrected",))
        self.tree_item_map[item] = row
        self.tree.selection_set(item)
        self.tree.see(item)
        # Renumber scan_index for all rows below the insertion point so order stays consistent
        self._renumber_tree_from(item)

    def delete_material_rows(self) -> None:
        for item in self.tree.selection():
            self.tree_item_map.pop(item, None)
            self.tree.delete(item)
        self._renumber_tree_from(None)  # keep scan_index sequential after deletion

    def _renumber_tree_from(self, start_item: Optional[str]) -> None:
        """Re-assign sequential scan_index values to every row after start_item.
        If start_item is None, renumber all rows from 1.
        """
        children = list(self.tree.get_children())
        if start_item is None:
            start_pos = 0
            next_num = 1
        else:
            try:
                start_pos = children.index(start_item)
                ref = self.tree_item_map.get(start_item)
                next_num = (ref.scan_index if ref else start_pos) + 1
            except ValueError:
                start_pos = 0
                next_num = 1
        for item in children[start_pos + (0 if start_item is None else 1):]:
            row = self.tree_item_map.get(item)
            if row:
                row.scan_index = next_num
            vals = list(self.tree.item(item, "values"))
            if vals:
                vals[0] = str(next_num)
                self.tree.item(item, values=vals)
            next_num += 1

    def edit_tree_cell(self, event) -> None:
        item = self.tree.identify_row(event.y)
        col = self.tree.identify_column(event.x)
        if not item or not col:
            return
        col_idx = int(col.replace("#", "")) - 1
        bbox = self.tree.bbox(item, col)
        if not bbox:
            return
        x, y, w, h = bbox
        values = list(self.tree.item(item, "values"))
        old = values[col_idx] if col_idx < len(values) else ""
        entry = ttk.Entry(self.tree)
        entry.insert(0, old)
        entry.place(x=x, y=y, width=w, height=h)
        entry.focus_set()

        def save_edit(event=None):
            new = entry.get()
            values[col_idx] = new
            self.tree.item(item, values=values)
            row = self.tree_item_map.get(item)
            if row is not None:
                attrs = ["scan_index", "review_status", "ai_status", "qty", "unit", "description", "part_number", "manufacturer", "confidence", "notes"]
                if 0 <= col_idx < len(attrs):
                    attr = attrs[col_idx]
                    try:
                        setattr(row, attr, int(new) if attr in {"confidence", "scan_index"} else str(new))
                    except Exception:
                        setattr(row, attr, new)
                    if attr in {"qty", "unit", "description", "part_number", "manufacturer", "notes"}:
                        row.snapshot_original()  # v2.0: freeze original before overwriting
                        row.user_decision = "manual"
                        if row.review_status == "pending":
                            row.review_status = "corrected"
                            values[1] = "corrected"
                        # Smart-review the edited row immediately so the highlight color
                        # changes as the user fixes it.
                        if getattr(self.settings, "enable_rule_review", True):
                            AIReviewer(self.settings, log=self.log).review([row], self.current_result.merged_text if self.current_result else "")
                            values[2] = self.ai_label(row)
                        # Update highlight colour after re-review
                        st = (row.ai_status or "").lower()
                        row.highlight_color = "green" if st == "ok" else ("red" if st == "error" else "yellow")
                        self.tree.item(item, values=values, tags=(self.row_tag(row),))
                    elif attr == "review_status":
                        self.tree.item(item, tags=(self.row_tag(row),))
            entry.destroy()
            self.highlight_selected_evidence()
            if self.current_result:
                self.render_preview(page=self.preview_page or (row.source_page if row else 1), highlight=row)
            self.update_review_status_bar()

        entry.bind("<Return>", save_edit)
        entry.bind("<FocusOut>", save_edit)

    def selected_material_line(self) -> Optional[MaterialLine]:
        sel = self.tree.selection()
        if not sel:
            return None
        return self.tree_item_map.get(sel[0])

    def on_material_select(self, event=None) -> None:
        row = self.selected_material_line()
        if not row:
            return

        # Update span label (v2.1)
        span = row.source_end - row.source_start if row.source_start >= 0 else 0
        self.hl_span_var.set(f"span: {span} chars  [{row.source_start}:{row.source_end}]")

        # v2.2: show current bbox coords in the preview toolbar
        if row.source_bbox:
            self.bbox_coords_var.set(f"bbox: {row.source_bbox}  (page {row.source_page or '?'})")
        else:
            self.bbox_coords_var.set("")

        # v2.5.6: keep selection feedback tied to the active evidence tab.
        # - Extracted/Merged Text tab: show a light-yellow source row guide.
        # - File Preview tab: move through the preview highlights instead.
        active_tab = 0
        try:
            active_tab = self.evidence_tabs.index(self.evidence_tabs.select())
        except Exception:
            active_tab = 0

        if active_tab == 1:
            page = row.source_page or self.infer_page_for_row(row) or self.preview_page or 1
            same_page = (page == self.preview_page)
            self.render_preview(page=page, highlight=row, keep_scroll=same_page)
            self.status_var.set("Highlighted preview location for selected material row.")
        else:
            self.highlight_selected_evidence()
            self.status_var.set("Highlighted source text row for selected material row.")

    # ── Highlight region editing (v2.1) ──────────────────────────────────────────

    def _tk_idx_to_char(self, tk_idx: str) -> int:
        """Convert a Tk Text widget index like '3.12' to an absolute character offset.
        Uses the widget's own count() method — the only reliable approach.
        Returns -1 on error.
        """
        try:
            result = self.text_box.count("1.0", tk_idx, "chars")
            # count() returns a tuple on some Tk versions, int on others
            if isinstance(result, (list, tuple)):
                return int(result[0])
            return int(result)
        except Exception:
            return -1

    def _use_text_selection_as_highlight(self) -> None:
        """Snap the selected row's source_start/end to the text currently
        selected (blue highlight) in the merged-text viewer.

        Must be in merged view — source_start/end are offsets into merged_text.
        """
        row = self.selected_material_line()
        if not row:
            messagebox.showinfo("SmartScan",
                "Select a material row in the table first, then select text "
                "in the source view and click Draw Highlight.")
            return

        # Guard: offsets only make sense when showing the merged text
        if self.current_text_source != "merged":
            messagebox.showinfo("SmartScan",
                "Switch to the Merged Text view first (click Show Merged), "
                "then select the text you want to highlight and click Draw Highlight.")
            return

        try:
            sel_start_idx = self.text_box.index("sel.first")
            sel_end_idx   = self.text_box.index("sel.last")
        except tk.TclError:
            messagebox.showinfo("SmartScan",
                "No text selected.  Click and drag to select the correct text "
                "in the source view, then click Draw Highlight.")
            return

        start_char = self._tk_idx_to_char(sel_start_idx)
        end_char   = self._tk_idx_to_char(sel_end_idx)

        if start_char < 0 or end_char <= start_char:
            messagebox.showinfo("SmartScan", "Could not read selection — try again.")
            return

        # Verify offsets are inside merged_text
        merged = (self.current_result.merged_text or "") if self.current_result else ""
        if end_char > len(merged):
            end_char = len(merged)
        if start_char >= end_char:
            return

        row.source_start = start_char
        row.source_end   = end_char
        row.source_line  = merged[start_char:end_char]
        self._apply_confirmed_source_text_to_material_row(row, row.source_line)
        self._refresh_current_tree_row(row)

        self._refresh_highlight_display(row)

        # Re-run rule review so the highlight colour updates immediately
        if getattr(self.settings, "enable_rule_review", True) and self.current_result:
            AIReviewer(self.settings, log=self.log).review([row], merged)
            st = (row.ai_status or "").lower()
            row.highlight_color = "green" if st == "ok" else ("red" if st == "error" else "yellow")

        self.apply_all_source_highlights()

        # v2.4: auto-cross-reference — switch to preview and attempt to locate
        # the selected text on the scanned page so the user can confirm/adjust.
        self._xref_text_to_preview(row)

    def _reset_highlight(self) -> None:
        """Revert the selected row's highlight span back to its original parsed position."""
        row = self.selected_material_line()
        if not row:
            return
        # Re-annotate just this row using EvidenceMapper
        if self.current_result:
            mapper = EvidenceMapper(Path(self.current_result.file_path), self.current_result.merged_text)
            mapper.annotate([row])
        self._refresh_highlight_display(row)

    def _refresh_highlight_display(self, row: MaterialLine) -> None:
        """Update span label, repaint all highlights, and scroll to the active span."""
        span = row.source_end - row.source_start if row.source_start >= 0 else 0
        self.hl_span_var.set(f"span: {span} chars  [{row.source_start}:{row.source_end}]")
        self.apply_all_source_highlights()
        self.highlight_selected_evidence()

    def apply_all_source_highlights(self) -> None:
        """Paint green/yellow/red behind every recognised material span in the merged text view."""
        if self.current_text_source != "merged":
            return
        for tag in ("hl_green", "hl_yellow", "hl_red"):
            self.text_box.tag_remove(tag, "1.0", "end")
        if not self.current_result:
            return
        # Use the widget's own character count as the authoritative length —
        # never trust len(merged_text) because the widget may hold a slightly
        # different string after tag operations or encoding differences.
        widget_len = self._tk_idx_to_char("end-1c")
        if widget_len <= 0:
            return
        for row in self.current_result.materials:
            start = row.source_start
            end   = row.source_end
            if start < 0 or end <= start or end > widget_len:
                continue
            color = row.highlight_color or ""
            if not color:
                st = (row.ai_status or "").lower()
                color = ("green" if st == "ok"
                         else "red" if st == "error"
                         else "yellow" if st in {"warning", "suggestion"}
                         else "")
            hl_tag = {"green": "hl_green", "yellow": "hl_yellow", "red": "hl_red"}.get(color)
            if hl_tag:
                self.text_box.tag_add(hl_tag, f"1.0+{start}c", f"1.0+{end}c")

    def highlight_selected_evidence(self) -> None:
        for tag in ("evidence_row_select", "evidence_current", "evidence_soft", "evidence_ok",
                    "evidence_warning", "evidence_error", "evidence_suggestion",
                    "evidence_confirm"):
            self.text_box.tag_remove(tag, "1.0", "end")
        row = self.selected_material_line()
        if not row:
            return
        tag = self.text_tag_for_row(row)
        start = row.source_start
        end   = row.source_end

        # Offsets are valid only when the text widget is showing merged text
        # AND the offsets sit inside the actual widget content.
        if self.current_text_source == "merged" and start >= 0 and end > start:
            # The text widget's character count is the authoritative length.
            widget_len = self._tk_idx_to_char("end-1c")
            if 0 <= start < end <= widget_len:
                idx = f"1.0+{start}c"
                end_idx = f"1.0+{end}c"
                self._highlight_full_source_line(idx)
                self.text_box.tag_add(tag, idx, end_idx)
                self.text_box.see(idx)
                return

        # Fallback: search by text for native/ocr views or when offsets are stale.
        needle = self._needle_for_text_view_row(row)
        if needle:
            idx = self.text_box.search(needle, "1.0", nocase=True, stopindex="end")
            if idx:
                self._highlight_full_source_line(idx)
                self.text_box.tag_add(tag, idx, f"{idx}+{len(needle)}c")
                self.text_box.see(idx)

    def _highlight_full_source_line(self, idx: str) -> None:
        """Add the bright yellow full-pane row band on the selected source-text line.

        v2.5.7 note:
        Tk Text tag backgrounds only paint the tagged characters unless the line's
        newline character is included.  By tagging through ``lineend + 1c`` when
        possible, the selected row guide fills the rest of the visible text pane,
        making the OCR/source line much easier to match to the blue material row.
        """
        try:
            line_start = self.text_box.index(f"{idx} linestart")
            line_end = self.text_box.index(f"{idx} lineend")
            full_line_end = line_end
            # Include the newline character when it exists; this makes Tk paint
            # the background across the full width of the text pane.  For the
            # final line, lineend may already be at end-1c, so fall back safely.
            try:
                next_char = self.text_box.index(f"{line_end}+1c")
                if self.text_box.compare(next_char, "<=", "end"):
                    full_line_end = next_char
            except Exception:
                full_line_end = line_end
            if self.text_box.compare(full_line_end, ">", line_start):
                self.text_box.tag_add("evidence_row_select", line_start, full_line_end)
                self.text_box.tag_raise("evidence_row_select")
        except Exception:
            pass

    def _needle_for_text_view_row(self, row: MaterialLine) -> str:
        """Best single-line string to locate the row in OCR/native/merged text views."""
        for value in (row.source_line, row.description, row.part_number):
            text = (value or "").strip()
            if not text:
                continue
            for part in re.split(r"\r?\n", text):
                part = part.strip()
                if len(part) >= 2:
                    return part
        return ""

    def _on_text_box_click_select_row(self, event=None) -> None:
        """Let a user click a source-text line and select the matching material row.

        Skips normal drag selections, so it does not interfere with the existing
        select-text-then-Draw-Highlight correction workflow.
        """
        if not self.current_result:
            return
        try:
            if self.text_box.tag_ranges("sel"):
                return
        except Exception:
            pass

        try:
            click_idx = self.text_box.index(f"@{event.x},{event.y}" if event else "insert")
        except Exception:
            return

        row = None
        if self.current_text_source == "merged":
            pos = self._tk_idx_to_char(click_idx)
            line_start = self._tk_idx_to_char(self.text_box.index(f"{click_idx} linestart"))
            line_end = self._tk_idx_to_char(self.text_box.index(f"{click_idx} lineend"))
            for candidate in self.current_result.materials:
                start = candidate.source_start
                end = candidate.source_end
                if start >= 0 and end > start and (start <= pos <= end or (line_start <= start <= line_end)):
                    row = candidate
                    break

        if row is None:
            try:
                line_text = self.text_box.get(f"{click_idx} linestart", f"{click_idx} lineend").strip()
            except Exception:
                line_text = ""
            norm_line = normalize_space(line_text).lower()
            if norm_line:
                for candidate in self.current_result.materials:
                    for value in (candidate.source_line, candidate.description, candidate.part_number):
                        needle = normalize_space(value or "").lower()
                        if needle and (needle in norm_line or norm_line in needle):
                            row = candidate
                            break
                    if row is not None:
                        break

        if row is None:
            return

        for item, candidate in self.tree_item_map.items():
            if candidate is row:
                self.tree.selection_set(item)
                self.tree.focus(item)
                self.tree.see(item)
                self.highlight_selected_evidence()
                break

    def text_tag_for_row(self, row: MaterialLine) -> str:
        st = (row.ai_status or "").lower()
        if st == "ok":
            return "evidence_ok"
        if st == "error":
            return "evidence_error"
        if st == "warning":
            return "evidence_warning"
        if st == "suggestion":
            return "evidence_suggestion"
        return "evidence_current"

    def show_preview_tab(self) -> None:
        if self.current_result:
            row = self.selected_material_line()
            page = (row.source_page if row and row.source_page else None) or self.preview_page or 1
            if row:
                page = self.infer_page_for_row(row) or page
            self.render_preview(page=page, highlight=row)
        self.evidence_tabs.select(1)

    def preview_page_count(self, path: Path) -> int:
        ext = path.suffix.lower()
        if ext == ".pdf" and fitz is not None:
            try:
                doc = fitz.open(str(path))
                n = len(doc)
                doc.close()
                return max(1, n)
            except Exception:
                return 1
        if ext in SUPPORTED_IMAGE_EXTS:
            return 1
        return 1

    def update_preview_page_controls(self) -> None:
        try:
            self.preview_page_var.set(str(max(1, int(self.preview_page or 1))))
            self.preview_total_var.set(f"/ {max(1, int(self.preview_total_pages or 1))}")
        except Exception:
            pass

    def prev_preview_page(self) -> None:
        if not self.current_result:
            return
        row = self.selected_material_line()
        self.render_preview(page=max(1, (self.preview_page or 1) - 1), highlight=row)
        self.evidence_tabs.select(1)

    def next_preview_page(self) -> None:
        if not self.current_result:
            return
        row = self.selected_material_line()
        max_page = max(1, self.preview_total_pages or 1)
        self.render_preview(page=min(max_page, (self.preview_page or 1) + 1), highlight=row)
        self.evidence_tabs.select(1)

    def go_preview_page(self) -> None:
        if not self.current_result:
            return
        try:
            page = int(self.preview_page_var.get().strip())
        except Exception:
            page = self.preview_page or 1
        max_page = max(1, self.preview_total_pages or 1)
        page = max(1, min(max_page, page))
        self.render_preview(page=page, highlight=self.selected_material_line())
        self.evidence_tabs.select(1)

    def infer_page_for_row(self, row: Optional[MaterialLine]) -> int:
        if not row or not self.current_result:
            return 0
        if row.source_page:
            return row.source_page
        if row.source_start is not None and row.source_start >= 0:
            prefix = (self.current_result.merged_text or "")[:row.source_start]
            matches = list(re.finditer(r"---\s+(?:PAGE\s+)?(\d+)\s+(?:NATIVE TEXT|POSITIONAL LINES)|---\s+OCR\s+page_(\d+)", prefix, flags=re.I))
            if matches:
                m = matches[-1]
                try:
                    return int(m.group(1) or m.group(2) or 0)
                except Exception:
                    pass
        return 0

    def _read_grid_for_preview(self, path: Path) -> List[List[str]]:
        """Re-read the spreadsheet/Word table as a cell grid for the visual preview.

        Mirrors FileExtractor's readers but is self-contained so the preview never
        depends on scan-time state. Returns [] if it cannot read the grid.
        """
        ext = path.suffix.lower()
        try:
            extractor = FileExtractor(self.settings, log=self.log)
            if ext in SUPPORTED_SPREADSHEET_EXTS:
                _nt, _tt, _warn, matrix = extractor.extract_spreadsheet(path)
                return matrix or []
            if ext in SUPPORTED_WORD_EXTS:
                _nt, _tt, _warn, matrices = extractor.extract_docx(path)
                if matrices:
                    return max(matrices, key=lambda mx: (len(mx), max((len(r) for r in mx), default=0)))
        except Exception as e:
            self.log(f"Grid preview read failed: {e}")
        return []

    def _render_table_preview_image(self, path: Path):
        """Draw a spreadsheet / Word-table grid to a PIL image and record each
        row's pixel rectangle so material rows can be highlighted on it.

        Sets:
          self._table_row_boxes  -> list of (row_index, x0, y0, x1, y1) in image px
          self._table_row_text   -> list of normalized joined row text (same order)
          self._table_grid       -> the cell grid used
        Returns the PIL image, or None on failure.
        """
        self._table_row_boxes = []
        self._table_row_text = []
        self._table_grid = []
        self._grid_row_assignment = None  # rebuilt lazily for the new grid
        grid = self._read_grid_for_preview(path)
        if not grid:
            return None

        # Normalize to a rectangular grid of strings.
        rows = [[("" if c is None else str(c)).strip() for c in r] for r in grid if any(
            ("" if c is None else str(c)).strip() for c in r)]
        if not rows:
            return None
        ncols = max(len(r) for r in rows)
        rows = [r + [""] * (ncols - len(r)) for r in rows]
        self._table_grid = rows

        # Font (fall back to PIL default if truetype unavailable).
        try:
            font = ImageFont.truetype("arial.ttf", 14)
            bold = ImageFont.truetype("arialbd.ttf", 14)
        except Exception:
            try:
                font = ImageFont.truetype("DejaVuSans.ttf", 14)
                bold = ImageFont.truetype("DejaVuSans-Bold.ttf", 14)
            except Exception:
                font = ImageFont.load_default()
                bold = font

        pad_x, pad_y = 8, 6
        line_h = 22
        min_col_w, max_col_w = 40, 460

        # Measure column widths from content (cap very wide columns).
        def text_w(s: str, f) -> int:
            try:
                box = f.getbbox(s)
                return box[2] - box[0]
            except Exception:
                try:
                    return f.getsize(s)[0]
                except Exception:
                    return 8 * len(s)

        col_w = [min_col_w] * ncols
        for r in rows:
            for ci in range(ncols):
                w = text_w(r[ci], font) + pad_x * 2
                if w > col_w[ci]:
                    col_w[ci] = min(max_col_w, w)
        total_w = sum(col_w) + 1
        total_h = line_h * len(rows) + 1

        # Cap overall size for responsiveness.
        max_w = 1700
        scale_down = 1.0
        if total_w > max_w:
            scale_down = max_w / total_w
        img_w = max(60, int(total_w * scale_down))
        img_h = max(40, int(total_h))

        img = Image.new("RGB", (int(total_w), int(total_h)), "#ffffff")
        draw = ImageDraw.Draw(img)

        # Decide which row looks like the header (reuse structured detector).
        try:
            header_idx, _cmap = detect_structured_header(rows)
        except Exception:
            header_idx = -1

        y = 0
        for ri, r in enumerate(rows):
            x = 0
            is_header = (ri == header_idx)
            row_bg = "#dfe7f3" if is_header else ("#f7f9fc" if ri % 2 == 0 else "#ffffff")
            draw.rectangle([0, y, total_w - 1, y + line_h], fill=row_bg)
            for ci in range(ncols):
                cell_text = r[ci]
                # clip text to column width
                avail = col_w[ci] - pad_x * 2
                t = cell_text
                if text_w(t, font) > avail and t:
                    while t and text_w(t + "…", font) > avail:
                        t = t[:-1]
                    if t != cell_text:
                        t = t + "…"
                draw.text((x + pad_x, y + pad_y), t, fill="#202020", font=(bold if is_header else font))
                # vertical gridline
                draw.line([x, y, x, y + line_h], fill="#e2e2e2")
                x += col_w[ci]
            # horizontal gridline
            draw.line([0, y, total_w - 1, y], fill="#d8d8d8")
            # record this row's rectangle and text (image/original coords)
            self._table_row_boxes.append((ri, 0, y, total_w - 1, y + line_h))
            self._table_row_text.append(_struct_norm_header(" ".join(c for c in r if c)))
            y += line_h
        draw.line([0, total_h - 1, total_w - 1, total_h - 1], fill="#d8d8d8")
        # outer border
        draw.rectangle([0, 0, total_w - 1, total_h - 1], outline="#b8b8b8")
        return img

    def _find_table_box_for_row(self, row: "MaterialLine"):
        """Return (x0,y0,x1,y1) in image-original px for the grid row matching this
        material line, or None.

        Uses a precomputed one-to-one assignment (built in _build_grid_row_assignment)
        so each material row maps to a DISTINCT grid row.  This prevents the failure
        the user saw where a few rows never highlighted on the preview because an
        earlier row had already 'claimed' their grid line via a loose text match.
        """
        boxes = getattr(self, "_table_row_boxes", None)
        if not boxes:
            return None
        assign = getattr(self, "_grid_row_assignment", None)
        if assign is None:
            assign = self._build_grid_row_assignment()
        ri = assign.get(id(row))
        if ri is not None:
            for (gri, x0, y0, x1, y1) in boxes:
                if gri == ri:
                    return (x0, y0, x1, y1)
        # Fallback (row not in the assignment, e.g. added after render): best single match.
        return self._match_single_grid_box(row)

    def _match_single_grid_box(self, row: "MaterialLine"):
        boxes = getattr(self, "_table_row_boxes", None)
        texts = getattr(self, "_table_row_text", None)
        if not boxes or not texts:
            return None
        targets = []
        if row.source_line:
            targets.append(_struct_norm_header(row.source_line))
        sig = _struct_norm_header(" ".join(x for x in [row.qty, row.description, row.part_number] if x))
        if sig:
            targets.append(sig)
        desc_norm = _struct_norm_header(row.description or "")
        for tgt in targets:
            if not tgt:
                continue
            for (ri, x0, y0, x1, y1), txt in zip(boxes, texts):
                if txt == tgt:
                    return (x0, y0, x1, y1)
        if desc_norm:
            for (ri, x0, y0, x1, y1), txt in zip(boxes, texts):
                if desc_norm and desc_norm in txt:
                    return (x0, y0, x1, y1)
            want = set(desc_norm.split())
            best, best_score = None, 0
            for (ri, x0, y0, x1, y1), txt in zip(boxes, texts):
                score = len(want & set(txt.split()))
                if score > best_score:
                    best_score, best = score, (x0, y0, x1, y1)
            if best_score >= max(1, len(want) // 2):
                return best
        return None

    def _grid_row_match_score(self, row: "MaterialLine", grid_text: str) -> int:
        """Score how well a material row matches a grid row's normalized text."""
        if not grid_text:
            return 0
        score = 0
        src = _struct_norm_header(row.source_line or "")
        if src and src == grid_text:
            return 1000  # exact source-line identity is the strongest signal
        sig = _struct_norm_header(" ".join(x for x in [row.qty, row.description, row.part_number] if x))
        if sig and sig == grid_text:
            return 900
        desc_norm = _struct_norm_header(row.description or "")
        if desc_norm:
            if desc_norm == grid_text:
                score = max(score, 800)
            elif desc_norm in grid_text:
                score = max(score, 500 + len(desc_norm))
            want = set(desc_norm.split())
            have = set(grid_text.split())
            shared = len(want & have)
            if shared:
                score = max(score, 100 + shared * 10)
            # quantity agreement is a good tie-breaker
            if row.qty and _struct_norm_header(str(row.qty)) in have:
                score += 5
        return score

    def _build_grid_row_assignment(self) -> Dict[int, int]:
        """Assign each material row to a unique grid row index via greedy best-score
        matching (highest-confidence pairs first), so every row gets its own line and
        no grid row is double-claimed.  Returns {id(material_row): grid_row_index}."""
        assignment: Dict[int, int] = {}
        boxes = getattr(self, "_table_row_boxes", None)
        texts = getattr(self, "_table_row_text", None)
        if not boxes or not texts:
            self._grid_row_assignment = assignment
            return assignment
        rows = self.tree_rows() if hasattr(self, "tree_rows") else []
        grid_indices = [b[0] for b in boxes]
        text_by_ri = {b[0]: t for b, t in zip(boxes, texts)}
        # Build all (score, row, grid_index) candidates.
        candidates = []
        for row in rows:
            for ri in grid_indices:
                sc = self._grid_row_match_score(row, text_by_ri.get(ri, ""))
                if sc > 0:
                    candidates.append((sc, id(row), ri))
        candidates.sort(key=lambda t: t[0], reverse=True)
        used_rows: set = set()
        used_grid: set = set()
        for sc, rid, ri in candidates:
            if rid in used_rows or ri in used_grid:
                continue
            assignment[rid] = ri
            used_rows.add(rid)
            used_grid.add(ri)
        self._grid_row_assignment = assignment
        return assignment

    def render_preview(self, page: int = 1, highlight: Optional[MaterialLine] = None,
                       keep_scroll: bool = False) -> None:
        # Save current scroll position before wiping the canvas so we can
        # restore it when we are NOT scrolling to a specific highlight.
        try:
            saved_x = self.preview_canvas.xview()[0]
            saved_y = self.preview_canvas.yview()[0]
        except Exception:
            saved_x = saved_y = 0.0
        self.preview_canvas.delete("all")
        self.preview_photo = None
        self.preview_pil_image = None
        self.preview_page = page or 1
        self.preview_scale = 1.0
        self.preview_original_size = (0, 0)
        self.preview_display_size = (0, 0)
        self._render_scroll_target: Optional[Tuple[float, float]] = None  # set by draw_canvas_box
        self._render_saved_scroll = (saved_x, saved_y)
        self._render_keep_scroll = keep_scroll
        if not self.current_result:
            return
        path = Path(self.current_result.file_path)
        if not path.exists():
            self.preview_canvas.create_text(20, 20, anchor="nw", text="Preview unavailable: file not found.")
            return
        self.preview_total_pages = self.preview_page_count(path)
        self.preview_page = max(1, min(self.preview_page, self.preview_total_pages))
        self.update_preview_page_controls()
        ext = path.suffix.lower()
        try:
            if ext == ".pdf":
                if fitz is None or Image is None or ImageTk is None:
                    self.preview_canvas.create_text(20, 20, anchor="nw", text="PDF preview needs PyMuPDF and Pillow installed.")
                    return
                doc = fitz.open(str(path))
                page_idx = max(0, min((self.preview_page or 1) - 1, len(doc) - 1))
                pdf_page = doc[page_idx]
                zoom = 1.6
                mat = fitz.Matrix(zoom, zoom)
                pix = pdf_page.get_pixmap(matrix=mat, alpha=False)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                doc.close()
                self.preview_scale = zoom
                self.preview_original_size = (pix.width, pix.height)
            elif ext in SUPPORTED_IMAGE_EXTS:
                if Image is None or ImageTk is None:
                    self.preview_canvas.create_text(20, 20, anchor="nw", text="Image preview needs Pillow installed.")
                    return
                img = Image.open(str(path)).convert("RGB")
                self.preview_total_pages = 1
                self.preview_page = 1
                self.update_preview_page_controls()
                self.preview_original_size = img.size
            elif ext in (SUPPORTED_SPREADSHEET_EXTS | SUPPORTED_WORD_EXTS):
                # Render the table grid to an image so spreadsheets/Word tables get a
                # real visual preview with per-row highlighting (matches PDF/image UX).
                if Image is None or ImageTk is None:
                    self.preview_canvas.create_text(20, 20, anchor="nw",
                        text="Grid preview needs Pillow installed (pip install pillow).")
                    return
                img = self._render_table_preview_image(path)
                if img is None:
                    self.preview_canvas.create_text(20, 20, anchor="nw",
                        text="Could not build a grid preview for this file. Use the text tab.")
                    return
                self.preview_total_pages = 1
                self.preview_page = 1
                self.update_preview_page_controls()
                self.preview_original_size = img.size
            else:
                self.preview_canvas.create_text(20, 20, anchor="nw", text="Preview is available for PDF/image/spreadsheet files. Use the text tab for this file type.")
                return

            # Keep very large pages responsive.
            max_side = 1800
            if max(img.size) > max_side:
                shrink = max_side / max(img.size)
                new_size = (max(1, int(img.size[0] * shrink)), max(1, int(img.size[1] * shrink)))
                img = img.resize(new_size)
                self.preview_scale *= shrink
            self.preview_pil_image = img
            self.preview_display_size = img.size
            self.preview_photo = ImageTk.PhotoImage(img)
            self.preview_canvas.create_image(0, 0, anchor="nw", image=self.preview_photo)
            self.preview_canvas.configure(scrollregion=(0, 0, img.size[0], img.size[1]))

            is_grid = ext in (SUPPORTED_SPREADSHEET_EXTS | SUPPORTED_WORD_EXTS)
            if is_grid:
                self.preview_canvas.create_text(8, 8, anchor="nw",
                    text="Table preview — click a material row to highlight its source row.",
                    fill="#333333")
                self.draw_grid_row_highlights(selected=highlight)
            else:
                self.preview_canvas.create_text(8, 8, anchor="nw", text=f"Page {self.preview_page} of {self.preview_total_pages}", fill="#333333")
                # Keep the file preview focused on the single selected material row.
                # The user reviews one row at a time; drawing every mapped box at once
                # makes image/PDF tests look noisy and can obscure the selected row.
                if highlight:
                    # Only highlight a row on the page it actually lives on.  When the
                    # user manually pages away (Prev/Next/Go) while a row is selected,
                    # the row's stored bbox is in ITS page's coordinates — drawing it
                    # here would put a box in the wrong place, and OCR-locating would
                    # search the wrong page.  Guard on source_page so multi-page review
                    # stays correct.
                    row_pg = highlight.source_page or 0
                    on_row_page = (row_pg == 0) or (row_pg == self.preview_page)
                    if not on_row_page:
                        self.preview_canvas.create_text(
                            20, 34, anchor="nw",
                            text=f"Selected row is on page {row_pg}. "
                                 f"Showing page {self.preview_page}.",
                            fill="#666666",
                        )
                    else:
                        # Auto-locate: image files (and image PDFs) have no precomputed
                        # bbox, so a plain row click previously showed the page with NO
                        # highlight.  If this row has no box yet, find it now — first via
                        # the cached visual-OCR locator, then by interpolating from
                        # neighbors — so the highlight simply appears on click.
                        if not highlight.source_bbox:
                            located = False
                            if pytesseract is not None:
                                try:
                                    located = bool(self.locate_row_on_current_preview(highlight))
                                except Exception:
                                    located = False
                            if not located:
                                try:
                                    self._guess_preview_bbox_from_neighbors(highlight)
                                except Exception:
                                    pass
                        if highlight.source_bbox:
                            self.draw_preview_highlight(highlight)
                        else:
                            self.preview_canvas.create_text(
                                20, 34, anchor="nw",
                                text=("Could not auto-locate this row on the scan. Use “Find Highlight”, "
                                      "or drag a box on the preview to set it."),
                                fill="#8a5a00",
                            )
                else:
                    self.preview_canvas.create_text(
                        20, 34, anchor="nw",
                        text="Select a material row to highlight its preview location.",
                        fill="#666666",
                    )
            # If no highlight scrolled us anywhere, restore the previous position.
            if self._render_scroll_target is None:
                sx, sy = self._render_saved_scroll
                if keep_scroll or not highlight:
                    self.preview_canvas.after_idle(
                        lambda x=sx, y=sy: (
                            self.preview_canvas.xview_moveto(x),
                            self.preview_canvas.yview_moveto(y),
                        )
                    )
        except Exception as e:
            self.preview_canvas.create_text(20, 20, anchor="nw", text=f"Preview failed: {e}")

    def draw_grid_row_highlights(self, selected: Optional[MaterialLine] = None) -> None:
        """Highlight ONLY the currently selected material row on the grid preview
        (bold box, scrolled into view).  Per user request, other rows are not boxed
        so the preview stays clean and the user can check items one at a time.
        Coordinates are image-original px; draw_canvas_box converts via preview_scale."""
        if not getattr(self, "_table_row_boxes", None):
            if selected:
                self.preview_canvas.create_text(20, 34, anchor="nw",
                    text="No grid rows mapped for this file.", fill="#666666")
            return
        if selected is None:
            return
        scale = self.preview_scale or 1.0
        selected_box = self._find_table_box_for_row(selected)
        if selected_box:
            x0, y0, x1, y1 = [v * scale for v in selected_box]
            self.draw_canvas_box(x0, y0, x1, y1,
                                 label=self.preview_label_for_row(selected) + " — selected",
                                 color=self.preview_color_for_row(selected),
                                 selected=True, scroll=True)
        else:
            self.preview_canvas.create_text(20, 34, anchor="nw",
                text="Could not locate this row in the table grid. Use the text tab highlight.",
                fill="#8a5a00")

    # ── v2.2: User-drawn bbox on File Preview ────────────────────────────────

    def _on_bbox_draw_toggle(self) -> None:
        """Toggle draw mode cursor/status."""
        if self.bbox_draw_var.get():
            self.preview_canvas.config(cursor="crosshair")
            self.status_var.set("Draw mode ON — drag a rectangle on the preview to set this item's location.")
        else:
            self.preview_canvas.config(cursor="")
            self.status_var.set("Draw mode OFF.")

    def _current_preview_bbox_kind(self) -> str:
        """Return the stable coordinate kind for the currently rendered preview.

        PDF highlights are stored in PDF/page coordinates. Image highlights are
        stored in the image's original pixel coordinates. The canvas may be a
        zoomed PDF render or a shrunken image, so we always convert through
        preview_scale instead of saving raw screen/canvas pixels.
        """
        try:
            if self.current_result and Path(self.current_result.file_path).suffix.lower() == ".pdf":
                return "pdf"
        except Exception:
            pass
        return "image"

    def _canvas_to_doc_coords(self, cx: float, cy: float) -> Tuple[float, float]:
        """Convert rendered canvas coordinates to stable document coordinates."""
        scale = self.preview_scale or 1.0
        if scale <= 0:
            scale = 1.0
        return cx / scale, cy / scale

    def _doc_to_canvas_coords(self, dx: float, dy: float, kind: str = "") -> Tuple[float, float]:
        """Convert stored PDF/image coordinates back into rendered canvas coordinates."""
        kind = (kind or "").lower().strip()
        # Compatibility: very old v2.5 test builds sometimes stored OCR/display
        # pixels as kind=image. New saves use image-original or PDF-page coords.
        # The normal path for both pdf and image is therefore scale-by-preview_scale.
        if kind in {"canvas", "display", "preview"}:
            return dx, dy
        scale = self.preview_scale or 1.0
        if scale <= 0:
            scale = 1.0
        return dx * scale, dy * scale

    def _bbox_area(self, box: Tuple[float, float, float, float]) -> float:
        x0, y0, x1, y1 = box
        return max(0.0, x1 - x0) * max(0.0, y1 - y0)

    def _bbox_overlap_ratio(self, a: Tuple[float, float, float, float], b: Tuple[float, float, float, float]) -> float:
        """Intersection over the smaller box area. High means one highlight sits on another."""
        ax0, ay0, ax1, ay1 = a
        bx0, by0, bx1, by1 = b
        ix0, iy0 = max(ax0, bx0), max(ay0, by0)
        ix1, iy1 = min(ax1, bx1), min(ay1, by1)
        inter = self._bbox_area((ix0, iy0, ix1, iy1))
        if inter <= 0:
            return 0.0
        denom = max(1.0, min(self._bbox_area(a), self._bbox_area(b)))
        return inter / denom

    def _bbox_conflicts_with_existing(self, row: "MaterialLine", canvas_box: Tuple[float, float, float, float], threshold: float = 0.42) -> bool:
        """True when this proposed preview box would be drawn over another row's box."""
        if not self.current_result:
            return False
        page = self.preview_page or row.source_page or 1
        for other in self.tree_rows():
            if other is row or not other.source_bbox:
                continue
            opage = other.source_page or self.infer_page_for_row(other) or page
            if opage != page:
                continue
            obox = self._parse_row_bbox_canvas(other)
            if obox and self._bbox_overlap_ratio(canvas_box, obox) >= threshold:
                return True
        return False

    def _clear_conflicting_bboxes_for_row(self, row: "MaterialLine", canvas_box: Tuple[float, float, float, float], threshold: float = 0.42) -> int:
        """When the user manually corrects a box, remove stale boxes it replaces."""
        if not self.current_result:
            return 0
        page = self.preview_page or row.source_page or 1
        cleared = 0
        for other in self.tree_rows():
            if other is row or not other.source_bbox:
                continue
            opage = other.source_page or self.infer_page_for_row(other) or page
            if opage != page:
                continue
            obox = self._parse_row_bbox_canvas(other)
            if obox and self._bbox_overlap_ratio(canvas_box, obox) >= threshold:
                other.source_bbox = ""
                other.source_bbox_kind = "pdf"
                cleared += 1
        return cleared

    def _store_display_bbox_on_row(self, row: "MaterialLine", x0: float, y0: float, x1: float, y1: float, *, page: Optional[int] = None, replace_conflicts: bool = False) -> None:
        """Save a rendered preview bbox on a row using stable document coords.

        x/y values passed here are canvas/display pixels for the currently
        rendered page. They are immediately converted to PDF points (for PDFs)
        or original image pixels (for image files) so redraw, scrolling, saving,
        and reload all keep the same physical location.
        """
        cx0, cy0, cx1, cy1 = min(x0, x1), min(y0, y1), max(x0, x1), max(y0, y1)
        if replace_conflicts:
            self._clear_conflicting_bboxes_for_row(row, (cx0, cy0, cx1, cy1))
        dx0, dy0 = self._canvas_to_doc_coords(cx0, cy0)
        dx1, dy1 = self._canvas_to_doc_coords(cx1, cy1)
        row.source_bbox = f"{dx0:.2f},{dy0:.2f},{dx1:.2f},{dy1:.2f}"
        row.source_bbox_kind = self._current_preview_bbox_kind()
        row.source_page = page or self.preview_page or row.source_page or 1

    def _bbox_mouse_press(self, event) -> None:
        if not self.bbox_draw_var.get():
            return
        # Convert event coords to canvas coords (scroll-adjusted)
        cx = self.preview_canvas.canvasx(event.x)
        cy = self.preview_canvas.canvasy(event.y)
        self._bbox_draw_start = (cx, cy)
        # Remove any existing temp rect
        if self._bbox_rect_id is not None:
            try:
                self.preview_canvas.delete(self._bbox_rect_id)
            except Exception:
                pass
            self._bbox_rect_id = None

    def _bbox_mouse_drag(self, event) -> None:
        if not self.bbox_draw_var.get() or self._bbox_draw_start is None:
            return
        cx = self.preview_canvas.canvasx(event.x)
        cy = self.preview_canvas.canvasy(event.y)
        x0, y0 = self._bbox_draw_start
        if self._bbox_rect_id is not None:
            try:
                self.preview_canvas.delete(self._bbox_rect_id)
            except Exception:
                pass
        self._bbox_rect_id = self.preview_canvas.create_rectangle(
            x0, y0, cx, cy,
            outline="#e65c00", width=2, dash=(6, 3),
        )
        # Live coord label in doc space
        dx0, dy0 = self._canvas_to_doc_coords(min(x0, cx), min(y0, cy))
        dx1, dy1 = self._canvas_to_doc_coords(max(x0, cx), max(y0, cy))
        self.bbox_coords_var.set(f"box: {dx0:.0f},{dy0:.0f} → {dx1:.0f},{dy1:.0f}")

    def _bbox_mouse_release(self, event) -> None:
        if not self.bbox_draw_var.get() or self._bbox_draw_start is None:
            return
        cx = self.preview_canvas.canvasx(event.x)
        cy = self.preview_canvas.canvasy(event.y)
        x0, y0 = self._bbox_draw_start
        self._bbox_draw_start = None
        # Remove temp rect
        if self._bbox_rect_id is not None:
            try:
                self.preview_canvas.delete(self._bbox_rect_id)
            except Exception:
                pass
            self._bbox_rect_id = None

        # Ignore tiny accidental clicks (< 5 px)
        if abs(cx - x0) < 5 or abs(cy - y0) < 5:
            return

        row = self.selected_material_line()
        if not row:
            messagebox.showinfo("SmartScan", "Select a material row in the table first, then draw its location on the preview.")
            return

        # Store the drawn canvas rectangle in stable document coordinates.
        # This fixes drift/misalignment caused by saving display pixels from a
        # zoomed PDF or shrunken image preview.
        self._store_display_bbox_on_row(row, x0, y0, cx, cy, page=self.preview_page or 1, replace_conflicts=True)
        dx0, dy0, dx1, dy1 = [float(v) for v in row.source_bbox.split(",")]
        row.user_decision = row.user_decision or "manual"
        if row.review_status == "pending":
            row.review_status = "corrected"
        # Refresh item in tree
        for item, r in self.tree_item_map.items():
            if r is row:
                self.refresh_tree_row(item)
                break

        self.bbox_coords_var.set(f"✓ saved  {dx0:.0f},{dy0:.0f} → {dx1:.0f},{dy1:.0f}")
        self.status_var.set(f"Highlight box set for row #{row.scan_index}. Click 'Save Corrections' to persist it.")
        # Turn off draw mode and redraw page with new box (keep current scroll)
        self.bbox_draw_var.set(False)
        self._on_bbox_draw_toggle()
        self.render_preview(page=self.preview_page or 1, highlight=row, keep_scroll=True)
        # v2.4: auto-cross-reference — switch to extracted text tab and highlight
        # the matching source text so the user can confirm/adjust the text span.
        self.root.after(300, lambda r=row: self._xref_preview_to_text(r))

    def _clear_preview_bbox(self) -> None:
        """Remove the stored bbox for the selected row so auto-detection takes over."""
        row = self.selected_material_line()
        if not row:
            return
        row.source_bbox = ""
        row.source_bbox_kind = "pdf"
        self.bbox_coords_var.set("")
        self.render_preview(page=self.preview_page or 1, highlight=row)
        self.status_var.set(f"Bbox cleared for row #{row.scan_index}. Auto-detection will re-run on next highlight.")

    # ── v2.4: bidirectional cross-reference helpers ──────────────────────────

    def _xref_text_to_preview(self, row: "MaterialLine") -> None:
        """After the user sets a text highlight, switch to the File Preview tab,
        guess the matching visual box, and ask the user to confirm it.

        The locator tries exact visual OCR first. If that fails, it estimates the
        missing row position from the closest already-known rows above/below.
        """
        if not self.current_result:
            return
        page = row.source_page or self.infer_page_for_row(row) or self.preview_page or 1
        self.evidence_tabs.select(1)
        self.render_preview(page=page, highlight=row)

        guessed = False
        guess_source = "existing"
        if row.source_bbox:
            guessed = True
        elif pytesseract is not None:
            located = self.locate_row_on_current_preview(row)
            if located:
                guessed = True
                guess_source = "visual OCR"
        if not guessed:
            guessed = self._guess_preview_bbox_from_neighbors(row)
            if guessed:
                guess_source = "nearby rows"

        if guessed:
            self.render_preview(page=row.source_page or page, highlight=row)
            self.status_var.set(
                f"Row #{row.scan_index}: guessed preview highlight from {guess_source}. Confirm it or draw a corrected box."
            )
            self.root.after(150, lambda r=row: self._confirm_preview_guess_popup(r))
        else:
            self.status_var.set(
                f"Row #{row.scan_index}: text highlight set, but SmartScan could not guess the preview box. "
                "Draw the box manually on the preview."
            )
            self.root.after(150, lambda r=row: self._offer_manual_preview_box_popup(r))

    def _xref_preview_to_text(self, row: "MaterialLine") -> None:
        """After the user draws a bbox on the preview, switch to the Merged Text
        tab, highlight the best matching text in orange, and ask for confirmation.
        """
        if not self.current_result:
            return
        self.show_text_source("merged")
        self.evidence_tabs.select(0)

        span = self._best_text_span_for_row(row)
        if span:
            start, end = span
            self._show_orange_text_confirmation(row, start, end)
            self.status_var.set(
                f"Row #{row.scan_index}: preview box saved. Orange text is SmartScan's linked text guess."
            )
            self.root.after(150, lambda r=row, s=start, e=end: self._confirm_text_guess_popup(r, s, e))
        else:
            self.status_var.set(
                f"Row #{row.scan_index}: preview box saved. SmartScan could not find matching extracted text. "
                "Select the correct text and click Draw Highlight."
            )
            self.root.after(150, lambda r=row: self._offer_manual_text_link_popup(r))

    def _best_text_span_for_row(self, row: "MaterialLine") -> Optional[Tuple[int, int]]:
        """Return the most reasonable merged-text span for a row.
        Prefer a valid existing span, otherwise search by source line/description.
        """
        widget_len = self._tk_idx_to_char("end-1c")
        if widget_len <= 0:
            return None
        start = row.source_start
        end = row.source_end
        if start >= 0 and end > start and end <= widget_len:
            return start, end
        needles = [row.source_line, row.description, row.part_number]
        if row.qty and row.description:
            needles.insert(0, f"{row.qty} {row.description}")
        for needle in needles:
            needle = (needle or "").strip()
            if not needle:
                continue
            idx = self.text_box.search(needle, "1.0", nocase=True, stopindex="end")
            if idx:
                end_idx = f"{idx}+{len(needle)}c"
                s = self._tk_idx_to_char(idx)
                e = self._tk_idx_to_char(end_idx)
                if s >= 0 and e > s:
                    return s, e
        return None

    def _show_orange_text_confirmation(self, row: "MaterialLine", start: int, end: int) -> None:
        """Show a temporary orange confirmation tag for guessed text."""
        for tag in ("evidence_current", "evidence_soft", "evidence_ok",
                    "evidence_warning", "evidence_error", "evidence_suggestion",
                    "evidence_confirm"):
            self.text_box.tag_remove(tag, "1.0", "end")
        self.text_box.tag_add("evidence_confirm", f"1.0+{start}c", f"1.0+{end}c")
        self.text_box.see(f"1.0+{start}c")
        self.hl_span_var.set(f"span: {end - start} chars  [{start}:{end}]")

    def _confirm_preview_guess_popup(self, row: "MaterialLine") -> None:
        """Ask whether the guessed preview box is correct. No enables editing."""
        if not row or not self.current_result:
            return
        ok = messagebox.askyesno(
            "Confirm preview highlight",
            "SmartScan guessed the location on the scan preview.\n\nIs the highlighted box correct?"
        )
        if ok:
            row.user_decision = row.user_decision or "confirmed_preview_highlight"
            if row.review_status == "pending":
                row.review_status = "corrected"
            self._refresh_current_tree_row(row)
            self.status_var.set(f"Preview highlight confirmed for row #{row.scan_index}. Click Save Corrections to persist it.")
        else:
            self.bbox_draw_var.set(True)
            self._on_bbox_draw_toggle()
            self.status_var.set(
                f"Adjust row #{row.scan_index}: drag a new highlight box on the preview, then release to save it."
            )

    def _confirm_text_guess_popup(self, row: "MaterialLine", start: int, end: int) -> None:
        """Ask whether the orange extracted-text guess is correct. No leaves it editable."""
        if not row or not self.current_result:
            return
        ok = messagebox.askyesno(
            "Confirm extracted text",
            "SmartScan highlighted the matching extracted text in orange.\n\nIs this the correct text for the item?"
        )
        if ok:
            merged = self.current_result.merged_text or ""
            row.source_start = start
            row.source_end = end
            row.source_line = merged[start:end] if end <= len(merged) else self.text_box.get(f"1.0+{start}c", f"1.0+{end}c")
            self._apply_confirmed_source_text_to_material_row(row, row.source_line)
            row.user_decision = row.user_decision or "confirmed_text_highlight"
            if row.review_status == "pending":
                row.review_status = "corrected"
            self._refresh_current_tree_row(row)
            self.apply_all_source_highlights()
            self.highlight_selected_evidence()
            self.status_var.set(f"Text highlight confirmed for row #{row.scan_index}. Click Save Corrections to persist it.")
        else:
            # Keep the orange tag visible so the user can see what was guessed,
            # but let them select the correct text and click Draw Highlight.
            self.status_var.set(
                f"Adjust row #{row.scan_index}: select the correct extracted text, then click Draw Highlight."
            )

    def _offer_manual_preview_box_popup(self, row: "MaterialLine") -> None:
        if not row:
            return
        messagebox.showinfo(
            "Draw preview highlight",
            "SmartScan could not confidently guess the preview box.\n\nDraw a box around the item on the preview, then release the mouse to save it."
        )
        self.bbox_draw_var.set(True)
        self._on_bbox_draw_toggle()

    def _offer_manual_text_link_popup(self, row: "MaterialLine") -> None:
        if not row:
            return
        messagebox.showinfo(
            "Select extracted text",
            "SmartScan could not confidently find the matching extracted text.\n\nSelect the correct text in the extracted-text view, then click Draw Highlight."
        )

    def _refresh_current_tree_row(self, row: "MaterialLine") -> None:
        for item, r in self.tree_item_map.items():
            if r is row:
                self.refresh_tree_row(item)
                break
        self.update_review_status_bar()

    def _apply_confirmed_source_text_to_material_row(self, row: "MaterialLine", text: str) -> bool:
        """Use confirmed highlighted text to populate Qty/UOM/Description.

        This is intentionally run only after the user confirms or manually links
        the source text, so it helps corrected/added rows without constantly
        rewriting rows from uncertain OCR guesses.
        """
        if not row:
            return False
        raw = normalize_space(text or "")
        raw = re.sub(r"\s*\n\s*", "  ", raw).strip(" |\t")
        if not raw:
            return False

        parser = MaterialParser(self.db, self.settings)
        parsed = parser.parse_line(raw)
        if parsed is None:
            # Some highlighted table selections contain pipes/tabs or a row number
            # split from the qty. Try the same text with table separators preserved.
            tableish = re.sub(r"\s{2,}", " | ", raw)
            parsed = parser.parse_line(tableish) or parser.parse_numbered_table_row(raw)
        if parsed is None:
            return False

        row.snapshot_original()
        if parsed.qty:
            row.qty = parsed.qty
        if parsed.unit:
            row.unit = parsed.unit
        if parsed.description:
            row.description = parsed.description
        if parsed.part_number:
            row.part_number = parsed.part_number
        if parsed.manufacturer:
            row.manufacturer = parsed.manufacturer
        row.confidence = max(row.confidence or 0, parsed.confidence or 0, 80)
        row.notes = append_note(row.notes, "Fields updated from confirmed highlight text")
        row.user_decision = row.user_decision or "confirmed_highlight_parse"
        if row.review_status == "pending":
            row.review_status = "corrected"

        if getattr(self.settings, "enable_rule_review", True) and self.current_result:
            AIReviewer(self.settings, log=self.log).review([row], self.current_result.merged_text or "")
            st = (row.ai_status or "").lower()
            row.highlight_color = "green" if st == "ok" else ("red" if st == "error" else "yellow")
        return True

    def _parse_row_bbox_canvas(self, row: "MaterialLine") -> Optional[Tuple[float, float, float, float]]:
        """Return row bbox in current rendered canvas/display coordinates."""
        if not row.source_bbox:
            return None
        try:
            x0, y0, x1, y1 = [float(x) for x in row.source_bbox.split(",")]
            cx0, cy0 = self._doc_to_canvas_coords(x0, y0, row.source_bbox_kind)
            cx1, cy1 = self._doc_to_canvas_coords(x1, y1, row.source_bbox_kind)
            return cx0, cy0, cx1, cy1
        except Exception:
            return None

    def _guess_preview_bbox_from_neighbors(self, row: "MaterialLine") -> bool:
        """Estimate a missing visual box from nearby rows that already have boxes.

        This is intentionally conservative: it only uses rows on the same page as
        the selected row/current preview and stores the estimate as image/canvas
        coordinates so the existing drawing logic can show it immediately.
        """
        if not row or not self.current_result or self.preview_pil_image is None:
            return False
        rows = self.tree_rows()
        if row not in rows:
            return False
        page = row.source_page or self.infer_page_for_row(row) or self.preview_page or 1
        idx = rows.index(row)

        def usable(r: "MaterialLine") -> Optional[Tuple[float, float, float, float]]:
            rpage = r.source_page or self.infer_page_for_row(r) or page
            if rpage != page:
                return None
            box = self._parse_row_bbox_canvas(r)
            if box:
                return box
            # Give nearby rows a chance to get their own visual OCR box first;
            # then the selected row can be interpolated between them.
            if pytesseract is not None:
                try:
                    located = self.locate_row_on_current_preview(r)
                    if located:
                        return self._parse_row_bbox_canvas(r)
                except Exception:
                    return None
            return None

        prev = None
        for r in reversed(rows[:idx]):
            box = usable(r)
            if box:
                prev = (r, box)
                break
        nxt = None
        for r in rows[idx + 1:]:
            box = usable(r)
            if box:
                nxt = (r, box)
                break
        if not prev and not nxt:
            return False

        img_w, img_h = self.preview_display_size or self.preview_original_size or (1, 1)
        if prev and nxt:
            pbox = prev[1]
            nbox = nxt[1]
            py = (pbox[1] + pbox[3]) / 2.0
            ny = (nbox[1] + nbox[3]) / 2.0
            y_center = (py + ny) / 2.0
            h = max(18.0, min(70.0, ((pbox[3] - pbox[1]) + (nbox[3] - nbox[1])) / 2.0))
            x0 = min(pbox[0], nbox[0])
            x1 = max(pbox[2], nbox[2])
        elif prev:
            pbox = prev[1]
            h = max(18.0, min(70.0, pbox[3] - pbox[1]))
            y_center = min(img_h - h / 2.0, ((pbox[1] + pbox[3]) / 2.0) + h * 1.35)
            x0, x1 = pbox[0], pbox[2]
        else:
            nbox = nxt[1]
            h = max(18.0, min(70.0, nbox[3] - nbox[1]))
            y_center = max(h / 2.0, ((nbox[1] + nbox[3]) / 2.0) - h * 1.35)
            x0, x1 = nbox[0], nbox[2]

        # Make the box large enough to be editable/visible, and keep it on page.
        min_w = max(160.0, img_w * 0.18)
        if (x1 - x0) < min_w:
            cx = (x0 + x1) / 2.0
            x0, x1 = cx - min_w / 2.0, cx + min_w / 2.0
        pad_x = 12.0
        pad_y = 5.0
        x0 = max(0.0, x0 - pad_x)
        x1 = min(float(img_w), x1 + pad_x)
        y0 = max(0.0, y_center - h / 2.0 - pad_y)
        y1 = min(float(img_h), y_center + h / 2.0 + pad_y)
        if x1 <= x0 or y1 <= y0:
            return False
        if self._bbox_conflicts_with_existing(row, (x0, y0, x1, y1)):
            return False
        self._store_display_bbox_on_row(row, x0, y0, x1, y1, page=page)
        return True

    def draw_preview_highlight(self, row: MaterialLine) -> None:
        if not row:
            return
        color = self.preview_color_for_row(row)
        if row.source_bbox and (not row.source_page or row.source_page == self.preview_page):
            if self._draw_known_bbox(row, label="selected evidence", color=color, selected=True):
                return
        located = self.locate_row_on_current_preview(row)
        if located:
            x0, y0, x1, y1 = located
            self.draw_canvas_box(x0, y0, x1, y1, "selected visual OCR match", color=color, selected=True)
            return
        self.preview_canvas.create_text(
            20, 34, anchor="nw",
            text="No exact visual box found on this page. Use Prev/Next or the text tab highlight. Install/enable Tesseract for scanned-page visual locating.",
            fill="#8a5a00",
        )

    def draw_all_review_highlights_on_page(self, selected: Optional[MaterialLine] = None) -> None:
        """Draw evidence boxes for the current page without stacking boxes.

        Rows are drawn in visual reading order (top-left to right, then down).
        If two stored boxes substantially overlap, only the first/selected one is
        shown. This keeps stale or duplicate scan guesses from covering each
        other on the preview.
        """
        if not self.current_result:
            return
        selected_id = id(selected) if selected else None
        any_box = False
        drawn_boxes: List[Tuple[float, float, float, float]] = []

        selected_box = self._parse_row_bbox_canvas(selected) if selected and selected.source_bbox else None
        candidates: List[Tuple[float, float, float, float, MaterialLine]] = []
        for row in self.tree_rows():
            if not row.source_bbox:
                continue
            if row.source_page and row.source_page != self.preview_page:
                continue
            box = self._parse_row_bbox_canvas(row)
            if not box:
                continue
            x0, y0, x1, y1 = box
            candidates.append((y0, x0, x1, y1, row))

        # Draw non-selected boxes first, in reading order. Do not draw boxes that
        # would sit under/over the selected box.
        for _y0, _x0, _x1, _y1, row in sorted(candidates, key=lambda t: (int(round(t[0] / 8.0)), t[1])):
            if selected_id is not None and id(row) == selected_id:
                continue
            box = self._parse_row_bbox_canvas(row)
            if not box:
                continue
            if selected_box and self._bbox_overlap_ratio(box, selected_box) >= 0.42:
                continue
            if any(self._bbox_overlap_ratio(box, drawn) >= 0.42 for drawn in drawn_boxes):
                continue
            if self._draw_known_bbox(row, label=self.preview_label_for_row(row), color=self.preview_color_for_row(row), selected=False, scroll=False):
                drawn_boxes.append(box)
                any_box = True

        if selected:
            # Redraw the selected row on top and scroll to it. If it has no known
            # bbox, this method will try the OCR visual locator.
            self.draw_preview_highlight(selected)
        elif not any_box:
            self.preview_canvas.create_text(
                20, 34, anchor="nw",
                text="No mapped material boxes on this page yet. Select a row or click Find Highlight.",
                fill="#666666",
            )

    def preview_label_for_row(self, row: MaterialLine) -> str:
        idx = row.scan_index or ""
        status = (row.ai_status or "pending").lower()
        if status == "ok":
            return f"#{idx} OK"
        if status == "error":
            return f"#{idx} ERROR"
        if status == "warning":
            return f"#{idx} REVIEW"
        if status == "suggestion":
            return f"#{idx} FIX?"
        return f"#{idx}"

    def preview_color_for_row(self, row: MaterialLine) -> str:
        st = (row.ai_status or "").lower()
        if st == "ok":
            return "green"
        if st == "error":
            return "red"
        if st == "warning":
            return "yellow"
        if st == "suggestion":
            return "blue"
        if row.review_status == "rejected":
            return "red"
        return "yellow"

    def _draw_known_bbox(self, row: MaterialLine, label: str = "", color: str = "yellow", selected: bool = False, scroll: bool = True) -> bool:
        try:
            box = self._parse_row_bbox_canvas(row)
            if not box:
                return False
            x0, y0, x1, y1 = box
            self.draw_canvas_box(x0, y0, x1, y1, label, color=color, selected=selected, scroll=scroll)
            return True
        except Exception:
            return False

    def draw_canvas_box(self, x0: float, y0: float, x1: float, y1: float, label: str = "", color: str = "yellow", selected: bool = False, scroll: bool = True) -> None:
        palette = {
            "green": ("#244a32", "#5db87a", "#d7f7df"),
            "yellow": ("#5a4a22", "#c99632", "#fff4cc"),
            "red": ("#5a2a2a", "#d96a6a", "#ffdede"),
            "blue": ("#2d4260", "#6aa0ff", "#dcecff"),
        }
        fill, outline, text_color = palette.get(color, palette["yellow"])
        pad = 5 if selected else 3
        x0, y0, x1, y1 = x0 - pad, y0 - pad, x1 + pad, y1 + pad
        width = 5 if selected else 3
        self.preview_canvas.create_rectangle(x0, y0, x1, y1, fill=fill, stipple="gray25", outline=outline, width=width)
        self.preview_canvas.create_rectangle(x0, y0, x1, y1, outline=text_color, width=1)
        if label:
            self.preview_canvas.create_text(x0, max(0, y0 - 18), anchor="nw", text=label, fill=text_color)
        if scroll:
            w, h = self.preview_display_size or self.preview_original_size or (1, 1)
            tx = max(0.0, (x0 - 80) / max(1, w))
            ty = max(0.0, (y0 - 80) / max(1, h))
            self.preview_canvas.xview_moveto(tx)
            self.preview_canvas.yview_moveto(ty)
            # Signal that we scrolled so render_preview won't override us
            self._render_scroll_target = (tx, ty)

    def find_highlight_on_current_page(self) -> None:
        row = self.selected_material_line()
        if not row:
            messagebox.showinfo("SmartScan", "Select a material row first.")
            return
        self.render_preview(page=self.preview_page or row.source_page or 1, highlight=row)
        self.evidence_tabs.select(1)

    def _visible_line_number_for_row(self, row: MaterialLine) -> str:
        """Return the visible LINE column number when the source row exposes it.

        For image/PDF material tables, the visual preview contains both LINE and QTY.
        Highlight matching must use the LINE token as an anchor when available; otherwise
        two neighboring rows with the same Qty/size words can be confused.
        """
        if not row:
            return ""
        for value in (row.source_line,):
            src = normalize_space(value or "")
            if not src:
                continue
            m = re.match(r"^\s*\|?\s*(\d{1,4})\b", src)
            if m:
                return m.group(1)
        try:
            # Fallback only for clean generated/test lists where review_order follows
            # the visible line numbers.  Source-line match above is preferred.
            if row.scan_index and 0 < int(row.scan_index) <= 999:
                return str(int(row.scan_index))
        except Exception:
            pass
        return ""

    def _preview_word_bands(self, words: List[Tuple[str, float, float, float, float]]) -> List[dict]:
        """Cluster visual OCR words into horizontal table rows."""
        bands: List[dict] = []
        for (txt, x, y, w, h) in words:
            cy = y + h / 2.0
            band = None
            for b in bands:
                if abs(cy - b["cy"]) <= max(14.0, h * 1.15):
                    band = b
                    break
            if band is None:
                band = {"cy": cy, "items": []}
                bands.append(band)
            band["items"].append((txt, x, y, w, h))
            centers = [it[2] + it[4] / 2.0 for it in band["items"]]
            band["cy"] = sum(centers) / len(centers)
        for b in bands:
            b["items"].sort(key=lambda it: it[1])
        bands.sort(key=lambda b: b["cy"])
        return bands

    def _match_preview_table_row_band(self, words: List[Tuple[str, float, float, float, float]], row: MaterialLine) -> Optional[Tuple[float, float, float, float]]:
        """Prefer a full visual table-row band anchored by LINE + QTY + description.

        This fixes the off-by-one highlight issue seen on the 50-item OCR test:
        row 14 and 15 both had Qty=15 and nearly identical EMT descriptions, so a
        description-only/partial-token match could land on the neighbor row.  The
        visible LINE column is the decisive anchor when available.
        """
        if not words or not row:
            return None
        line_no = self._norm_preview_token(self._visible_line_number_for_row(row))
        qty_tok = self._norm_preview_token(row.qty or "")
        unit_tok = self._norm_preview_token(row.unit or "")
        desc_tokens = [self._norm_preview_token(t) for t in re.findall(r"[A-Za-z0-9#/\'\".-]+", row.description or "")]
        desc_tokens = [t for t in desc_tokens if t and t not in {"in", "inch", "ft", "ea"}]
        # Keep enough descriptors to distinguish connector vs coupling, strap sizes, etc.
        desc_want = set(desc_tokens[:10])
        if not (line_no or qty_tok or desc_want):
            return None
        bands = self._preview_word_bands(words)
        best_band = None
        best_score = 0.0
        best_line_hit = False
        for b in bands:
            items = b.get("items", [])
            toks = [it[0] for it in items]
            if not toks:
                continue
            score = 0.0
            # LINE number should be the left-most/early token in these tables. Give it
            # a large bonus and require it when we know it, so row 14 cannot select row 15.
            line_hit = bool(line_no and line_no in toks[:4])
            if line_no and line_hit:
                score += 80.0
            # If the LINE anchor was available but OCR missed/merged it, do not let
            # common Qty/UOM tokens dominate; use description words to distinguish
            # neighboring same-qty rows.
            allow_qty_unit_bonus = not (line_no and not line_hit)
            if allow_qty_unit_bonus and qty_tok and qty_tok in toks[:6]:
                score += 22.0
            if allow_qty_unit_bonus and unit_tok and unit_tok in toks[:8]:
                score += 12.0
            shared_desc = desc_want & set(toks)
            for t in shared_desc:
                if len(t) >= 7:
                    score += 8.0
                elif len(t) >= 4:
                    score += 5.0
                else:
                    score += 2.0
            # Must have at least a little description agreement; LINE+QTY alone is not
            # enough when a page has many repeated quantities.
            if desc_want and len(shared_desc) < min(2, len(desc_want)):
                continue
            if score > best_score:
                best_score = score
                best_band = items
                best_line_hit = bool(line_hit)
        # Normal case: LINE anchor hit, so require the high anchored score.  If OCR
        # merged/misread the LINE+QTY cells (example: line 14/qty 15 became "415"),
        # allow a description-only band only when it is strong enough to beat the
        # neighboring near-duplicate row.
        min_score = 82.0 if (line_no and best_line_hit) else (18.0 if line_no else 12.0)
        if not best_band or best_score < min_score:
            return None
        # Build the visible highlight box from real text/number tokens, not table
        # border tokens. Tesseract often reads vertical grid lines as "|" with a
        # very tall bbox spanning the whole cell; including those made early rows
        # such as line 4 look over-tall / slightly off even when the correct row
        # was selected. This keeps the highlight tight around the actual material
        # row text while preserving the exact row match logic above.
        # Build the box from the tokens that actually matched THIS row's
        # description (plus anchors on the same visual row), not from every token
        # in the band.  When OCR mis-reads the LINE/QTY column, _preview_word_bands
        # can fuse two physically-adjacent rows into one band; using all tokens then
        # stretched the highlight onto the neighbor (the off-by-one droop seen on the
        # 100-line image PDF, e.g. line 73 highlighting line 74).  Restricting the box
        # to the vertical extent of the matched description tokens keeps it on the
        # correct row.  Falls back to the old behavior when no description matched.
        matched_items = [it for it in best_band if it[0] in desc_want]
        if matched_items:
            row_top = min(it[2] for it in matched_items)
            row_bot = max(it[2] + it[4] for it in matched_items)
            row_pad = max(4.0, (row_bot - row_top) * 0.4)
            lo, hi = row_top - row_pad, row_bot + row_pad
            text_band = [
                it for it in best_band
                if lo <= (it[2] + it[4] / 2.0) <= hi
                and re.search(r"[a-z0-9]", str(it[0]), flags=re.I)
            ] or matched_items
        else:
            text_band = [it for it in best_band if re.search(r"[a-z0-9]", str(it[0]), flags=re.I)] or best_band
        xs0 = [it[1] for it in text_band]
        ys0 = [it[2] for it in text_band]
        xs1 = [it[1] + it[3] for it in text_band]
        ys1 = [it[2] + it[4] for it in text_band]
        x0, y0, x1, y1 = min(xs0), min(ys0), max(xs1), max(ys1)
        heights = sorted([it[4] for it in text_band]) or [12.0]
        mh = heights[len(heights) // 2]
        max_h = max(18.0, mh * 2.8)
        if (y1 - y0) > max_h:
            cy = (y0 + y1) / 2.0
            y0, y1 = cy - max_h / 2.0, cy + max_h / 2.0
        return (x0, y0, x1, y1)

    def locate_row_on_current_preview(self, row: MaterialLine) -> Optional[Tuple[float, float, float, float]]:
        if not row or self.preview_pil_image is None:
            return None
        if pytesseract is None:
            return None
        words = self._preview_ocr_words()
        if not words:
            return None
        candidates = self._preview_candidate_tokens(row)
        all_tokens: List[str] = []
        seen_t = set()
        for tokens in candidates:
            for t in tokens:
                if t not in seen_t:
                    seen_t.add(t)
                    all_tokens.append(t)
        part_tok = self._norm_preview_token(row.part_number or "")

        # First try a table-row band anchored by the visible LINE number.  This is
        # more accurate than description-only matching for multi-page image tables
        # with repeated quantities/sizes and nearly identical connector/coupling rows.
        table_box = self._match_preview_table_row_band(words, row)
        if table_box and not self._bbox_conflicts_with_existing(row, table_box):
            self._store_display_bbox_on_row(row, *table_box, page=self.preview_page or row.source_page or 1)
            return table_box

        # Part number is the unique row identifier.  If we have one, try the decisive
        # part-anchored band match FIRST: it lands on the exact row even when two items
        # have near-identical descriptions (e.g. "3/4 EMT compression connector CC-075"
        # vs "... coupling CPL-075"), which the strict in-order matcher can confuse.
        if part_tok and len(part_tok) >= 4 and all_tokens:
            box = self._match_preview_tokens_band(words, all_tokens, part_token=part_tok)
            if box and not self._bbox_conflicts_with_existing(row, box):
                self._store_display_bbox_on_row(row, *box, page=self.preview_page or row.source_page or 1)
                return box

        for tokens in candidates:
            box = self._match_preview_tokens(words, tokens)
            if box:
                if self._bbox_conflicts_with_existing(row, box):
                    continue
                # OCR returns coordinates in the currently rendered preview image.
                # Save them through the same stable conversion path as manual
                # drawing so future redraws/reloads keep the correct location.
                self._store_display_bbox_on_row(row, *box, page=self.preview_page or row.source_page or 1)
                return box
        # Final fallback: tolerant band match without a part anchor.
        if all_tokens:
            box = self._match_preview_tokens_band(words, all_tokens, part_token=part_tok)
            if box and not self._bbox_conflicts_with_existing(row, box):
                self._store_display_bbox_on_row(row, *box, page=self.preview_page or row.source_page or 1)
                return box
        return None

    def _preview_ocr_words(self) -> List[Tuple[str, float, float, float, float]]:
        """OCR the current preview image once and cache the word boxes.

        Per-row visual locating previously re-OCR'd the entire page on every call,
        which made clicking through rows slow.  We cache the normalized word list
        keyed by (file, page, display-size) — NOT id(img) — because on multi-page
        documents every page renders to the same size and CPython reuses memory
        addresses, so an id()-based key could hand page 2 the cached words from
        page 1 and send highlights to the wrong place.  Returns a list of
        (norm_text, left, top, width, height)."""
        img = self.preview_pil_image
        if img is None or pytesseract is None:
            return []
        try:
            src = str(self.current_result.file_path) if self.current_result else ""
        except Exception:
            src = ""
        cache_key = (src, int(self.preview_page or 1), getattr(img, "size", None))
        cached = getattr(self, "_preview_ocr_cache", None)
        if cached is not None and cached.get("key") == cache_key:
            return cached["words"]
        try:
            _smartscan_configure_pytesseract(self.settings)
            data = pytesseract.image_to_data(
                img,
                lang=self.settings.tesseract_lang or "eng",
                config="--psm 6 --oem 3",
                output_type=pytesseract.Output.DICT,
            )
        except Exception as e:
            self.log(f"Visual highlight OCR failed: {e}")
            return []
        words: List[Tuple[str, float, float, float, float]] = []
        n = len(data.get("text", []))
        for i in range(n):
            text = str(data["text"][i] or "").strip()
            norm = self._norm_preview_token(text)
            if not norm:
                continue
            try:
                conf = float(data.get("conf", [0] * n)[i])
            except Exception:
                conf = 0
            if conf < 0:
                continue
            words.append((norm, float(data["left"][i]), float(data["top"][i]),
                          float(data["width"][i]), float(data["height"][i])))
        # Read/order words the way the reviewer sees the page: left-to-right on each
        # horizontal band, then top-to-bottom.  Reduces duplicate row grabs on tables.
        words.sort(key=lambda w: (int(round(w[2] / 8.0)), w[1]))
        self._preview_ocr_cache = {"key": cache_key, "words": words}
        return words

    def _preview_candidate_tokens(self, row: MaterialLine) -> List[List[str]]:
        raw = [row.part_number, row.description, row.source_line]
        if row.qty and row.description:
            raw.insert(0, f"{row.qty} {row.description}")
        out: List[List[str]] = []
        seen = set()
        for value in raw:
            tokens = [self._norm_preview_token(t) for t in re.findall(r"[A-Za-z0-9#/\'\".-]+", value or "")]
            tokens = [t for t in tokens if t]
            if len(tokens) > 10:
                tokens = tokens[:10]
            if len(tokens) >= 2:
                key = tuple(tokens)
                if key not in seen:
                    seen.add(key)
                    out.append(tokens)
                if len(tokens) > 5:
                    short = tuple(tokens[:5])
                    if short not in seen:
                        seen.add(short)
                        out.append(list(short))
        return out

    def _match_preview_tokens(self, words: List[Tuple[str, float, float, float, float]], tokens: List[str]) -> Optional[Tuple[float, float, float, float]]:
        if len(tokens) < 2:
            return None
        best: Optional[Tuple[float, float, float, float]] = None
        best_score = 0
        first = tokens[0]
        for i, word in enumerate(words):
            if word[0] != first:
                continue
            # Keep visual OCR matches on one table row.  Older matching allowed
            # gaps across the whole page, which could union row 4 + row 5 into one
            # highlight when neighboring material descriptions shared tokens.
            base_y = word[2] + word[4] / 2.0
            row_band = max(18.0, min(34.0, word[4] * 2.2))
            matched = [word]
            token_idx = 1
            j = i + 1
            gaps = 0
            while j < len(words) and token_idx < len(tokens) and gaps <= 8:
                wy = words[j][2] + words[j][4] / 2.0
                if abs(wy - base_y) > row_band:
                    # Once the OCR stream has moved clearly to a new row, stop.
                    # Do not keep searching downward for remaining tokens.
                    if wy > base_y:
                        break
                    j += 1
                    continue
                if words[j][0] == tokens[token_idx]:
                    matched.append(words[j])
                    token_idx += 1
                    gaps = 0
                else:
                    gaps += 1
                j += 1
            score = token_idx
            if score >= min(3, len(tokens)) and score > best_score:
                xs0 = [m[1] for m in matched]
                ys0 = [m[2] for m in matched]
                xs1 = [m[1] + m[3] for m in matched]
                ys1 = [m[2] + m[4] for m in matched]
                x0, y0, x1, y1 = min(xs0), min(ys0), max(xs1), max(ys1)
                # Tighten the highlight to the detected OCR row instead of padding
                # enough to cover a neighboring row. draw_canvas_box adds its own
                # small visual padding when it draws the selected rectangle.
                median_h = sorted([m[4] for m in matched])[len(matched)//2] if matched else 12.0
                max_h = max(18.0, median_h * 2.4)
                if (y1 - y0) > max_h:
                    cy = (y0 + y1) / 2.0
                    y0, y1 = cy - max_h / 2.0, cy + max_h / 2.0
                best = (x0, y0, x1, y1)
                best_score = score
                if best_score >= min(6, len(tokens)):
                    break
        return best

    def _match_preview_tokens_band(self, words: List[Tuple[str, float, float, float, float]],
                                   tokens: List[str], part_token: str = "") -> Optional[Tuple[float, float, float, float]]:
        """Robust fallback: group OCR words into horizontal rows (bands), then pick the
        band that shares the most tokens with the row's tokens.  Unlike the strict
        in-order matcher, this tolerates OCR re-ordering, a stray leading line number,
        and small gaps — which is what made rows like "2,000 EA Deep kindorf ... KOS-158"
        or "... gray WPB-2G" fail to highlight even though every token was present.
        Bounds the result to the matched words on that band.

        When part_token is given and it appears in exactly one band, that band wins
        decisively — the part number is the unique row identifier, so this prevents two
        near-identical descriptions (e.g. CC-075 vs CPL-075) from landing on one line."""
        if not words or not tokens:
            return None
        want = set(t for t in tokens if t)
        if not want:
            return None
        # Cluster words into bands by vertical center.
        bands: List[dict] = []
        for (txt, x, y, w, h) in words:
            cy = y + h / 2.0
            band = None
            for b in bands:
                if abs(cy - b["cy"]) <= max(16.0, h * 1.3):
                    band = b
                    break
            if band is None:
                band = {"cy": cy, "items": []}
                bands.append(band)
            band["items"].append((txt, x, y, w, h))
            ys = [it[2] + it[4] / 2.0 for it in band["items"]]
            band["cy"] = sum(ys) / len(ys)

        def bound(items, fallback_items):
            # Vertical extent comes from the matched tokens (keeps it on one row), but
            # the HORIZONTAL extent spans the whole band so the highlight covers the
            # entire line — not just the few words that matched (which made short rows
            # like "4 SET tool for..." highlight only "SET").
            use = items or fallback_items
            ys0 = [it[2] for it in use]; ys1 = [it[2] + it[4] for it in use]
            row_words = fallback_items or use
            xs0 = [it[1] for it in row_words]; xs1 = [it[1] + it[3] for it in row_words]
            x0, y0, x1, y1 = min(xs0), min(ys0), max(xs1), max(ys1)
            mh = sorted([it[4] for it in use])[len(use) // 2]
            max_h = max(18.0, mh * 2.6)
            if (y1 - y0) > max_h:
                cy = (y0 + y1) / 2.0
                y0, y1 = cy - max_h / 2.0, cy + max_h / 2.0
            return (x0, y0, x1, y1)

        # Decisive part-number anchor: if the part token sits in exactly one band, use it.
        pt = (part_token or "").strip()
        if pt:
            hit_bands = [b for b in bands if any(it[0] == pt for it in b["items"])]
            if len(hit_bands) == 1:
                b = hit_bands[0]
                matched_items = [it for it in b["items"] if it[0] in want] or b["items"]
                return bound(matched_items, b["items"])

        best = None
        best_score = 0.0
        def tok_weight(t: str) -> float:
            if pt and t == pt:
                return 6.0  # exact part-number token is the strongest signal
            if re.search(r"\d", t) and re.search(r"[a-z]", t):
                return 3.0  # part-number-like
            if len(t) >= 6:
                return 2.0
            return 1.0
        for b in bands:
            band_tokens = set(it[0] for it in b["items"])
            shared = want & band_tokens
            if not shared:
                continue
            score = sum(tok_weight(t) for t in shared)
            if score > best_score:
                matched_items = [it for it in b["items"] if it[0] in want] or b["items"]
                best = bound(matched_items, b["items"])
                best_score = score
        if best is not None and best_score >= 2.0:
            return best
        return None

    def _norm_preview_token(self, text: str) -> str:
        return re.sub(r"[^a-z0-9]+", "", (text or "").lower())

    def copy_text(self) -> None:
        text = self.text_box.get("1.0", "end").strip()
        self.root.clipboard_clear()
        self.root.clipboard_append(text)
        self.status_var.set("Copied extracted text.")

    def copy_materials(self) -> None:
        rows = self.tree_rows()
        out = []
        for r in rows:
            out.append(f"{r.qty}\t{r.unit}\t{r.description}\t{r.part_number}\t{r.manufacturer}\t{r.confidence}\t{r.notes}")
        self.root.clipboard_clear()
        self.root.clipboard_append("\n".join(out))
        self.status_var.set("Copied material lines.")

    def export_csv(self) -> None:
        rows = self.tree_rows()
        if not rows:
            messagebox.showinfo("SmartScan", "No material rows to export.")
            return
        default_dir = self.settings.output_dir or str(Path.home() / "Desktop")
        path = filedialog.asksaveasfilename(
            initialdir=default_dir,
            title="Export materials CSV",
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv")],
        )
        if not path:
            return
        with open(path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.DictWriter(f, fieldnames=["scan_index", "review_status", "ai_status", "ai_message", "user_decision", "qty", "unit", "description", "part_number", "manufacturer", "notes", "confidence", "source_line"])
            writer.writeheader()
            for r in rows:
                writer.writerow({
                    "scan_index": r.scan_index,
                    "review_status": r.review_status,
                    "ai_status": r.ai_status,
                    "ai_message": r.ai_message,
                    "user_decision": r.user_decision,
                    "qty": r.qty,
                    "unit": r.unit,
                    "description": r.description,
                    "part_number": r.part_number,
                    "manufacturer": r.manufacturer,
                    "notes": r.notes,
                    "confidence": r.confidence,
                    "source_line": r.source_line,
                })
        self.status_var.set(f"Exported CSV: {path}")

    def export_json(self) -> None:
        if not self.current_result:
            messagebox.showinfo("SmartScan", "No result to export.")
            return
        self.current_result.materials = self.tree_rows()
        default_dir = self.settings.output_dir or str(Path.home() / "Desktop")
        path = filedialog.asksaveasfilename(
            initialdir=default_dir,
            title="Export scan JSON",
            defaultextension=".json",
            filetypes=[("JSON", "*.json")],
        )
        if not path:
            return
        Path(path).write_text(json.dumps(self.current_result.as_dict(), indent=2, ensure_ascii=False), encoding="utf-8")
        self.status_var.set(f"Exported JSON: {path}")

    def run(self) -> None:
        self.root.mainloop()


def run_mainbox_headless_scan(file_path: str, json_out: str = "") -> Dict[str, Any]:
    """Headless SmartScan entry point used by MaINbox.

    This intentionally runs the same SmartScan engine as the standalone app, but
    without creating any Tk windows.  MaINbox calls this in a separate Python
    process so SmartScan behaves like the standalone scanner and cannot be
    affected by MaINbox imports, COM state, or UI globals.
    """
    logs: List[str] = []

    def hlog(msg: Any) -> None:
        try:
            logs.append(str(msg))
        except Exception:
            logs.append(repr(msg))

    db = SmartScanDB()
    try:
        settings = db.load_settings()
    except Exception:
        settings = ScanSettings()

    # Match the standalone "Normal" SmartScan path for RFQ use while avoiding
    # stale cached rows.  The smart router inside ScanEngine still chooses the
    # native-text vs OCR/watchdog path exactly like the standalone app.
    try:
        settings.mode = "Normal"
        settings.enable_cache = False
        settings.auto_ai_review_after_scan = True
        settings.smart_review_during_scan = True
    except Exception:
        pass

    engine = ScanEngine(db, settings, log=hlog)
    result = engine.scan_file(Path(file_path), force_rescan=True)
    payload: Dict[str, Any] = {
        "ok": True,
        "file_path": str(file_path),
        "rows": [m.as_dict() if hasattr(m, "as_dict") else dataclasses.asdict(m) for m in (result.materials or [])],
        "merged_text": result.merged_text or "",
        "native_text": result.native_text or "",
        "ocr_text": result.ocr_text or "",
        "table_text": result.table_text or "",
        "confidence": result.confidence,
        "warnings": list(result.warnings or []),
        "engine_summary": result.engine_summary or "",
        "elapsed_sec": result.elapsed_sec,
        "logs": logs[-200:],
    }

    if json_out:
        Path(json_out).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return payload


def main() -> None:
    # MaINbox integration: run SmartScan in a true headless subprocess.  This
    # makes MaINbox use the same SmartScan engine that works standalone, without
    # importing the SmartScan GUI module into MaINbox's process.
    if "--mainbox-headless" in sys.argv:
        try:
            idx = sys.argv.index("--mainbox-headless")
            file_path = sys.argv[idx + 1]
            json_out = ""
            if "--json-out" in sys.argv:
                jidx = sys.argv.index("--json-out")
                json_out = sys.argv[jidx + 1]
            payload = run_mainbox_headless_scan(file_path, json_out=json_out)
            if not json_out:
                print(json.dumps(payload, ensure_ascii=False))
            return
        except Exception as exc:
            err_payload = {
                "ok": False,
                "error": str(exc),
                "traceback": traceback.format_exc(),
            }
            if "--json-out" in sys.argv:
                try:
                    jidx = sys.argv.index("--json-out")
                    Path(sys.argv[jidx + 1]).write_text(json.dumps(err_payload, ensure_ascii=False, indent=2), encoding="utf-8")
                except Exception:
                    pass
            print(json.dumps(err_payload, ensure_ascii=False))
            raise SystemExit(2)

    if "--mainbox-review" in sys.argv:
        try:
            ridx = sys.argv.index("--mainbox-review")
            json_out = ""
            if "--json-out" in sys.argv:
                jidx = sys.argv.index("--json-out")
                json_out = sys.argv[jidx + 1]
            stop_flags = {"--json-out"}
            files: List[str] = []
            i = ridx + 1
            while i < len(sys.argv):
                arg = sys.argv[i]
                if arg == "--json-out":
                    i += 2
                    continue
                if arg.startswith("--"):
                    i += 1
                    continue
                files.append(arg)
                i += 1
            app = SmartScanApp(mainbox_files=files, mainbox_json_out=json_out, mainbox_review_mode=True)
            app.run()
            return
        except Exception as exc:
            err_payload = {"ok": False, "error": str(exc), "traceback": traceback.format_exc()}
            if "--json-out" in sys.argv:
                try:
                    jidx = sys.argv.index("--json-out")
                    Path(sys.argv[jidx + 1]).write_text(json.dumps(err_payload, ensure_ascii=False, indent=2), encoding="utf-8")
                except Exception:
                    pass
            print(json.dumps(err_payload, ensure_ascii=False))
            raise SystemExit(2)

    app = SmartScanApp()
    app.run()


if __name__ == "__main__":
    main()
